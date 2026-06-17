from __future__ import annotations

from copy import deepcopy
from datetime import datetime
from pathlib import Path
import re
import unicodedata
from typing import Any, Iterable, Mapping

from docx import Document
from docx.table import Table
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


BASE_DIR = Path(__file__).resolve().parent
TEMPLATE_PLANO_PATH = BASE_DIR / "modelos" / "modelo_plano_acao.docx"
TEMPLATE_PGR_PATH = BASE_DIR / "modelos" / "modelo_risco_pgr.docx"
TEMPLATE_PCMSO_PATH = BASE_DIR / "modelos" / "modelo_pcmso.docx"
TEMPLATE_RELACAO_PATH = BASE_DIR / "modelos" / "modelo_relacao_funcao_atividade.docx"
TEMPLATE_DESCRITIVO_SETOR_PATH = BASE_DIR / "modelos" / "modelo_descritivo_setor.docx"
TEMPLATE_PGR_COMPLETO_PATH = BASE_DIR / "modelos" / "modelo_pgr_completo.docx"
TEMPLATE_PCMSO_COMPLETO_PATH = BASE_DIR / "modelos" / "modelo_pcmso_completo.docx"
TEMPLATE_RISCOS_PCMSO_PATH = BASE_DIR / "modelos" / "modelo_riscos_pcmso.docx"
TEMPLATE_LTCAT_COMPLETO_PATH = BASE_DIR / "modelos" / "modelo_ltcat_completo.docx"
TEMPLATE_AET_BASE_PATH = BASE_DIR / "modelos" / "modelo_aet_base.docx"
DEFAULT_DOCX_FONT = "Arial Narrow"

TIPO_RISCO_COLORS = {
    "ACIDENTE(MECÂNICO)": "0066FF",
    "ERGONÔMICO": "FFFF00",
    "ERGONÔMICO PSICOSSOCIAL": "FFFF00",
    "FÍSICO": "009933",
    "BIOLÓGICO": "E36C0A",
    "QUÍMICO": "EE0000",
}

SEVERIDADE_COLORS = {
    "INSIGNIFICANTE": "009900",
    "PEQUENO": "009900",
    "MÉDIO": "FFFF00",
    "GRANDE": "336699",
    "CATASTRÓFICO": "FF0000",
}

POSSIBILIDADE_COLORS = {
    "IMPROVÁVEL": "009900",
    "POSSÍVEL": "336699",
    "PROVÁVEL": "FF0000",
}

NIVEL_RISCO_COLORS = {
    "TRIVIAL": "009900",
    "BAIXO": "FFFF00",
    "MODERADO": "336699",
    "ALTO": "990066",
    "MUITO ALTO": "FF0000",
}

AET_ACTIVITY_PROFILES: list[dict[str, Any]] = [
    {
        "label": "Serviços funerários",
        "keywords": ["96.03", "9603", "funer", "somatoconservacao", "somatoconservação", "assistencia postuma", "assistência póstuma"],
        "contexto": "atividade de serviços funerários, com rotinas administrativas, atendimento a familiares, apoio operacional, higienização, transporte e organização de serviços, exigindo controle técnico, postura profissional e equilíbrio emocional diante de situações sensíveis.",
        "demanda": "aprofundar a avaliação das exigências físicas, cognitivas, organizacionais e psicossociais associadas ao atendimento humanizado, à exposição a situações de luto, à organização de demandas variáveis e ao suporte operacional.",
        "tarefas": "atendimento e orientação de clientes/familiares, organização documental, transporte, apoio operacional, preparação de ambientes, higienização e suporte às rotinas internas da empresa.",
        "exigencias": "exigência emocional elevada, atenção contínua, comunicação cuidadosa, controle de prazos, deslocamentos, alternância postural e necessidade de tomada de decisão compatível com a sensibilidade dos serviços prestados.",
        "favoraveis": "estrutura formal definida, possibilidade de padronização dos atendimentos, divisão de funções por setor e integração das ações de saúde ocupacional ao PGR.",
        "recomendacoes": "padronizar fluxos de atendimento, orientar lideranças e equipes sobre comunicação, pausas, suporte emocional, organização das demandas e prevenção de conflitos.",
    },
    {
        "label": "Comércio varejista e atendimento",
        "keywords": ["47.", "comercio", "comércio", "varejista", "loja", "vestuario", "vestuário", "mercado", "supermercado", "atacado", "atacarejo"],
        "contexto": "atividade comercial com atendimento ao público, organização de mercadorias, operação de caixa, reposição, estoque e suporte administrativo, sujeita a variações de fluxo, metas e demandas de clientes.",
        "demanda": "avaliar as condições ergonômicas relacionadas à postura em pé, deslocamentos, repetitividade, atendimento ao público, organização de mercadorias e pressão por produtividade.",
        "tarefas": "vendas, atendimento, operação de caixa, reposição, organização de prateleiras, controle de estoque, conferência de mercadorias e apoio administrativo.",
        "exigencias": "permanência em pé, deslocamentos frequentes, atenção ao cliente, controle de prioridades, manuseio eventual de cargas e comunicação constante com equipe e consumidores.",
        "favoraveis": "atividades com rotinas conhecidas, possibilidade de rodízio, padronização de tarefas e implantação de pausas e orientações ergonômicas.",
        "recomendacoes": "promover alternância postural, pausas breves, organização do posto, rodízio de tarefas quando aplicável e orientação de NR-17 para trabalhadores e liderança.",
    },
    {
        "label": "Administrativo e escritório",
        "keywords": ["82.", "69.", "70.", "administrativo", "escritorio", "escritório", "advocacia", "contabilidade", "consultoria"],
        "contexto": "atividade predominantemente administrativa, com uso de computador, atendimento interno/externo, organização documental, controle de informações e cumprimento de prazos.",
        "demanda": "avaliar postos informatizados, postura sentada, movimentos repetitivos, carga cognitiva, organização de prioridades, comunicação e autonomia na execução das atividades.",
        "tarefas": "digitação, registros em sistema, atendimento, análise de documentos, controle de planilhas, arquivamento, comunicação interna e suporte à gestão.",
        "exigencias": "atenção contínua, permanência sentada, uso de tela, teclado e mouse, prazos, controle de informações e comunicação com diferentes áreas.",
        "favoraveis": "ambiente controlado, baixo esforço físico predominante e maior facilidade para ajustes de mobiliário, pausas e organização do trabalho.",
        "recomendacoes": "adequar cadeira, mesa e monitor, instituir pausas visuais, alternância postural, organização de demandas e orientação ergonômica.",
    },
    {
        "label": "Limpeza, conservação e serviços gerais",
        "keywords": ["81.21", "81.22", "81.29", "limpeza", "conservacao", "conservação", "servicos gerais", "serviços gerais"],
        "contexto": "atividade de limpeza e conservação, com deslocamentos, uso de ferramentas manuais, manuseio de materiais e execução de rotinas em diferentes ambientes.",
        "demanda": "avaliar posturas incômodas, esforço físico, deslocamentos, ritmo de trabalho, organização das rotinas e condições de ferramentas/equipamentos.",
        "tarefas": "limpeza, higienização, conservação de ambientes, recolhimento de resíduos, organização de materiais e apoio às áreas comuns.",
        "exigencias": "postura em pé, flexão de tronco, agachamentos eventuais, deslocamento frequente, uso de força moderada e atenção às condições do piso e circulação.",
        "favoraveis": "possibilidade de padronização da rotina, uso de ferramentas adequadas e implantação de pausas e alternância de tarefas.",
        "recomendacoes": "adequar ferramentas, organizar carrinhos/materiais, orientar sobre postura, pausas, alternância de tarefas e comunicação com a liderança.",
    },
    {
        "label": "Condomínio, portaria e serviços prediais",
        "keywords": ["81.12", "8112", "condominio", "condomínio", "portaria", "porteiro", "zelador", "predial"],
        "contexto": "atividade predial com controle de acesso, atendimento a moradores/visitantes, monitoramento, limpeza, manutenção e suporte administrativo.",
        "demanda": "avaliar atenção contínua, postura sentada/em pé, comunicação, controle de conflitos, organização das escalas e condições dos postos de trabalho.",
        "tarefas": "controle de entrada e saída, atendimento, monitoramento, limpeza, conservação, pequenos apoios operacionais e comunicação com administração.",
        "exigencias": "atenção permanente, comunicação com público, permanência em posto, alternância postural, rotina por escala e resposta a demandas variáveis.",
        "favoraveis": "tarefas passíveis de padronização, definição de procedimentos e ajustes pontuais de posto.",
        "recomendacoes": "adequar posto de portaria, orientar comunicação, definir procedimentos, garantir pausas conforme escala e acompanhar riscos psicossociais.",
    },
    {
        "label": "Alimentação, restaurante e cozinha",
        "keywords": ["56.", "alimentacao", "alimentação", "restaurante", "lanchonete", "cozinha", "refeicao", "refeição"],
        "contexto": "atividade de preparo, atendimento e serviço de alimentação, com fluxo variável, permanência em pé, calor, manipulação de utensílios e organização de pedidos.",
        "demanda": "avaliar postura em pé, repetitividade, calor, ritmo de produção, organização do posto, manuseio de materiais e atendimento ao público.",
        "tarefas": "preparo de alimentos, montagem, atendimento, caixa, higienização, armazenamento e organização de utensílios e insumos.",
        "exigencias": "permanência em pé, deslocamentos, atenção ao preparo, ritmo por demanda, uso de utensílios, calor e organização de prioridades.",
        "favoraveis": "rotinas operacionais definidas, possibilidade de organização de bancadas, pausas e alternância de atividades.",
        "recomendacoes": "adequar bancadas, melhorar organização do fluxo, orientar postura, pausas, hidratação e alternância de tarefas.",
    },
    {
        "label": "Saúde, clínica, consultório e laboratório",
        "keywords": ["86.", "clinica", "clínica", "consultorio", "consultório", "laboratorio", "laboratório", "saude", "saúde"],
        "contexto": "atividade de saúde com atendimento a pacientes, recepção, apoio técnico, procedimentos, registros e organização de documentos e exames.",
        "demanda": "avaliar exigências cognitivas, atendimento ao público, postura, movimentos repetitivos, organização de agendas, comunicação e possível exposição emocional.",
        "tarefas": "recepção, cadastro, atendimento, apoio a procedimentos, coleta/exames, limpeza, registros em sistema e organização documental.",
        "exigencias": "atenção contínua, comunicação clara, permanência sentada/em pé, precisão nos registros, controle de agenda e interação frequente com pacientes.",
        "favoraveis": "ambientes com rotinas técnicas definidas e possibilidade de padronização de fluxos, mobiliário e pausas.",
        "recomendacoes": "adequar mobiliário, organizar fluxo de atendimento, promover pausas, comunicação clara e orientação ergonômica.",
    },
    {
        "label": "Oficina, manutenção e serviços técnicos",
        "keywords": ["45.20", "4520", "33.", "oficina", "mecanica", "mecânica", "manutencao", "manutenção", "refrigeracao", "refrigeração"],
        "contexto": "atividade técnica de manutenção, inspeção, reparo e organização de peças/equipamentos, com uso de ferramentas e posturas variadas.",
        "demanda": "avaliar posturas incômodas, esforço físico, ferramentas, levantamento de cargas, atenção técnica, organização do posto e pausas.",
        "tarefas": "inspeção, manutenção, reparo, montagem/desmontagem, uso de ferramentas, organização de peças e apoio administrativo.",
        "exigencias": "posturas em flexão, agachamento eventual, permanência em pé, atenção técnica, manuseio de ferramentas e demandas por ordem de serviço.",
        "favoraveis": "possibilidade de organização de bancadas, ferramentas adequadas e padronização de procedimentos.",
        "recomendacoes": "organizar ferramentas, ajustar bancadas, orientar postura, alternância de tarefas, pausas e uso de meios auxiliares.",
    },
    {
        "label": "Construção civil e obras",
        "keywords": ["41.", "42.", "43.", "construcao", "construção", "obra", "edificacao", "edificação", "engenharia"],
        "contexto": "atividade de obra e apoio operacional, com exigência física relevante, deslocamentos, frentes de trabalho, ferramentas e variação ambiental.",
        "demanda": "avaliar esforço físico, posturas incômodas, levantamento de cargas, ritmo de produção, organização das frentes de serviço e áreas de descanso.",
        "tarefas": "execução de serviços de obra, apoio operacional, transporte de materiais, organização de ferramentas, almoxarifado e acompanhamento técnico.",
        "exigencias": "elevada exigência física, atenção à segurança, deslocamentos, manuseio de materiais, calor/ruído/poeira e coordenação de equipes.",
        "favoraveis": "possibilidade de planejamento de frentes, pausas, rodízios e uso de equipamentos auxiliares.",
        "recomendacoes": "planejar tarefas, usar meios auxiliares, orientar levantamento de cargas, pausas, hidratação e organização do canteiro.",
    },
    {
        "label": "Transporte, logística e táxi",
        "keywords": ["49.", "52.", "transporte", "logistica", "logística", "taxi", "táxi", "motorista", "garagem", "deposito", "depósito", "entrega"],
        "contexto": "atividade de transporte, deslocamento, logística ou armazenagem, com condução, atenção contínua, prazos, movimentação e comunicação operacional.",
        "demanda": "avaliar postura sentada prolongada, atenção, fadiga, vibração, ritmo de rotas, movimentação de materiais e organização das pausas.",
        "tarefas": "condução de veículos, atendimento/entrega, carga e descarga, conferência, expedição, organização de depósito e suporte administrativo.",
        "exigencias": "atenção contínua, postura sentada ou em pé, deslocamentos, prazos, interação com trânsito/clientes e possível manuseio de cargas.",
        "favoraveis": "rotas e processos passíveis de planejamento, pausas e orientações de ergonomia e direção segura.",
        "recomendacoes": "organizar pausas, orientar postura ao dirigir, controlar fadiga, planejar rotas e melhorar comunicação operacional.",
    },
    {
        "label": "Educação e ensino",
        "keywords": ["85.", "educacao", "educação", "escola", "ensino", "professor", "creche", "curso"],
        "contexto": "atividade educacional com interação contínua, planejamento pedagógico, comunicação, permanência em pé/sentado e exigência cognitiva.",
        "demanda": "avaliar carga cognitiva, comunicação, organização de aulas, postura, voz, demandas emocionais e apoio às rotinas escolares.",
        "tarefas": "aulas, atendimento a alunos/responsáveis, planejamento, registros, coordenação, limpeza e apoio escolar.",
        "exigencias": "atenção contínua, comunicação verbal, controle de turma, prazos, postura alternada e interação social intensa.",
        "favoraveis": "rotinas definidas por calendário e possibilidade de planejamento, pausas e organização do posto.",
        "recomendacoes": "organizar pausas, alternância postural, suporte à comunicação, planejamento de demandas e acompanhamento psicossocial quando aplicável.",
    },
    {
        "label": "Agropecuária e atividade rural",
        "keywords": ["01.", "agro", "agricultura", "pecuaria", "pecuária", "rural", "fazenda", "plantio"],
        "contexto": "atividade rural ou agropecuária com tarefas operacionais, deslocamentos, variação climática, manuseio de materiais/equipamentos e organização de rotinas de campo.",
        "demanda": "avaliar esforço físico, posturas incômodas, jornada, calor, ferramentas, deslocamentos e organização das frentes de trabalho.",
        "tarefas": "manejo, plantio, apoio operacional, manutenção, organização de insumos, atividades administrativas e controle de produção.",
        "exigencias": "esforço físico, deslocamentos, variação ambiental, atenção a equipamentos, manuseio de cargas e demandas sazonais.",
        "favoraveis": "possibilidade de planejamento de tarefas, pausas, ferramentas adequadas e capacitação operacional.",
        "recomendacoes": "planejar tarefas, orientar postura, pausas, hidratação, ferramentas adequadas e rodízio quando aplicável.",
    },
]


# ---------------------------------------------------------------------------
# Utilitários gerais de OOXML
# ---------------------------------------------------------------------------

def _normalize_option(value: Any) -> str:
    return str(value or "").strip().upper()


def _set_run_properties_font(run_properties, font_name: str = DEFAULT_DOCX_FONT, size_pt: float | None = None) -> None:
    """Força a fonte no OOXML, inclusive em textos inseridos diretamente nas tabelas."""
    if run_properties is None:
        return
    r_fonts = run_properties.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        run_properties.insert(0, r_fonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        r_fonts.set(qn(f"w:{attr}"), font_name)
    if size_pt is not None:
        size_val = str(int(float(size_pt) * 2))
        for tag in ("w:sz", "w:szCs"):
            size_el = run_properties.find(qn(tag))
            if size_el is None:
                size_el = OxmlElement(tag)
                run_properties.append(size_el)
            size_el.set(qn("w:val"), size_val)


def _apply_run_font(run, font_name: str = DEFAULT_DOCX_FONT) -> None:
    """Aplica a fonte padrão preservando tamanho, negrito, cor e demais estilos."""
    try:
        run.font.name = font_name
        r_pr = run._element.get_or_add_rPr()
        _set_run_properties_font(r_pr, font_name)
    except Exception:
        pass


def _apply_part_font(part, font_name: str = DEFAULT_DOCX_FONT) -> None:
    for paragraph in getattr(part, "paragraphs", []) or []:
        for run in paragraph.runs:
            _apply_run_font(run, font_name)
    for table in getattr(part, "tables", []) or []:
        for row in table.rows:
            for cell in row.cells:
                _apply_part_font(cell, font_name)


def apply_default_docx_font(doc: Document, font_name: str = DEFAULT_DOCX_FONT) -> Document:
    """Padroniza todos os textos editáveis do DOCX gerado em Arial Narrow."""
    for style in doc.styles:
        try:
            style.font.name = font_name
        except Exception:
            pass
    _apply_part_font(doc, font_name)
    for section in doc.sections:
        for part in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ):
            _apply_part_font(part, font_name)
    return doc


def save_docx_with_default_font(doc: Document, output_path: str | Path, font_name: str = DEFAULT_DOCX_FONT) -> None:
    apply_default_docx_font(doc, font_name)
    doc.save(str(output_path))


def _clean_cargo_and_cbo(raw_cargo: Any, raw_cbo: Any = "") -> tuple[str, str]:
    """Separa cargo e CBO quando laudos importados trouxerem CBO duplicado no cargo."""
    raw_cargo = re.sub(r"\s+", " ", str(raw_cargo or "").strip())
    raw_cbo = re.sub(r"\s+", " ", str(raw_cbo or "").strip())
    combined = f"{raw_cargo} {raw_cbo}"
    cbo_matches = re.findall(r"\b\d{4}-\d{2}\b", combined)
    cbo = cbo_matches[-1] if cbo_matches else ""
    cargo = re.split(r"\s*[–—-]\s*CBO\s*:|\bCBO\s*:", raw_cargo, maxsplit=1, flags=re.I)[0]
    cargo = re.sub(r"\s*[–—-]\s*$", "", cargo).strip()
    cargo = re.sub(r"\s+", " ", cargo)
    if not cbo:
        m = re.search(r"\b\d{4}-\d{2}\b", raw_cbo)
        cbo = m.group(0) if m else "A DEFINIR"
    if not cargo:
        cargo = "A DEFINIR"
    return cargo, cbo


def _make_text(value: Any) -> list:
    """Cria nós OOXML de texto, preservando quebras de linha."""
    value = "" if value is None else str(value)
    lines = value.splitlines() or [""]
    nodes = []
    for index, line in enumerate(lines):
        if index > 0:
            nodes.append(OxmlElement("w:br"))
        text_node = OxmlElement("w:t")
        if line != line.strip():
            text_node.set(qn("xml:space"), "preserve")
        text_node.text = line
        nodes.append(text_node)
    return nodes


def _set_font_color(run_properties, color_hex: str | None) -> None:
    if not color_hex:
        return
    existing = run_properties.find(qn("w:color"))
    if existing is None:
        existing = OxmlElement("w:color")
        run_properties.append(existing)
    existing.set(qn("w:val"), color_hex.replace("#", "").upper())


def _set_cell_text(cell_xml, value: Any, font_color: str | None = "000000") -> None:
    """Substitui o texto visível da célula preservando estilo básico."""
    old_paragraphs = cell_xml.findall(qn("w:p"))
    first_paragraph = old_paragraphs[0] if old_paragraphs else None
    paragraph_properties = None
    run_properties = None

    if first_paragraph is not None:
        p_pr = first_paragraph.find(qn("w:pPr"))
        if p_pr is not None:
            paragraph_properties = deepcopy(p_pr)
        first_run = first_paragraph.find(qn("w:r"))
        if first_run is not None:
            r_pr = first_run.find(qn("w:rPr"))
            if r_pr is not None:
                run_properties = deepcopy(r_pr)

    if run_properties is None:
        run_properties = OxmlElement("w:rPr")
    _set_run_properties_font(run_properties)
    _set_font_color(run_properties, font_color)

    for child in list(cell_xml):
        if child.tag != qn("w:tcPr"):
            cell_xml.remove(child)

    paragraph = OxmlElement("w:p")
    if paragraph_properties is not None:
        paragraph.append(paragraph_properties)
    run = OxmlElement("w:r")
    run.append(run_properties)
    for node in _make_text(value):
        run.append(node)
    paragraph.append(run)
    cell_xml.append(paragraph)



def _set_cell_text_with_font(
    cell_xml,
    value: Any,
    font_name: str = "Arial Narrow",
    font_size_pt: int = 10,
    font_color: str | None = "000000",
    bold: bool = False,
    align: str | None = None,
    vertical_align: str | None = None,
) -> None:
    """Substitui o texto da célula e força fonte/tamanho/alinhamento no OOXML.

    Usado especialmente na coluna GES do Plano de Ação, porque no PGR completo
    a tabela é manipulada por XML puro e não passa pelo loop final do python-docx.
    """
    old_paragraphs = cell_xml.findall(qn("w:p"))
    first_paragraph = old_paragraphs[0] if old_paragraphs else None
    paragraph_properties = None
    run_properties = None

    if first_paragraph is not None:
        p_pr = first_paragraph.find(qn("w:pPr"))
        if p_pr is not None:
            paragraph_properties = deepcopy(p_pr)
        first_run = first_paragraph.find(qn("w:r"))
        if first_run is not None:
            r_pr = first_run.find(qn("w:rPr"))
            if r_pr is not None:
                run_properties = deepcopy(r_pr)

    if paragraph_properties is None:
        paragraph_properties = OxmlElement("w:pPr")
    if align:
        jc = paragraph_properties.find(qn("w:jc"))
        if jc is None:
            jc = OxmlElement("w:jc")
            paragraph_properties.append(jc)
        jc.set(qn("w:val"), align)

    if run_properties is None:
        run_properties = OxmlElement("w:rPr")

    # Remove formatação antiga que possa vir do modelo para não ficar com
    # duas marcações conflitantes de fonte/tamanho no Word.
    for tag in ("w:rFonts", "w:sz", "w:szCs", "w:b", "w:bCs"):
        for el in list(run_properties.findall(qn(tag))):
            run_properties.remove(el)

    r_fonts = OxmlElement("w:rFonts")
    run_properties.append(r_fonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        r_fonts.set(qn(f"w:{attr}"), font_name)

    size_val = str(int(font_size_pt * 2))
    sz = OxmlElement("w:sz")
    run_properties.append(sz)
    sz.set(qn("w:val"), size_val)
    sz_cs = OxmlElement("w:szCs")
    run_properties.append(sz_cs)
    sz_cs.set(qn("w:val"), size_val)

    if bold:
        run_properties.append(OxmlElement("w:b"))
    _set_font_color(run_properties, font_color)

    tc_pr = cell_xml.find(qn("w:tcPr"))
    if vertical_align:
        if tc_pr is None:
            tc_pr = OxmlElement("w:tcPr")
            cell_xml.insert(0, tc_pr)
        v_align = tc_pr.find(qn("w:vAlign"))
        if v_align is None:
            v_align = OxmlElement("w:vAlign")
            tc_pr.append(v_align)
        v_align.set(qn("w:val"), vertical_align)

    for child in list(cell_xml):
        if child.tag != qn("w:tcPr"):
            cell_xml.remove(child)

    paragraph = OxmlElement("w:p")
    paragraph.append(paragraph_properties)
    run = OxmlElement("w:r")
    run.append(run_properties)
    for node in _make_text(value):
        run.append(node)
    paragraph.append(run)
    cell_xml.append(paragraph)

def _set_cell_shading(cell_xml, fill_hex: str | None) -> None:
    if not fill_hex:
        return
    fill_hex = fill_hex.replace("#", "").upper()
    tc_pr = cell_xml.find(qn("w:tcPr"))
    if tc_pr is None:
        tc_pr = OxmlElement("w:tcPr")
        cell_xml.insert(0, tc_pr)
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)


def _value_cell(rows, row_index: int, cell_index: int):
    return rows[row_index].findall(qn("w:tc"))[cell_index]


def _row_cell(row_xml, cell_index: int):
    return row_xml.findall(qn("w:tc"))[cell_index]


def _colored_value(cell_xml, value: Any, color_map: Mapping[str, str]) -> None:
    option = _normalize_option(value)
    fill = color_map.get(option)
    _set_cell_shading(cell_xml, fill)
    # Todas as letras devem permanecer pretas, inclusive nas células coloridas.
    _set_cell_text(cell_xml, option, font_color="000000")


def _insert_before_section_or_end(body, element) -> None:
    children = list(body)
    for index, child in enumerate(children):
        if child.tag == qn("w:sectPr"):
            body.insert(index, element)
            return
    body.append(element)


def _blank_paragraph() -> OxmlElement:
    paragraph = OxmlElement("w:p")
    p_pr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    p_pr.append(spacing)
    paragraph.append(p_pr)
    return paragraph


def _page_break_paragraph() -> OxmlElement:
    paragraph = OxmlElement("w:p")
    p_pr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    p_pr.append(spacing)
    paragraph.append(p_pr)
    run = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run.append(br)
    paragraph.append(run)
    return paragraph



def _set_row_cant_split(row_xml) -> None:
    """Evita que uma linha de tabela seja quebrada entre páginas."""
    tr_pr = row_xml.find(qn("w:trPr"))
    if tr_pr is None:
        tr_pr = OxmlElement("w:trPr")
        row_xml.insert(0, tr_pr)
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def _set_row_min_height(row_xml, minimum_twips: int) -> None:
    """Define altura mínima da linha, permitindo expansão se o conteúdo precisar."""
    tr_pr = row_xml.find(qn("w:trPr"))
    if tr_pr is None:
        tr_pr = OxmlElement("w:trPr")
        row_xml.insert(0, tr_pr)
    height = tr_pr.find(qn("w:trHeight"))
    if height is None:
        height = OxmlElement("w:trHeight")
        tr_pr.append(height)
    old_val = height.get(qn("w:val"))
    try:
        old_int = int(old_val or 0)
    except ValueError:
        old_int = 0
    height.set(qn("w:val"), str(max(old_int, minimum_twips)))
    height.set(qn("w:hRule"), "atLeast")


def _insert_page_break_before_previous_paragraph(element, text_marker: str) -> None:
    """Insere quebra de página antes do parágrafo anterior que contenha o marcador."""
    parent = element.getparent()
    if parent is None:
        return
    children = list(parent)
    try:
        index = children.index(element)
    except ValueError:
        return
    normalized_marker = text_marker.strip().upper()
    for prev_index in range(index - 1, -1, -1):
        candidate = children[prev_index]
        if candidate.tag != qn("w:p"):
            continue
        text = "".join(node.text or "" for node in candidate.iter(qn("w:t"))).strip().upper()
        if normalized_marker in text:
            # Evita inserir duas vezes caso o modelo já tenha uma quebra logo antes.
            if prev_index > 0 and children[prev_index - 1].tag == qn("w:p"):
                has_break = children[prev_index - 1].find(f".//{qn('w:br')}") is not None
                if has_break:
                    return
            parent.insert(prev_index, _page_break_paragraph())
            return

# ---------------------------------------------------------------------------
# PLANO DE AÇÃO
# ---------------------------------------------------------------------------

def _fill_plano_row(row_xml, risk: Mapping[str, Any], setor: str = "", data_atual: str = "", data_final: str = "") -> None:
    cells = row_xml.findall(qn("w:tc"))
    if len(cells) >= 7:
        setor_lines = [line for line in str(setor or "").splitlines() if line.strip()]
        _set_row_cant_split(row_xml)
        _set_row_min_height(row_xml, max(1025, 420 + (max(1, len(setor_lines)) * 260)))
        # Novo modelo: GES recebe o setor.
        # Datas do plano de ação:
        # - regra geral: prazo de implantação = Data atual/início da vigência;
        #                prazo de reavaliação = Data final da vigência.
        # - risco ERGONÔMICO PSICOSSOCIAL: implantação em 30 DIAS e reavaliação em 180 DIAS.
        is_psychosocial = _normalize_option(risk.get("tipo_risco")) == "ERGONÔMICO PSICOSSOCIAL"
        prazo_implantacao = "30 DIAS" if is_psychosocial else (data_atual or "")
        prazo_reavaliacao = "180 DIAS" if is_psychosocial else (data_final or "")

        _set_cell_text_with_font(cells[0], setor, font_name="Arial Narrow", font_size_pt=10, font_color="000000", bold=False, align="center", vertical_align="center")
        _set_cell_text(cells[1], risk.get("risco", ""))
        _set_cell_text(cells[2], risk.get("acoes", ""))
        # cells[3] mantém o responsável fixo do modelo: ADMINISTRAÇÃO.
        # Alguns modelos do PGR completo têm uma célula extra/mesclada antes dos prazos.
        if len(cells) >= 8:
            _set_cell_text(cells[5], prazo_implantacao)
            _set_cell_text(cells[6], prazo_reavaliacao)
            _set_cell_text(cells[7], risk.get("indicador", ""))
        else:
            _set_cell_text(cells[4], prazo_implantacao)
            _set_cell_text(cells[5], prazo_reavaliacao)
            _set_cell_text(cells[6], risk.get("indicador", ""))
        return

    replacements = {
        "{{SETOR}}": setor,
        "{{risco}}": risk.get("risco", ""),
        "{{AÇÕES PREVENTIVA / CORRETIVA}}": risk.get("acoes", ""),
        "{{INDICADOR DE EFETIVIDADE}}": risk.get("indicador", ""),
    }
    for node in list(row_xml.iter(qn("w:t"))):
        text = node.text or ""
        if "{{mês/ano}}" in text:
            node.text = text.replace("{{mês/ano}}", "")
        else:
            for placeholder, value in replacements.items():
                if placeholder in text:
                    node.text = text.replace(placeholder, str(value or ""))
                    break


def _plano_entries_from_groups_or_risks(items: list[Mapping[str, Any]]) -> list[tuple[str, Mapping[str, Any]]]:
    """Converte a seleção em linhas de Plano de Ação: (GES, risco).

    Quando o mesmo risco aparece em mais de um setor, o plano de ação deve
    gerar uma única linha para o risco e colocar todos os setores na coluna GES.
    Isso deixa o PGR mais limpo e evita várias linhas repetidas mudando apenas
    o setor.
    """
    if _is_grouped_payload(items):
        grouped: dict[str, dict[str, Any]] = {}
        order: list[str] = []
        for group in _sanitize_sector_risk_groups(items):
            setor = str(group["sector"].get("setor", "")).strip()
            for risk in group.get("risks", []):
                risk_key = str(risk.get("id") or risk.get("risco") or "").strip().lower()
                if not risk_key:
                    continue
                if risk_key not in grouped:
                    grouped[risk_key] = {"risk": risk, "setores": []}
                    order.append(risk_key)
                if setor and setor not in grouped[risk_key]["setores"]:
                    grouped[risk_key]["setores"].append(setor)
        entries: list[tuple[str, Mapping[str, Any]]] = []
        for key in order:
            setores = sorted(grouped[key]["setores"], key=lambda value: value.lower())
            ges = "\n".join(setores)
            entries.append((ges, grouped[key]["risk"]))
        return entries
    return [("", risk) for risk in items]


def _find_template_row(table_xml, marker: str = "{{risco}}"):
    for row in table_xml.tr_lst:
        text = "".join(node.text or "" for node in row.iter(qn("w:t")))
        if marker in text:
            return row
    return None


def generate_action_plan_docx(groups_or_risks: Iterable[Mapping[str, Any]], output_path: str | Path, data_atual: str | None = None, data_final: str | None = None) -> Path:
    items = list(groups_or_risks)
    data_atual = (data_atual or "").strip()
    data_final = (data_final or "").strip()
    entries = _plano_entries_from_groups_or_risks(items)
    if not entries:
        raise ValueError("Selecione pelo menos um setor e um risco para gerar o Plano de Ação.")
    if not TEMPLATE_PLANO_PATH.exists():
        raise FileNotFoundError(f"Modelo Word não encontrado: {TEMPLATE_PLANO_PATH}")

    doc = Document(str(TEMPLATE_PLANO_PATH))
    if not doc.tables:
        raise ValueError("O modelo do Plano de Ação precisa ter uma tabela com a linha-modelo.")

    table = doc.tables[0]
    template_row_xml = _find_template_row(table._tbl)
    if template_row_xml is None:
        raise ValueError("O modelo do Plano de Ação precisa ter uma linha-modelo com {{risco}}.")
    template_row_copy = deepcopy(template_row_xml)
    table._tbl.remove(template_row_xml)

    # Mantém a primeira linha (PLANO DE AÇÃO) e a segunda linha (cabeçalho) apenas uma vez.
    # Depois, insere todos os setores/riscos um abaixo do outro, sem espaçamento.
    for setor, risk in entries:
        new_row = deepcopy(template_row_copy)
        _fill_plano_row(new_row, risk, setor=setor, data_atual=data_atual, data_final=data_final)
        table._tbl.append(new_row)

    # Padroniza a visualização da coluna GES e demais células do plano.
    # Os setores agrupados ficam um abaixo do outro, em Arial Narrow, sem ocupar várias linhas repetidas.
    for row_index, row in enumerate(table.rows):
        for cell_index, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Arial Narrow"
                    # Na coluna GES, quando o mesmo risco está em vários setores,
                    # cada setor fica em uma linha, com leitura melhor no Word.
                    run.font.size = Pt(10 if cell_index == 0 and row_index > 1 else 7)
                    if cell_index == 0 and row_index > 1:
                        run.bold = False

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    save_docx_with_default_font(doc, output_path)
    return output_path


# Compatibilidade com a primeira versão do site.
generate_docx_from_risks = generate_action_plan_docx


# ---------------------------------------------------------------------------
# RISCO PGR POR SETOR
# ---------------------------------------------------------------------------

def _sanitize_sector(sector: Mapping[str, Any]) -> dict[str, Any]:
    cargos = sector.get("cargos") or []
    clean_cargos = []
    for cargo in cargos:
        cargo_nome, cbo_limpo = _clean_cargo_and_cbo(cargo.get("cargo", ""), cargo.get("cbo", ""))
        clean_cargos.append(
            {
                "cargo": cargo_nome,
                "cbo": cbo_limpo,
                "n_func": str(cargo.get("n_func", "")).strip(),
                "descricao": str(cargo.get("descricao", "")).strip(),
            }
        )
    return {"setor": str(sector.get("setor", "")).strip(), "cargos": clean_cargos}


def _sector_cargos_text(sector: Mapping[str, Any]) -> str:
    cargos = sector.get("cargos") or []
    names = [str(cargo.get("cargo", "")).strip() for cargo in cargos if str(cargo.get("cargo", "")).strip()]
    return ", ".join(names)


def _is_grouped_payload(items: list[Mapping[str, Any]]) -> bool:
    return bool(items) and isinstance(items[0], Mapping) and "risks" in items[0]


def _sanitize_sector_risk_groups(groups: Iterable[Mapping[str, Any]]) -> list[dict[str, Any]]:
    clean_groups: list[dict[str, Any]] = []
    for group in groups:
        sector = _sanitize_sector(group.get("sector") or {})
        risks = [risk for risk in group.get("risks", []) if risk]
        if sector.get("setor") and risks:
            clean_groups.append({"sector": sector, "risks": risks})
    return clean_groups


def _fill_pgr_risk_rows(risk_rows: list, risk: Mapping[str, Any]) -> None:
    if len(risk_rows) < 11:
        raise ValueError("O bloco de risco do PGR precisa manter 11 linhas.")

    tipo = risk.get("tipo_risco", "")
    descricao = risk.get("descricao_agente") or risk.get("risco", "")
    possiveis_lesoes = risk.get("possiveis_lesoes", "")
    fontes = risk.get("fontes_circunstancias") or "Durante o processo de trabalho."
    epis = risk.get("epis", "")
    epcs = risk.get("epcs", "")
    severidade = risk.get("grau_severidade", "")
    possibilidade = risk.get("grau_possibilidade", "")
    nivel = risk.get("grau_nivel_risco", "")

    _colored_value(_row_cell(risk_rows[0], 1), tipo, TIPO_RISCO_COLORS)
    _set_cell_text(_row_cell(risk_rows[1], 1), descricao)
    _set_cell_text(_row_cell(risk_rows[2], 1), possiveis_lesoes)
    _set_cell_text(_row_cell(risk_rows[3], 1), fontes)
    _set_cell_text(_row_cell(risk_rows[5], 1), epis)
    _set_cell_text(_row_cell(risk_rows[6], 1), epcs)

    _colored_value(_row_cell(risk_rows[9], 1), severidade, SEVERIDADE_COLORS)
    _colored_value(_row_cell(risk_rows[9], 3), possibilidade, POSSIBILIDADE_COLORS)
    _colored_value(_row_cell(risk_rows[9], 5), nivel, NIVEL_RISCO_COLORS)


def _fill_pgr_sector_table(table_xml, sector: Mapping[str, Any], risks: list[Mapping[str, Any]]) -> None:
    rows = table_xml.findall(qn("w:tr"))
    if len(rows) < 18:
        raise ValueError("O novo modelo Risco PGR precisa manter a tabela original com 18 linhas.")
    if not risks:
        raise ValueError("Cada setor selecionado precisa ter pelo menos um risco.")

    # 4 primeiras linhas do setor.
    _set_cell_text(_row_cell(rows[0], 0), sector.get("setor", ""))
    _set_cell_text(_row_cell(rows[2], 0), _sector_cargos_text(sector))

    risk_template_rows = [deepcopy(row) for row in rows[4:15]]
    _fill_pgr_risk_rows(rows[4:15], risks[0])

    # A partir do segundo risco, entram apenas as linhas do bloco do risco, sem repetir cabeçalho do setor.
    footer_first_row = rows[15]
    insertion_index = list(table_xml).index(footer_first_row)
    for risk in risks[1:]:
        block_rows = [deepcopy(row) for row in risk_template_rows]
        _fill_pgr_risk_rows(block_rows, risk)
        for block_row in block_rows:
            table_xml.insert(insertion_index, block_row)
            insertion_index += 1

    has_psychosocial = any(_normalize_option(risk.get("tipo_risco")) == "ERGONÔMICO PSICOSSOCIAL" for risk in risks)
    # O inventário deve sempre manter as linhas fixas:
    # CONTROLES EXISTENTES NO GES E SUA EFICÁCIA
    # Monitoramento da saúde do trabalhador através de exames ocupacionais.
    # A frase de ausência psicossocial só entra quando não houver risco psicossocial
    # e deve ser a última linha da tabela, abaixo do monitoramento.
    #
    # Importante: alguns modelos antigos deixam a frase antes das linhas de controle,
    # enquanto o PGR completo novo deixa a frase após o monitoramento. Portanto, nunca
    # usamos índice fixo para esta linha; buscamos a linha pelo próprio texto, removemos
    # todas as ocorrências e só depois recolocamos uma única vez no final quando necessário.
    phrase_template = None
    for row in list(table_xml.findall(qn("w:tr"))):
        if "NENHUM FATOR DE RISCO PSICOSSOCIAL" in _xml_text(row).upper():
            if phrase_template is None:
                phrase_template = deepcopy(row)
            table_xml.remove(row)

    if not has_psychosocial and phrase_template is not None:
        table_xml.append(phrase_template)


def generate_pgr_docx(groups_or_risks: Iterable[Mapping[str, Any]], output_path: str | Path) -> Path:
    items = list(groups_or_risks)
    if not items:
        raise ValueError("Selecione pelo menos um setor e um risco para gerar o PGR.")
    if not TEMPLATE_PGR_PATH.exists():
        raise FileNotFoundError(f"Modelo Risco PGR não encontrado: {TEMPLATE_PGR_PATH}")

    # Compatibilidade: se a chamada antiga passar apenas riscos, cria um setor genérico.
    if _is_grouped_payload(items):
        groups = _sanitize_sector_risk_groups(items)
    else:
        groups = [{"sector": {"setor": "", "cargos": []}, "risks": items}]

    if not groups:
        raise ValueError("Selecione pelo menos um setor e um risco para gerar o PGR.")

    doc = Document(str(TEMPLATE_PGR_PATH))
    if not doc.tables:
        raise ValueError("O modelo Risco PGR precisa ter uma tabela-modelo.")

    body = doc._body._element
    original_table_xml = doc.tables[0]._tbl
    template_table_copy = deepcopy(original_table_xml)
    body.remove(original_table_xml)

    for index, group in enumerate(groups):
        if index > 0:
            _insert_before_section_or_end(body, _page_break_paragraph())
        new_table = deepcopy(template_table_copy)
        _fill_pgr_sector_table(new_table, group["sector"], group["risks"])
        _insert_before_section_or_end(body, new_table)

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    save_docx_with_default_font(doc, output_path)
    return output_path


# ---------------------------------------------------------------------------
# PCMSO
# ---------------------------------------------------------------------------

def _fill_pcmso_table(table_xml, risk: Mapping[str, Any]) -> None:
    rows = table_xml.findall(qn("w:tr"))
    if len(rows) < 8:
        raise ValueError("O modelo PCMSO precisa manter a tabela original com 8 linhas.")

    tipo = risk.get("tipo_risco", "")
    descricao = risk.get("descricao_agente") or risk.get("risco", "")
    possiveis_lesoes = risk.get("possiveis_lesoes", "")
    fontes = risk.get("fontes_circunstancias") or "Durante o processo de trabalho."
    epis = risk.get("epis", "")
    epcs = risk.get("epcs", "")

    _colored_value(_value_cell(rows, 0, 1), tipo, TIPO_RISCO_COLORS)
    _set_cell_text(_value_cell(rows, 1, 1), descricao)
    _set_cell_text(_value_cell(rows, 2, 1), possiveis_lesoes)
    _set_cell_text(_value_cell(rows, 3, 1), fontes)
    _set_cell_text(_value_cell(rows, 5, 1), epis)
    _set_cell_text(_value_cell(rows, 6, 1), epcs)


def _flatten_risks_from_groups_or_risks(groups_or_risks: Iterable[Mapping[str, Any]]) -> list[Mapping[str, Any]]:
    items = list(groups_or_risks)
    if _is_grouped_payload(items):
        flattened: list[Mapping[str, Any]] = []
        for group in items:
            flattened.extend(group.get("risks", []))
        return flattened
    return items


def generate_pcmso_docx(groups_or_risks: Iterable[Mapping[str, Any]], output_path: str | Path) -> Path:
    risks = _flatten_risks_from_groups_or_risks(groups_or_risks)
    if not risks:
        raise ValueError("Selecione pelo menos um risco para gerar o PCMSO.")
    if not TEMPLATE_PCMSO_PATH.exists():
        raise FileNotFoundError(f"Modelo PCMSO não encontrado: {TEMPLATE_PCMSO_PATH}")

    doc = Document(str(TEMPLATE_PCMSO_PATH))
    if not doc.tables:
        raise ValueError("O modelo PCMSO precisa ter uma tabela-modelo.")

    body = doc._body._element
    original_table_xml = doc.tables[0]._tbl
    template_table_copy = deepcopy(original_table_xml)
    body.remove(original_table_xml)

    for risk in risks:
        new_table = deepcopy(template_table_copy)
        _fill_pcmso_table(new_table, risk)
        _insert_before_section_or_end(body, new_table)

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    save_docx_with_default_font(doc, output_path)
    return output_path


# ---------------------------------------------------------------------------
# RELAÇÃO FUNÇÃO X ATIVIDADE
# ---------------------------------------------------------------------------

def _fill_relacao_cargo_row(row_xml, cargo: Mapping[str, Any]) -> None:
    cells = row_xml.findall(qn("w:tc"))
    if len(cells) < 3:
        raise ValueError("A linha de cargo do modelo Relação Função x Atividade precisa ter 3 células.")
    cargo_text, cbo_text = _clean_cargo_and_cbo(cargo.get("cargo", ""), cargo.get("cbo", ""))
    func_text = str(cargo.get("n_func", "")).strip()
    desc_text = str(cargo.get("descricao", "")).strip()

    _set_cell_text(cells[0], f"{cargo_text} – CBO: {cbo_text}")
    _set_cell_text(cells[1], func_text)
    _set_cell_text(cells[2], desc_text)


def _fill_relacao_table(table_xml, sector: Mapping[str, Any], data_atual: str, data_final: str) -> None:
    sector = _sanitize_sector(sector)
    rows = table_xml.findall(qn("w:tr"))
    if len(rows) < 4:
        raise ValueError("O modelo Relação Função x Atividade precisa manter as 4 linhas originais.")

    _set_cell_text(_value_cell(rows, 0, 1), sector.get("setor", ""))
    _set_cell_text(_value_cell(rows, 1, 1), f"{data_atual} – {data_final}".strip(" –"))

    template_cargo_row = rows[3]
    template_cargo_copy = deepcopy(template_cargo_row)
    table_xml.remove(template_cargo_row)

    cargos = sector.get("cargos") or []
    if not cargos:
        cargos = [{"cargo": "", "cbo": "", "n_func": "", "descricao": ""}]
    for cargo in cargos:
        new_row = deepcopy(template_cargo_copy)
        _fill_relacao_cargo_row(new_row, cargo)
        table_xml.append(new_row)


def generate_relacao_funcao_atividade_docx(
    sectors: Iterable[Mapping[str, Any]],
    output_path: str | Path,
    data_atual: str | None = None,
    data_final: str | None = None,
    company_extra: Mapping[str, Any] | None = None,
) -> Path:
    sectors = [_sanitize_sector(sector) for sector in sectors]
    if not sectors:
        raise ValueError("Selecione pelo menos um setor para gerar o Word.")
    if not TEMPLATE_RELACAO_PATH.exists():
        raise FileNotFoundError(f"Modelo Relação Função x Atividade não encontrado: {TEMPLATE_RELACAO_PATH}")

    data_atual = (data_atual or datetime.now().strftime("%d/%m/%Y")).strip()
    data_final = (data_final or "").strip()

    doc = Document(str(TEMPLATE_RELACAO_PATH))
    if not doc.tables:
        raise ValueError("O modelo Relação Função x Atividade precisa ter uma tabela-modelo.")

    body = doc._body._element
    original_table_xml = doc.tables[0]._tbl
    template_table_copy = deepcopy(original_table_xml)
    body.remove(original_table_xml)

    for index, sector in enumerate(sectors):
        new_table = deepcopy(template_table_copy)
        _fill_relacao_table(new_table, sector, data_atual, data_final)
        _insert_before_section_or_end(body, new_table)
        if index < len(sectors) - 1:
            _insert_before_section_or_end(body, _blank_paragraph())

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    save_docx_with_default_font(doc, output_path)
    return output_path


# ---------------------------------------------------------------------------
# DESCRITIVO SETOR
# ---------------------------------------------------------------------------

def _fill_descritivo_setor_table(table_xml, sector: Mapping[str, Any]) -> None:
    rows = table_xml.findall(qn("w:tr"))
    if len(rows) < 1:
        raise ValueError("O modelo Descritivo Setor precisa manter a tabela original.")
    _set_cell_text(_value_cell(rows, 0, 1), str(sector.get("setor", "")).strip())


def generate_descritivo_setor_docx(sectors: Iterable[Mapping[str, Any]], output_path: str | Path) -> Path:
    sectors = [_sanitize_sector(sector) for sector in sectors]
    if not sectors:
        raise ValueError("Selecione pelo menos um setor para gerar o Word.")
    if not TEMPLATE_DESCRITIVO_SETOR_PATH.exists():
        raise FileNotFoundError(f"Modelo Descritivo Setor não encontrado: {TEMPLATE_DESCRITIVO_SETOR_PATH}")

    doc = Document(str(TEMPLATE_DESCRITIVO_SETOR_PATH))
    if not doc.tables:
        raise ValueError("O modelo Descritivo Setor precisa ter uma tabela-modelo.")

    body = doc._body._element
    original_table_xml = doc.tables[0]._tbl
    template_table_copy = deepcopy(original_table_xml)
    body.remove(original_table_xml)

    for sector in sectors:
        new_table = deepcopy(template_table_copy)
        _fill_descritivo_setor_table(new_table, sector)
        _insert_before_section_or_end(body, new_table)

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    save_docx_with_default_font(doc, output_path)
    return output_path


# ---------------------------------------------------------------------------
# PGR COMPLETO
# ---------------------------------------------------------------------------


def _xml_text(element) -> str:
    return "".join(node.text or "" for node in element.iter(qn("w:t")))


def _table_contains_all(table_xml, markers: Iterable[str]) -> bool:
    text = _xml_text(table_xml)
    return all(marker in text for marker in markers)


def _find_table_xml(doc: Document, markers: Iterable[str]):
    markers = list(markers)
    for table in doc.tables:
        if _table_contains_all(table._tbl, markers):
            return table._tbl
    raise ValueError(f"Não encontrei a tabela-modelo com os marcadores: {', '.join(markers)}")


def _replace_xml_element_with(parent_element, original_element, new_elements: Iterable) -> None:
    index = list(parent_element).index(original_element)
    parent_element.remove(original_element)
    for offset, element in enumerate(new_elements):
        parent_element.insert(index + offset, element)


def _replace_text_in_paragraph(paragraph, replacements: Mapping[str, str]) -> None:
    full_text = paragraph.text or ""
    if not any(marker in full_text for marker in replacements):
        return

    # Primeiro tenta substituir dentro dos próprios runs para preservar ao máximo a formatação.
    for run in paragraph.runs:
        text = run.text
        for marker, value in replacements.items():
            text = text.replace(marker, value)
        run.text = text

    # Se o Word tiver quebrado o marcador em vários runs, reconstrói o parágrafo usando o estilo do primeiro run.
    full_text = paragraph.text or ""
    if any(marker in full_text for marker in replacements):
        for marker, value in replacements.items():
            full_text = full_text.replace(marker, value)
        first_rpr = None
        if paragraph.runs and paragraph.runs[0]._r.rPr is not None:
            first_rpr = deepcopy(paragraph.runs[0]._r.rPr)
        paragraph.clear()
        run = paragraph.add_run(full_text)
        if first_rpr is not None:
            run._r.insert(0, first_rpr)


def _replace_doc_placeholders(doc: Document, replacements: Mapping[str, Any]) -> None:
    clean = {key: "" if value is None else str(value) for key, value in replacements.items()}

    def replace_in_table(table):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_text_in_paragraph(paragraph, clean)

    for paragraph in doc.paragraphs:
        _replace_text_in_paragraph(paragraph, clean)
    for table in doc.tables:
        replace_in_table(table)

    for section in doc.sections:
        parts = [section.header, section.footer]
        # Em alguns modelos o Word cria variações de primeira página/par/ímpar.
        for attr in ("first_page_header", "first_page_footer", "even_page_header", "even_page_footer"):
            part = getattr(section, attr, None)
            if part is not None:
                parts.append(part)
        for part in parts:
            for paragraph in part.paragraphs:
                _replace_text_in_paragraph(paragraph, clean)
            for table in part.tables:
                replace_in_table(table)



def _force_page_breaks_before_headings(doc: Document, headings: Iterable[str]) -> None:
    targets = tuple(headings)
    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").strip()
        if any(text.startswith(target) for target in targets):
            paragraph.paragraph_format.page_break_before = True


def _compact_first_page_date_block(doc: Document) -> None:
    """Mantém o bloco Data/Revisão compacto para caber na primeira página da capa."""
    date_index = None
    for index, paragraph in enumerate(doc.paragraphs):
        text = (paragraph.text or "").strip()
        if text.startswith("Data do documento:"):
            date_index = index
            break

    if date_index is None:
        return

    # Remove a quebra forçada antes da data, caso tenha sido aplicada em versões anteriores.
    for paragraph in doc.paragraphs[date_index: min(len(doc.paragraphs), date_index + 5)]:
        paragraph.paragraph_format.page_break_before = False
        paragraph.paragraph_format.space_before = 0
        paragraph.paragraph_format.space_after = 0

    # Compacta parágrafos vazios imediatamente anteriores ao bloco da data.
    for paragraph in reversed(doc.paragraphs[:date_index]):
        if (paragraph.text or "").strip():
            break
        paragraph.text = ""
        paragraph.paragraph_format.space_before = 0
        paragraph.paragraph_format.space_after = 0
        paragraph.paragraph_format.line_spacing = 1


def _clone_paragraph_style(new_p, source_p) -> None:
    p_pr = source_p.find(qn("w:pPr"))
    if p_pr is not None:
        new_p.insert(0, deepcopy(p_pr))


def _make_revision_paragraph(text: str, source_paragraph) -> OxmlElement:
    p = OxmlElement("w:p")
    _clone_paragraph_style(p, source_paragraph._p)
    r = OxmlElement("w:r")
    r_pr = None
    first_run = source_paragraph._p.find(qn("w:r"))
    if first_run is not None:
        maybe_rpr = first_run.find(qn("w:rPr"))
        if maybe_rpr is not None:
            r_pr = deepcopy(maybe_rpr)
    if r_pr is not None:
        r.append(r_pr)
    for node in _make_text(text):
        r.append(node)
    p.append(r)
    return p


def _insert_psychosocial_revision_line(doc: Document, company: Mapping[str, str]) -> None:
    """Insere a linha de revisão de inclusão de risco psicossocial entre Data do documento e Revisão periódica."""
    if company.get("ajuste_psicossocial") != "1":
        return
    data_revisao = (company.get("data_da_revisao") or "").strip()
    if not data_revisao:
        return
    text = f"Revisão periódica de inclusão de Risco Psicossocial: {data_revisao}"
    for paragraph in doc.paragraphs:
        p_text = (paragraph.text or "").strip()
        if p_text.startswith("Data do documento:"):
            # Evita duplicar a linha caso o documento seja processado novamente.
            next_el = paragraph._p.getnext()
            while next_el is not None and next_el.tag == qn("w:p"):
                next_text = "".join(t.text or "" for t in next_el.iter(qn("w:t"))).strip()
                if next_text:
                    if next_text.startswith("Revisão periódica de inclusão de Risco Psicossocial"):
                        return
                    break
                next_el = next_el.getnext()
            new_p = _make_revision_paragraph(text, paragraph)
            paragraph._p.addnext(new_p)
            return


def _compact_known_static_spacers(doc: Document) -> None:
    """Suaviza espaçamentos exagerados de páginas estáticas do modelo, sem alterar tabelas."""
    compact_starts = (
        "FICHA DE ENTREGA GRATUITA DE EPI",
        "TERMO DE RESPONSABILIDADE",
    )
    page_break_and_compact = (
        "IDENTIFICAÇÃO DA EMPRESA",
    )
    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").strip()
        if any(text.startswith(item) for item in compact_starts):
            paragraph.paragraph_format.space_before = 0
            paragraph.paragraph_format.space_after = 0
        if any(text.startswith(item) for item in page_break_and_compact):
            paragraph.paragraph_format.page_break_before = True
            paragraph.paragraph_format.space_before = 0
            paragraph.paragraph_format.space_after = 0


def _build_relacao_elements(template_table_xml, sectors: list[Mapping[str, Any]], data_atual: str, data_final: str) -> list:
    elements: list = []
    for index, sector in enumerate(sectors):
        new_table = deepcopy(template_table_xml)
        _fill_relacao_table(new_table, sector, data_atual, data_final)
        elements.append(new_table)
        if index < len(sectors) - 1:
            elements.append(_blank_paragraph())
    return elements


def _build_descritivo_elements(template_table_xml, sectors: list[Mapping[str, Any]]) -> list:
    elements: list = []
    for sector in sectors:
        new_table = deepcopy(template_table_xml)
        _fill_descritivo_setor_table(new_table, sector)
        elements.append(new_table)
    return elements


def _build_risco_pgr_elements(template_table_xml, groups: list[Mapping[str, Any]], break_before_first: bool = False) -> list:
    elements: list = []
    for index, group in enumerate(groups):
        if index > 0 or (index == 0 and break_before_first):
            elements.append(_page_break_paragraph())
        new_table = deepcopy(template_table_xml)
        _fill_pgr_sector_table(new_table, group["sector"], group["risks"])
        elements.append(new_table)
    return elements


def _fill_action_plan_table_xml(table_xml, groups: list[Mapping[str, Any]], data_atual: str = "", data_final: str = "") -> None:
    entries = _plano_entries_from_groups_or_risks(groups)
    if not entries:
        raise ValueError("Selecione pelo menos um setor e um risco para preencher o Plano de Ação no PGR.")

    template_row_xml = _find_template_row(table_xml)
    if template_row_xml is None:
        raise ValueError("A tabela de Plano de Ação do PGR precisa ter uma linha-modelo com {{risco}}.")
    template_row_copy = deepcopy(template_row_xml)
    table_xml.remove(template_row_xml)

    for setor, risk in entries:
        new_row = deepcopy(template_row_copy)
        _fill_plano_row(new_row, risk, setor=setor, data_atual=data_atual, data_final=data_final)
        table_xml.append(new_row)


def _build_action_plan_elements(template_table_xml, groups: list[Mapping[str, Any]], data_atual: str = "", data_final: str = "") -> list:
    new_table = deepcopy(template_table_xml)
    _fill_action_plan_table_xml(new_table, groups, data_atual=data_atual, data_final=data_final)
    return [new_table]


def _unique_sectors_from_groups(groups: list[Mapping[str, Any]]) -> list[dict[str, Any]]:
    sectors: list[dict[str, Any]] = []
    seen: set[str] = set()
    for group in groups:
        sector = _sanitize_sector(group.get("sector") or {})
        key = sector.get("setor", "").strip().lower()
        if key and key not in seen:
            sectors.append(sector)
            seen.add(key)
    return sectors




def _company_dict_for_docs(empresa: str, cnpj: str, data_atual: str, data_final: str, extra: Mapping[str, Any] | None = None) -> dict[str, str]:
    extra = extra or {}
    def pick(*keys: str) -> str:
        for key in keys:
            value = extra.get(key)
            if value not in (None, ""):
                return str(value).strip()
        return ""
    data_criacao = pick("data_criacao_laudo", "datacriacaolaudo", "data_avaliacao")
    return {
        "empresa": str(empresa or pick("empresa", "nome")).strip(),
        "cnpj": str(cnpj or pick("cnpj")).strip(),
        "data_atual": str(data_atual or pick("data_atual")).strip(),
        "data_final": str(data_final or pick("data_final")).strip(),
        "endereco": pick("endereco"),
        "bairro_cidade": pick("bairro_cidade", "bairro/cidade"),
        "cep": pick("cep"),
        "cnae": pick("cnae", "cnae1"),
        "descricao_atividade": pick("descricao_atividade", "descricao1"),
        "grau_risco": pick("grau_risco", "grau1"),
        "cnae_secundario": pick("cnae_secundario", "cnae2"),
        "descricao_atividade_secundaria": pick("descricao_atividade_secundaria", "descricao2"),
        "grau_risco_secundario": pick("grau_risco_secundario", "grau2"),
        "funcionarios": pick("funcionarios"),
        "email": pick("email"),
        "fone": pick("fone"),
        "data_avaliacao": data_criacao or str(data_atual or "").strip(),
        "data_criacao_laudo": data_criacao,
        "ajuste_psicossocial": pick("ajuste_psicossocial"),
        "data_da_revisao": pick("data_da_revisao", "data_revisao_ajuste"),
    }


def _total_funcionarios_from_sectors(sectors: list[Mapping[str, Any]]) -> str:
    total = 0
    for sector in sectors:
        for cargo in sector.get("cargos", []) or []:
            try:
                total += int(str(cargo.get("n_func", "")).strip() or "0")
            except ValueError:
                pass
    return str(total) if total else ""


def _company_replacements(company: Mapping[str, str]) -> dict[str, str]:
    funcionarios = company.get("funcionarios", "")
    return {
        "{{EMPRESA}}": company.get("empresa", ""),
        "{{empresa}}": company.get("empresa", ""),
        "{{CNPJ}}": company.get("cnpj", ""),
        "{{DATAATUAL}}": company.get("data_atual", ""),
        "{{DATAFINAL}}": company.get("data_final", ""),
        "{{DATA DO LAUDO}}": company.get("data_atual", ""),
        "{{ENDERECO}}": company.get("endereco", ""),
        "{{BAIRRO/CIDADE}}": company.get("bairro_cidade", ""),
        "{{CEP}}": company.get("cep", ""),
        "{{CNAE1}}": company.get("cnae", ""),
        "{{DESCRICAO1}}": company.get("descricao_atividade", ""),
        "{{GRAU1}}": company.get("grau_risco", ""),
        "{{CNAE2}}": company.get("cnae_secundario", ""),
        "{{DESCRICAO2}}": company.get("descricao_atividade_secundaria", ""),
        "{{GRAU2}}": company.get("grau_risco_secundario", ""),
        "{{N°FUNCIONARIOS}}": funcionarios,
        "{{EMAIL}}": company.get("email", ""),
        "{{FONE}}": company.get("fone", ""),
        "{{DATACRIACAOLAUDO}}": company.get("data_criacao_laudo") or company.get("data_avaliacao") or company.get("data_atual", ""),
        "{{datacriacaolaudo}}": company.get("data_criacao_laudo") or company.get("data_avaliacao") or company.get("data_atual", ""),
        "{{DATADAREVISAO}}": company.get("data_da_revisao", ""),
        "{{datadarevisao}}": company.get("data_da_revisao", ""),
    }


def _fill_company_identification_tables(doc: Document, company: Mapping[str, str], sectors: list[Mapping[str, Any]]) -> None:
    funcionarios = company.get("funcionarios") or _total_funcionarios_from_sectors(sectors)
    values_by_label = {
        "EMPRESA": company.get("empresa", ""),
        "ENDEREÇO": company.get("endereco", ""),
        "BAIRRO / CIDADE": company.get("bairro_cidade", ""),
        "CEP": company.get("cep", ""),
        "CNPJ": company.get("cnpj", ""),
        "CNAE": company.get("cnae", ""),
        "GRAU DE RISCO": company.get("grau_risco", ""),
        "CNAE (SECUNDÁRIO)": company.get("cnae_secundario", ""),
        "GRAU DE RISCO (SECUNDÁRIO)": company.get("grau_risco_secundario", ""),
        "FUNCIONÁRIOS": funcionarios,
        "VIGÊNCIA": f"{company.get('data_atual','')} – {company.get('data_final','')}",
        "EMAIL": company.get("email", ""),
        "FONE": company.get("fone", ""),
    }
    for table in doc.tables:
        labels = [row.cells[0].text.strip().upper() for row in table.rows if row.cells]
        if "EMPRESA" not in labels or not any(label in labels for label in ("CNPJ", "VIGÊNCIA", "CNAE", "FUNCIONÁRIOS")):
            continue
        descricao_seen = 0
        for row in table.rows:
            if len(row.cells) < 2:
                continue
            label = row.cells[0].text.strip().upper()
            if label == "DESCRIÇÃO DA ATIVIDADE":
                descricao_seen += 1
                value = company.get("descricao_atividade_secundaria", "") if descricao_seen == 2 else company.get("descricao_atividade", "")
                _ltcat_set_cell_text(row.cells[1], value)
            elif label in values_by_label:
                _ltcat_set_cell_text(row.cells[1], values_by_label[label])

def generate_complete_pgr_docx(
    groups_or_risks: Iterable[Mapping[str, Any]],
    output_path: str | Path,
    empresa: str,
    cnpj: str,
    data_atual: str | None = None,
    data_final: str | None = None,
    company_extra: Mapping[str, Any] | None = None,
) -> Path:
    """Gera o PGR completo preenchendo o modelo principal com todos os blocos já criados."""
    items = list(groups_or_risks)
    if not items:
        raise ValueError("Selecione pelo menos um setor e um risco para gerar o PGR completo.")
    if not empresa or not str(empresa).strip():
        raise ValueError("Preencha o nome da empresa para gerar o PGR completo.")
    if not cnpj or not str(cnpj).strip():
        raise ValueError("Preencha o CNPJ da empresa para gerar o PGR completo.")
    if not TEMPLATE_PGR_COMPLETO_PATH.exists():
        raise FileNotFoundError(f"Modelo PGR completo não encontrado: {TEMPLATE_PGR_COMPLETO_PATH}")

    if _is_grouped_payload(items):
        groups = _sanitize_sector_risk_groups(items)
    else:
        groups = [{"sector": {"setor": "", "cargos": []}, "risks": items}]
    if not groups:
        raise ValueError("Selecione pelo menos um setor e um risco para gerar o PGR completo.")

    empresa = str(empresa).strip()
    cnpj = str(cnpj).strip()
    data_atual = (data_atual or datetime.now().strftime("%d/%m/%Y")).strip()
    data_final = (data_final or "").strip()
    company = _company_dict_for_docs(empresa, cnpj, data_atual, data_final, company_extra)

    doc = Document(str(TEMPLATE_PGR_COMPLETO_PATH))

    relation_table = _find_table_xml(doc, ["{{CARGO}}", "{{CBOCARGO}}", "{{N°FUNC}}", "{{DESCRIÇÃO ATIVIDADE}}"])
    descritivo_table = _find_table_xml(doc, ["PAREDE:", "{{SETOR}}", "VENTILAÇÃO:"])
    risco_pgr_table = _find_table_xml(doc, ["{{CARGOS}}", "{{TIPO DE RISCO}}", "{{GRAU DE NIVEL DE RISCO}}"])
    plano_table = _find_table_xml(doc, ["PLANO", "{{risco}}", "{{AÇÕES PREVENTIVA / CORRETIVA}}", "{{INDICADOR DE EFETIVIDADE}}"])

    sectors = _unique_sectors_from_groups(groups)
    if not sectors:
        raise ValueError("Cadastre e selecione pelo menos um setor para gerar o PGR completo.")

    # Substitui os blocos-modelo por blocos gerados, mantendo a posição original dentro do PGR.
    _replace_xml_element_with(relation_table.getparent(), relation_table, _build_relacao_elements(relation_table, sectors, data_atual, data_final))
    _replace_xml_element_with(descritivo_table.getparent(), descritivo_table, _build_descritivo_elements(descritivo_table, sectors))
    _replace_xml_element_with(risco_pgr_table.getparent(), risco_pgr_table, _build_risco_pgr_elements(risco_pgr_table, groups, break_before_first=False))
    # O plano de ação fica em página própria para evitar que linhas com vários setores na coluna GES
    # sejam quebradas entre páginas no final do inventário.
    _insert_page_break_before_previous_paragraph(plano_table, "PLANO DE AÇÃO")
    _replace_xml_element_with(plano_table.getparent(), plano_table, _build_action_plan_elements(plano_table, groups, data_atual=data_atual, data_final=data_final))

    _fill_company_identification_tables(doc, company, sectors)
    _insert_psychosocial_revision_line(doc, company)
    _replace_doc_placeholders(doc, _company_replacements(company))

    # Ajustes de layout solicitados:
    # - manter o bloco Data do documento/Revisão compacto e ainda na primeira página;
    # - evitar páginas em branco entre o descritivo e o inventário de riscos.
    _compact_first_page_date_block(doc)
    _compact_known_static_spacers(doc)

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    save_docx_with_default_font(doc, output_path)
    return output_path


# ---------------------------------------------------------------------------
# PCMSO COMPLETO E BLOCO DE RISCOS/EXAMES
# ---------------------------------------------------------------------------


def _fill_pcmso_risk_rows(risk_rows: list, risk: Mapping[str, Any]) -> None:
    if len(risk_rows) != 4:
        raise ValueError("O bloco de risco do PCMSO precisa manter 4 linhas.")
    tipo = risk.get("tipo_risco", "")
    descricao = risk.get("descricao_agente") or risk.get("risco", "")
    possiveis_lesoes = risk.get("possiveis_lesoes", "")
    fontes = risk.get("fontes_circunstancias") or "Durante o processo de trabalho."

    _colored_value(_row_cell(risk_rows[0], 1), tipo, TIPO_RISCO_COLORS)
    _set_cell_text(_row_cell(risk_rows[1], 1), descricao)
    _set_cell_text(_row_cell(risk_rows[2], 1), possiveis_lesoes)
    _set_cell_text(_row_cell(risk_rows[3], 1), fontes)


def _fill_pcmso_exam_row(row_xml, exam: Mapping[str, Any]) -> None:
    cells = row_xml.findall(qn("w:tc"))
    if len(cells) < 7:
        raise ValueError("A linha de exame do modelo PCMSO precisa manter 7 células.")
    values = [
        exam.get("exame", ""),
        exam.get("periodicidade", ""),
        exam.get("admissional", ""),
        exam.get("periodico", ""),
        exam.get("retorno", ""),
        exam.get("mudanca", ""),
        exam.get("demissional", ""),
    ]
    for cell, value in zip(cells[:7], values):
        _set_cell_text(cell, value)


def _sanitize_exam(exam: Mapping[str, Any]) -> dict[str, Any]:
    return {
        "id": str(exam.get("id", "")).strip(),
        "exame": str(exam.get("exame", "")).strip(),
        "periodicidade": str(exam.get("periodicidade", "")).strip(),
        "admissional": str(exam.get("admissional", "")).strip(),
        "periodico": str(exam.get("periodico", "")).strip(),
        "retorno": str(exam.get("retorno", "")).strip(),
        "mudanca": str(exam.get("mudanca", "")).strip(),
        "demissional": str(exam.get("demissional", "")).strip(),
    }


def _sanitize_pcmso_groups(groups: Iterable[Mapping[str, Any]]) -> list[dict[str, Any]]:
    clean_groups: list[dict[str, Any]] = []
    for group in groups:
        sector = _sanitize_sector(group.get("sector") or {})
        risks = [risk for risk in (group.get("risks") or []) if risk]
        exams = [_sanitize_exam(exam) for exam in (group.get("exams") or []) if exam]
        if sector.get("setor") and risks:
            clean_groups.append({"sector": sector, "risks": risks, "exams": exams})
    return clean_groups


def _fill_pcmso_riscos_sector_table(table_xml, sector: Mapping[str, Any], risks: list[Mapping[str, Any]], exams: list[Mapping[str, Any]]) -> None:
    rows = table_xml.findall(qn("w:tr"))
    if len(rows) < 12:
        raise ValueError("O modelo riscos PCMSO precisa manter 12 linhas.")
    if not risks:
        raise ValueError("Cada setor selecionado precisa ter pelo menos um risco.")

    _set_cell_text(_row_cell(rows[0], 0), sector.get("setor", ""))
    _set_cell_text(_row_cell(rows[2], 0), _sector_cargos_text(sector))

    risk_template_rows = [deepcopy(row) for row in rows[4:8]]
    _fill_pcmso_risk_rows(rows[4:8], risks[0])

    phrase_row = rows[8]
    insertion_index = list(table_xml).index(phrase_row)
    for risk in risks[1:]:
        block_rows = [deepcopy(row) for row in risk_template_rows]
        _fill_pcmso_risk_rows(block_rows, risk)
        for block_row in block_rows:
            table_xml.insert(insertion_index, block_row)
            insertion_index += 1

    has_psychosocial = any(_normalize_option(risk.get("tipo_risco")) == "ERGONÔMICO PSICOSSOCIAL" for risk in risks)
    # A frase fica somente quando NÃO existe risco ERGONÔMICO PSICOSSOCIAL no setor.
    # Quando existir, ela é removida; quando não existir, ela é movida para ser a última linha da tabela.
    keep_phrase_row = not has_psychosocial
    if has_psychosocial:
        table_xml.remove(phrase_row)
    else:
        table_xml.remove(phrase_row)

    rows = table_xml.findall(qn("w:tr"))
    exam_template_row = None
    for row in rows:
        if "{{exame}}" in _xml_text(row):
            exam_template_row = row
            break
    if exam_template_row is None:
        raise ValueError("O modelo riscos PCMSO precisa ter uma linha de exame com {{exame}}.")
    exam_template_copy = deepcopy(exam_template_row)
    insert_index = list(table_xml).index(exam_template_row)
    table_xml.remove(exam_template_row)

    if not exams:
        exams = [{"exame": "", "periodicidade": "", "admissional": "", "periodico": "", "retorno": "", "mudanca": "", "demissional": ""}]
    for exam in exams:
        row = deepcopy(exam_template_copy)
        _fill_pcmso_exam_row(row, exam)
        table_xml.insert(insert_index, row)
        insert_index += 1

    if keep_phrase_row:
        table_xml.append(phrase_row)


def _build_pcmso_riscos_elements(template_table_xml, groups: list[Mapping[str, Any]], page_break_between_sectors: bool = True) -> list:
    elements: list = []
    for index, group in enumerate(groups):
        if index > 0 and page_break_between_sectors:
            elements.append(_page_break_paragraph())
        table = deepcopy(template_table_xml)
        _fill_pcmso_riscos_sector_table(table, group["sector"], group["risks"], group.get("exams", []))
        elements.append(table)
    return elements


def generate_riscos_pcmso_docx(groups: Iterable[Mapping[str, Any]], output_path: str | Path) -> Path:
    groups = _sanitize_pcmso_groups(groups)
    if not groups:
        raise ValueError("Selecione pelo menos um setor e um risco para gerar o bloco de riscos do PCMSO.")
    if not TEMPLATE_RISCOS_PCMSO_PATH.exists():
        raise FileNotFoundError(f"Modelo riscos PCMSO não encontrado: {TEMPLATE_RISCOS_PCMSO_PATH}")
    doc = Document(str(TEMPLATE_RISCOS_PCMSO_PATH))
    if not doc.tables:
        raise ValueError("O modelo riscos PCMSO precisa ter uma tabela-modelo.")
    body = doc._body._element
    original_table_xml = doc.tables[0]._tbl
    template_table_copy = deepcopy(original_table_xml)
    body.remove(original_table_xml)
    for element in _build_pcmso_riscos_elements(template_table_copy, groups):
        _insert_before_section_or_end(body, element)
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    save_docx_with_default_font(doc, output_path)
    return output_path


def _remove_element(element) -> None:
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _replace_pcmso_riscos_area(doc: Document, groups: list[Mapping[str, Any]]) -> None:
    risco_table = _find_table_xml(doc, ["{{CARGOS}}", "{{TIPO DE RISCO}}", "{{risco}}"])
    phrase_table = _find_table_xml(doc, ["NENHUM FATOR DE RISCO PSICOSSOCIAL"])
    exams_table = _find_table_xml(doc, ["EXAMES RECOMENDADOS", "{{tempo de peridiocidade}}"])

    template_doc = Document(str(TEMPLATE_RISCOS_PCMSO_PATH))
    template_table = deepcopy(template_doc.tables[0]._tbl)
    generated = [_page_break_paragraph()] + _build_pcmso_riscos_elements(template_table, groups)

    parent = risco_table.getparent()
    index = list(parent).index(risco_table)
    _remove_element(risco_table)
    _remove_element(phrase_table)
    _remove_element(exams_table)
    for offset, element in enumerate(generated):
        parent.insert(index + offset, element)


def generate_complete_pcmso_docx(
    groups: Iterable[Mapping[str, Any]],
    output_path: str | Path,
    empresa: str,
    cnpj: str,
    data_atual: str | None = None,
    data_final: str | None = None,
    company_extra: Mapping[str, Any] | None = None,
) -> Path:
    groups = _sanitize_pcmso_groups(groups)
    if not groups:
        raise ValueError("Selecione pelo menos um setor e um risco para gerar o PCMSO completo.")
    empresa = str(empresa or "").strip()
    cnpj = str(cnpj or "").strip()
    if not empresa:
        raise ValueError("Preencha o nome da empresa.")
    if not cnpj:
        raise ValueError("Preencha o CNPJ.")
    if not TEMPLATE_PCMSO_COMPLETO_PATH.exists():
        raise FileNotFoundError(f"Modelo PCMSO completo não encontrado: {TEMPLATE_PCMSO_COMPLETO_PATH}")
    data_atual = (data_atual or datetime.now().strftime("%m/%Y")).strip()
    data_final = (data_final or "").strip()

    company = _company_dict_for_docs(empresa, cnpj, data_atual, data_final, company_extra)
    doc = Document(str(TEMPLATE_PCMSO_COMPLETO_PATH))
    relation_table = _find_table_xml(doc, ["{{CARGO}}", "{{CBOCARGO}}", "{{N°FUNC}}", "{{DESCRIÇÃO ATIVIDADE}}"])
    sectors = _unique_sectors_from_groups(groups)
    _replace_xml_element_with(relation_table.getparent(), relation_table, _build_relacao_elements(relation_table, sectors, data_atual, data_final))
    _replace_pcmso_riscos_area(doc, groups)

    _fill_company_identification_tables(doc, company, sectors)
    _insert_psychosocial_revision_line(doc, company)
    _replace_doc_placeholders(doc, _company_replacements(company))
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    save_docx_with_default_font(doc, output_path)
    return output_path


# ---------------------------------------------------------------------------
# LTCAT COMPLETO
# ---------------------------------------------------------------------------

TIPOS_RISCO_LTCAT = {"FÍSICO", "QUÍMICO", "BIOLÓGICO"}
DEFAULT_LTCAT_ENQUADRAMENTO = (
    "Considerando a atividade que pode causar é recomendado adotar medidas preventivas "
    "até que as avaliações quantitativas sejam realizadas."
)
DEFAULT_LTCAT_PARECER = (
    "Considerando a atividade que pode causar é recomendado adotar medidas preventivas "
    "até que as avaliações quantitativas sejam realizadas."
)


def _ltcat_table_has_text(table, text: str) -> bool:
    return text in "\n".join(cell.text for row in table.rows for cell in row.cells)


def _ltcat_replace_in_table(table, replacements: Mapping[str, str]) -> None:
    clean = {key: "" if value is None else str(value) for key, value in replacements.items()}
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _replace_text_in_paragraph(paragraph, clean)


def _ltcat_set_cell_text(cell, text_value: Any) -> None:
    text_value = "" if text_value is None else str(text_value)
    if not cell.paragraphs:
        paragraph = cell.add_paragraph()
        paragraph.add_run(text_value)
        return
    paragraph = cell.paragraphs[0]
    if paragraph.runs:
        paragraph.runs[0].text = text_value
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text_value)
    for extra in cell.paragraphs[1:]:
        for run in extra.runs:
            run.text = ""


def _ltcat_shade_cell(cell, fill: str | None) -> None:
    if not fill:
        return
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _ltcat_set_cell_font_color(cell, rgb: str = "000000") -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor.from_string(rgb)


def _ltcat_insert_tbl_before(anchor_table, tbl_element) -> None:
    anchor_table._tbl.addprevious(tbl_element)


def _ltcat_remove_table(table) -> None:
    table._tbl.getparent().remove(table._tbl)


def _ltcat_insert_blank_before(anchor_table) -> None:
    p = OxmlElement("w:p")
    p_pr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    p_pr.append(spacing)
    p.append(p_pr)
    anchor_table._tbl.addprevious(p)


def _ltcat_insert_page_break_before(anchor_table) -> None:
    p = OxmlElement("w:p")
    p_pr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    p_pr.append(spacing)
    p.append(p_pr)
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    p.append(r)
    anchor_table._tbl.addprevious(p)


def _ltcat_remove_inventory_heading_paragraphs(doc: Document) -> None:
    """Remove títulos soltos de INVENTÁRIO DE RISCOS do modelo LTCAT.

    O título precisa ficar apenas imediatamente antes dos blocos de riscos
    gerados. Como o modelo pode ter esse texto em uma posição fixa antiga,
    removemos as ocorrências exatas e reinserimos no ponto correto.
    """
    for paragraph in list(doc.paragraphs):
        if paragraph.text.strip().upper() == "INVENTÁRIO DE RISCOS":
            paragraph._p.getparent().remove(paragraph._p)


def _ltcat_insert_inventory_heading_before(anchor_table) -> None:
    p = OxmlElement("w:p")
    p_pr = OxmlElement("w:pPr")

    keep_next = OxmlElement("w:keepNext")
    p_pr.append(keep_next)

    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "120")
    p_pr.append(spacing)

    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    p_pr.append(jc)

    p.append(p_pr)

    r = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Arial Narrow")
    r_fonts.set(qn("w:hAnsi"), "Arial Narrow")
    r_pr.append(r_fonts)
    bold = OxmlElement("w:b")
    r_pr.append(bold)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "24")
    r_pr.append(sz)
    r.append(r_pr)

    t = OxmlElement("w:t")
    t.text = "INVENTÁRIO DE RISCOS"
    r.append(t)
    p.append(r)
    anchor_table._tbl.addprevious(p)


def _ltcat_sector_cargos_text(sector: Mapping[str, Any]) -> str:
    return _sector_cargos_text(sector).upper()


def _ltcat_company_dict(empresa: str, cnpj: str, data_atual: str, data_final: str, extra: Mapping[str, Any] | None = None) -> dict[str, str]:
    extra = extra or {}
    data_criacao = str(extra.get("data_criacao_laudo") or extra.get("datacriacaolaudo") or extra.get("data_avaliacao") or "").strip()
    return {
        "empresa": str(empresa or "").strip(),
        "cnpj": str(cnpj or "").strip(),
        "data_atual": str(data_atual or "").strip(),
        "data_final": str(data_final or "").strip(),
        "endereco": str(extra.get("endereco", "") or "").strip(),
        "bairro_cidade": str(extra.get("bairro_cidade", "") or "").strip(),
        "cep": str(extra.get("cep", "") or "").strip(),
        "cnae": str(extra.get("cnae", extra.get("cnae1", "")) or "").strip(),
        "descricao_atividade": str(extra.get("descricao_atividade", extra.get("descricao1", "")) or "").strip(),
        "grau_risco": str(extra.get("grau_risco", extra.get("grau1", "")) or "").strip(),
        "cnae_secundario": str(extra.get("cnae_secundario", extra.get("cnae2", "")) or "").strip(),
        "descricao_atividade_secundaria": str(extra.get("descricao_atividade_secundaria", extra.get("descricao2", "")) or "").strip(),
        "grau_risco_secundario": str(extra.get("grau_risco_secundario", extra.get("grau2", "")) or "").strip(),
        "funcionarios": str(extra.get("funcionarios", "") or "").strip(),
        "email": str(extra.get("email", "") or "").strip(),
        "fone": str(extra.get("fone", "") or "").strip(),
        "data_avaliacao": data_criacao or str(data_atual or "").strip(),
        "data_criacao_laudo": data_criacao,
        "ajuste_psicossocial": str(extra.get("ajuste_psicossocial", "") or "").strip(),
        "data_da_revisao": str(extra.get("data_da_revisao", extra.get("data_revisao_ajuste", "")) or "").strip(),
    }


def _ltcat_total_funcionarios(sectors: list[Mapping[str, Any]]) -> str:
    total = 0
    for sector in sectors:
        for cargo in sector.get("cargos", []) or []:
            try:
                total += int(str(cargo.get("n_func", "")).strip() or "0")
            except ValueError:
                pass
    return str(total) if total else ""


def _ltcat_fill_company_table(doc: Document, company: Mapping[str, str], sectors: list[Mapping[str, Any]]) -> None:
    for table in doc.tables:
        if len(table.rows) >= 14 and table.cell(0, 0).text.strip().upper() == "EMPRESA":
            funcionarios = company.get("funcionarios") or _ltcat_total_funcionarios(sectors)
            values_by_label = {
                "EMPRESA": company.get("empresa", ""),
                "ENDEREÇO": company.get("endereco", ""),
                "BAIRRO / CIDADE": company.get("bairro_cidade", ""),
                "CEP": company.get("cep", ""),
                "CNPJ": company.get("cnpj", ""),
                "CNAE": company.get("cnae", ""),
                "GRAU DE RISCO": company.get("grau_risco", ""),
                "CNAE (SECUNDÁRIO)": company.get("cnae_secundario", ""),
                "GRAU DE RISCO (SECUNDÁRIO)": company.get("grau_risco_secundario", ""),
                "FUNCIONÁRIOS": funcionarios,
                "VIGÊNCIA": f"{company.get('data_atual','')} – {company.get('data_final','')}",
                "EMAIL": company.get("email", ""),
                "FONE": company.get("fone", ""),
            }
            descricao_seen = 0
            for row in table.rows:
                label = row.cells[0].text.strip().upper()
                if label == "DESCRIÇÃO DA ATIVIDADE":
                    descricao_seen += 1
                    value = company.get("descricao_atividade_secundaria", "") if descricao_seen == 2 else company.get("descricao_atividade", "")
                    _ltcat_set_cell_text(row.cells[1], value)
                elif label in values_by_label:
                    _ltcat_set_cell_text(row.cells[1], values_by_label[label])
            break


def _ltcat_fill_relacao_funcoes(doc: Document, sectors: list[Mapping[str, Any]], data_atual: str, data_final: str) -> None:
    relation_tables = [t for t in doc.tables if _ltcat_table_has_text(t, "{{CARGO}}") and _ltcat_table_has_text(t, "{{CBOCARGO}}")]
    if not relation_tables:
        return
    anchor = relation_tables[0]
    template_tbl = deepcopy(anchor._tbl)
    for index, sector in enumerate(sectors):
        new_tbl = deepcopy(template_tbl)
        _ltcat_insert_tbl_before(anchor, new_tbl)
        table = Table(new_tbl, doc)
        _fill_relacao_table(table._tbl, sector, data_atual, data_final)
        if index < len(sectors) - 1:
            _ltcat_insert_blank_before(anchor)
    _ltcat_remove_table(anchor)


def _ltcat_fill_descritivo_setores(doc: Document, sectors: list[Mapping[str, Any]]) -> None:
    desc_tables = [t for t in doc.tables if _ltcat_table_has_text(t, "PAREDE:") and _ltcat_table_has_text(t, "{{SETOR}}")]
    if not desc_tables:
        return
    anchor = desc_tables[0]
    template_tbl = deepcopy(anchor._tbl)
    for index, sector in enumerate(sectors):
        new_tbl = deepcopy(template_tbl)
        _ltcat_insert_tbl_before(anchor, new_tbl)
        table = Table(new_tbl, doc)
        _ltcat_replace_in_table(table, {"{{SETOR}}": str(sector.get("setor", ""))})
        if index < len(sectors) - 1:
            _ltcat_insert_blank_before(anchor)
    _ltcat_remove_table(anchor)


def _ltcat_risk_from_common(risk: Mapping[str, Any]) -> dict[str, Any]:
    return {
        "id": str(risk.get("id", "")),
        "risco": str(risk.get("risco", "") or ""),
        "tipo_risco": _normalize_option(risk.get("tipo_risco")),
        "meio_propagacao": str(risk.get("ltcat_meio_propagacao") or risk.get("fontes_circunstancias") or ""),
        "epis": str(risk.get("epis", "") or ""),
        "epcs": str(risk.get("epcs", "") or ""),
        "insalubridade": str(risk.get("ltcat_insalubridade") or "Não"),
        "grau_insalubridade": str(risk.get("ltcat_grau_insalubridade") or "Não aplicável"),
        "aposentadoria_especial": str(risk.get("ltcat_aposentadoria_especial") or "Não"),
        "enquadramento_tecnico": str(risk.get("ltcat_enquadramento_tecnico") or DEFAULT_LTCAT_ENQUADRAMENTO),
        "parecer_previdenciario": str(risk.get("ltcat_parecer_previdenciario") or DEFAULT_LTCAT_PARECER),
        "periodicidade_jornada": str(risk.get("ltcat_periodicidade_jornada") or "Mensal (<= 4 horas < 10% jornada)"),
    }


def _ltcat_set_repeated_cells(row, col_start: int, col_end: int, text_value: Any) -> None:
    for col in range(col_start, min(col_end + 1, len(row.cells))):
        _ltcat_set_cell_text(row.cells[col], text_value)


def _ltcat_fill_single_risk_block(table, start: int, risk: Mapping[str, Any], data_avaliacao: str) -> None:
    tipo = _normalize_option(risk.get("tipo_risco"))
    fill = TIPO_RISCO_COLORS.get(tipo)

    def row(offset: int):
        return table.rows[start + offset]

    _ltcat_set_repeated_cells(row(0), 3, 5, tipo)
    for col in range(3, min(6, len(row(0).cells))):
        _ltcat_shade_cell(row(0).cells[col], fill)
        _ltcat_set_cell_font_color(row(0).cells[col], "000000")

    _ltcat_set_repeated_cells(row(1), 3, 5, risk.get("risco", ""))
    _ltcat_set_repeated_cells(row(2), 3, 5, risk.get("meio_propagacao", ""))
    _ltcat_set_repeated_cells(row(8), 3, 5, risk.get("epis", ""))
    _ltcat_set_repeated_cells(row(9), 3, 5, risk.get("epcs", ""))

    if risk.get("periodicidade_jornada") and len(row(12).cells) > 1:
        _ltcat_set_cell_text(row(12).cells[1], risk.get("periodicidade_jornada", ""))
    _ltcat_set_repeated_cells(row(12), 5, 5, data_avaliacao)

    _ltcat_set_repeated_cells(row(14), 2, 5, risk.get("insalubridade", ""))
    _ltcat_set_repeated_cells(row(15), 2, 5, risk.get("grau_insalubridade", ""))
    _ltcat_set_repeated_cells(row(16), 2, 5, risk.get("enquadramento_tecnico", ""))
    _ltcat_set_repeated_cells(row(18), 2, 5, risk.get("aposentadoria_especial", ""))
    _ltcat_set_repeated_cells(row(19), 2, 5, risk.get("parecer_previdenciario", ""))

    replacements = {
        "{{TIPO DE RISCO}}": tipo,
        "{{risco}}": risk.get("risco", ""),
        "{{meiodepropagacao}}": risk.get("meio_propagacao", ""),
        "{{EPIs}}": risk.get("epis", ""),
        "{{EPCs}}": risk.get("epcs", ""),
        "{{datacriacaolaudo}}": data_avaliacao,
        "{{simounao}}": risk.get("insalubridade", ""),
        "{{graudeinsalubridade}}": risk.get("grau_insalubridade", ""),
    }
    for off in range(0, 25):
        if start + off >= len(table.rows):
            continue
        for cell in table.rows[start + off].cells:
            for paragraph in cell.paragraphs:
                _replace_text_in_paragraph(paragraph, replacements)


def _ltcat_fill_risk_table(table, sector: Mapping[str, Any], risks: list[Mapping[str, Any]], data_avaliacao: str) -> None:
    _ltcat_replace_in_table(table, {
        "{{SETOR}}": str(sector.get("setor", "")),
        "{{CARGOS}}": _ltcat_sector_cargos_text(sector),
    })
    original_block_rows = [deepcopy(table.rows[i]._tr) for i in range(4, len(table.rows))]
    for _ in range(max(0, len(risks) - 1)):
        for tr in original_block_rows:
            table._tbl.append(deepcopy(tr))
    block_size = len(original_block_rows)
    for index, risk in enumerate(risks):
        _ltcat_fill_single_risk_block(table, 4 + index * block_size, risk, data_avaliacao)


def _ltcat_fill_riscos_area(doc: Document, groups: list[Mapping[str, Any]], data_avaliacao: str) -> None:
    absence_tables = [t for t in doc.tables if _ltcat_table_has_text(t, "AUSÊNCIA DE RISCOS") and _ltcat_table_has_text(t, "{{SETOR}}")]
    risk_tables = [t for t in doc.tables if _ltcat_table_has_text(t, "{{meiodepropagacao}}") and _ltcat_table_has_text(t, "{{graudeinsalubridade}}")]
    if not absence_tables or not risk_tables:
        return
    absence_anchor = absence_tables[0]
    risk_anchor = risk_tables[0]
    absence_template = deepcopy(absence_anchor._tbl)
    risk_template = deepcopy(risk_anchor._tbl)

    _ltcat_remove_inventory_heading_paragraphs(doc)
    _ltcat_insert_inventory_heading_before(absence_anchor)

    for index, group in enumerate(groups):
        sector = group["sector"]
        if index > 0:
            _ltcat_insert_page_break_before(absence_anchor)
        raw_risks = group.get("risks") or []
        risks = [_ltcat_risk_from_common(risk) for risk in raw_risks]
        risks = [risk for risk in risks if _normalize_option(risk.get("tipo_risco")) in TIPOS_RISCO_LTCAT]
        if risks:
            new_tbl = deepcopy(risk_template)
            _ltcat_insert_tbl_before(absence_anchor, new_tbl)
            table = Table(new_tbl, doc)
            _ltcat_fill_risk_table(table, sector, risks, data_avaliacao)
        else:
            new_tbl = deepcopy(absence_template)
            _ltcat_insert_tbl_before(absence_anchor, new_tbl)
            table = Table(new_tbl, doc)
            _ltcat_replace_in_table(table, {
                "{{SETOR}}": str(sector.get("setor", "")),
                "{{CARGOS}}": _ltcat_sector_cargos_text(sector),
            })
    _ltcat_remove_table(absence_anchor)
    _ltcat_remove_table(risk_anchor)


def _ltcat_fill_signature(doc: Document, company: Mapping[str, str]) -> None:
    for table in doc.tables:
        if _ltcat_table_has_text(table, "RESPONSÁVEL LEGAL DA EMPRESA") and _ltcat_table_has_text(table, "{{empresa}}"):
            _ltcat_replace_in_table(table, {
                "{{empresa}}": company.get("empresa", ""),
                "{{DATACRIACAOLAUDO}}": company.get("data_criacao_laudo") or company.get("data_avaliacao") or company.get("data_atual", ""),
                "{{datacriacaolaudo}}": company.get("data_criacao_laudo") or company.get("data_avaliacao") or company.get("data_atual", ""),
            })


def _sanitize_ltcat_groups(groups: Iterable[Mapping[str, Any]]) -> list[dict[str, Any]]:
    clean_groups: list[dict[str, Any]] = []
    for group in groups:
        sector = _sanitize_sector(group.get("sector") or {})
        risks = [risk for risk in (group.get("risks") or []) if risk]
        if sector.get("setor"):
            clean_groups.append({"sector": sector, "risks": risks})
    return clean_groups


def generate_complete_ltcat_docx(
    groups: Iterable[Mapping[str, Any]],
    output_path: str | Path,
    empresa: str,
    cnpj: str,
    data_atual: str | None = None,
    data_final: str | None = None,
    company_extra: Mapping[str, Any] | None = None,
) -> Path:
    groups = _sanitize_ltcat_groups(groups)
    if not groups:
        raise ValueError("Selecione pelo menos um setor para gerar o LTCAT completo.")
    empresa = str(empresa or "").strip()
    cnpj = str(cnpj or "").strip()
    if not empresa:
        raise ValueError("Preencha o nome da empresa.")
    if not cnpj:
        raise ValueError("Preencha o CNPJ.")
    if not TEMPLATE_LTCAT_COMPLETO_PATH.exists():
        raise FileNotFoundError(f"Modelo LTCAT completo não encontrado: {TEMPLATE_LTCAT_COMPLETO_PATH}")

    data_atual = (data_atual or datetime.now().strftime("%m/%Y")).strip()
    data_final = (data_final or "").strip()
    company = _ltcat_company_dict(empresa, cnpj, data_atual, data_final, company_extra)

    doc = Document(str(TEMPLATE_LTCAT_COMPLETO_PATH))
    sectors = [group["sector"] for group in groups]
    data_avaliacao = company.get("data_avaliacao") or data_atual

    _insert_psychosocial_revision_line(doc, company)
    _replace_doc_placeholders(doc, _company_replacements(company))
    _ltcat_fill_company_table(doc, company, sectors)
    _ltcat_fill_relacao_funcoes(doc, sectors, data_atual, data_final)
    _ltcat_fill_descritivo_setores(doc, sectors)
    _ltcat_fill_riscos_area(doc, groups, data_avaliacao)
    _ltcat_fill_signature(doc, company)

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    save_docx_with_default_font(doc, output_path)
    return output_path

# ---------------------------------------------------------------------------
# AET - ANÁLISE ERGONÔMICA DO TRABALHO
# ---------------------------------------------------------------------------

def generate_aet_docx(groups: list[Mapping[str, Any]], output_path: Path, empresa: str = "", cnpj: str = "", data_atual: str = "", data_final: str = "", company: Mapping[str, Any] | None = None) -> Path:
    """Gera uma AET básica e editável a partir dos setores/riscos selecionados.

    O objetivo é entregar um módulo inicial de ergonomia usando a base já
    cadastrada no sistema. O documento é propositalmente conservador e deve ser
    revisado pelo responsável técnico quando a análise exigir medições, entrevistas
    ou observação presencial.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.shared import Inches, Pt

    company = company or {}
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)

    styles = doc.styles
    styles["Normal"].font.name = "Arial Narrow"
    styles["Normal"].font.size = Pt(10)

    def add_title(text: str, size: int = 18):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = True
        r.font.name = "Arial Narrow"
        r.font.size = Pt(size)
        return p

    def add_heading(text: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        r.bold = True
        r.font.name = "Arial Narrow"
        r.font.size = Pt(12)
        return p

    def add_text(text: str):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(text)
        r.font.name = "Arial Narrow"
        r.font.size = Pt(10)
        return p

    def set_table_font(table):
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs:
                        run.font.name = "Arial Narrow"
                        run.font.size = Pt(9)

    add_title("AET – ANÁLISE ERGONÔMICA DO TRABALHO", 17)
    add_title(empresa or "EMPRESA", 14)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Vigência: {data_atual or '-'} a {data_final or '-'}").bold = True
    doc.add_paragraph()

    add_heading("IDENTIFICAÇÃO DA EMPRESA")
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    ident = [
        ("EMPRESA", empresa or company.get("empresa", "")),
        ("CNPJ", cnpj or company.get("cnpj", "")),
        ("ENDEREÇO", company.get("endereco", "")),
        ("BAIRRO / CIDADE", company.get("bairro_cidade", "")),
        ("CEP", company.get("cep", "")),
        ("CNAE", company.get("cnae", company.get("cnae1", ""))),
        ("DESCRIÇÃO DA ATIVIDADE", company.get("descricao_atividade", company.get("descricao1", ""))),
        ("GRAU DE RISCO", company.get("grau_risco", company.get("grau1", ""))),
        ("FUNCIONÁRIOS", company.get("funcionarios", "")),
        ("VIGÊNCIA", f"{data_atual or ''} – {data_final or ''}"),
    ]
    for label, value in ident:
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = str(value or "")
        for run in row.cells[0].paragraphs[0].runs:
            run.bold = True
    set_table_font(table)

    add_heading("OBJETIVO")
    add_text("Esta Análise Ergonômica do Trabalho tem por objetivo registrar, de forma documental, as condições ergonômicas, organizacionais e psicossociais observadas nos setores selecionados, utilizando como base as informações cadastradas no sistema de SST da empresa.")

    add_heading("METODOLOGIA")
    add_text("A análise considera a relação função x atividade, os riscos ergonômicos e psicossociais vinculados aos setores, as possíveis fontes ou circunstâncias de exposição, os agravos esperados e as medidas preventivas/corretivas indicadas. Quando necessário, recomenda-se complementação por observação presencial, entrevistas, registros fotográficos e análise detalhada da tarefa.")

    add_heading("RELAÇÃO FUNÇÃO X ATIVIDADE")
    rel = doc.add_table(rows=1, cols=5)
    rel.style = "Table Grid"
    headers = ["SETOR", "CARGO", "CBO", "Nº FUNC.", "DESCRIÇÃO DA ATIVIDADE"]
    for i, h in enumerate(headers):
        rel.rows[0].cells[i].text = h
    for group in groups:
        sector = group.get("sector", {})
        cargos = sector.get("cargos", []) or []
        if not cargos:
            row = rel.add_row().cells
            row[0].text = sector.get("setor", "")
            row[1].text = "A DEFINIR"
            row[2].text = ""
            row[3].text = ""
            row[4].text = ""
        for cargo in cargos:
            row = rel.add_row().cells
            row[0].text = sector.get("setor", "")
            row[1].text = str(cargo.get("cargo", ""))
            row[2].text = str(cargo.get("cbo", ""))
            row[3].text = str(cargo.get("n_func", ""))
            row[4].text = str(cargo.get("descricao", ""))
    set_table_font(rel)

    doc.add_page_break()
    add_heading("ANÁLISE ERGONÔMICA POR SETOR")
    for group in groups:
        sector = group.get("sector", {})
        risks = [risk for risk in group.get("risks", []) if _normalize_option(risk.get("tipo_risco")) in {"ERGONÔMICO", "ERGONÔMICO PSICOSSOCIAL"}]
        add_heading(f"SETOR: {sector.get('setor', '')}")
        cargos = ", ".join([str(c.get("cargo", "")) for c in (sector.get("cargos", []) or []) if c.get("cargo")])
        if cargos:
            add_text(f"Cargos abrangidos: {cargos}.")
        if not risks:
            add_text("Não foram selecionados riscos ergonômicos ou psicossociais específicos para este setor no momento da geração. Recomenda-se manter acompanhamento das condições de trabalho e atualizar a AET quando houver alteração de processo, função, jornada ou queixas relacionadas.")
            continue
        tbl = doc.add_table(rows=1, cols=5)
        tbl.style = "Table Grid"
        for i, h in enumerate(["RISCO", "TIPO", "POSSÍVEIS AGRAVOS", "FONTES/CIRCUNSTÂNCIAS", "AÇÕES RECOMENDADAS"]):
            tbl.rows[0].cells[i].text = h
        for risk in risks:
            row = tbl.add_row().cells
            row[0].text = str(risk.get("risco", ""))
            row[1].text = str(risk.get("tipo_risco", ""))
            row[2].text = str(risk.get("possiveis_lesoes", ""))
            row[3].text = str(risk.get("fontes_circunstancias", ""))
            row[4].text = str(risk.get("acoes", ""))
        set_table_font(tbl)

    add_heading("CONCLUSÃO")
    add_text("Com base nas informações cadastradas e nos riscos selecionados, recomenda-se a implementação e o acompanhamento das medidas preventivas/corretivas indicadas, com prioridade para a adequação das condições de trabalho, organização das atividades, orientação dos trabalhadores e controle de fatores psicossociais quando identificados. Esta AET deve ser revisada sempre que houver mudanças significativas no processo de trabalho, organização, layout, mobiliário, número de trabalhadores, queixas recorrentes ou inclusão de novos riscos.")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Responsável Técnico ___________________________________________").bold = True

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    save_docx_with_default_font(doc, output_path)
    return output_path


# ---------------------------------------------------------------------------
# AET COMPLETA - modelo gerado pelo sistema com cabeçalho/rodapé do PGR
# ---------------------------------------------------------------------------

def _clear_document_body_keep_section(doc: Document) -> None:
    body = doc._element.body
    sect_pr = None
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            sect_pr = deepcopy(child)
        body.remove(child)
    if sect_pr is not None:
        body.append(sect_pr)


def _aet_list_text(values: Any, fallback: str = "Não informado") -> str:
    if isinstance(values, (list, tuple)):
        cleaned = [str(v).strip() for v in values if str(v).strip()]
        return ", ".join(cleaned) if cleaned else fallback
    value = str(values or "").strip()
    return value or fallback


def _aet_risk_recommendation(risk: Mapping[str, Any]) -> str:
    action = str(risk.get("acoes", "") or "").strip()
    if action:
        return action
    risk_name = str(risk.get("risco", "") or "").lower()
    if "postura" in risk_name or "sentado" in risk_name or "pé" in risk_name:
        return "Promover alternância postural, pausas programadas e orientação ergonômica.\nTREINAMENTO DE NR – NR-17"
    if "repet" in risk_name:
        return "Promover pausas, rodízio de tarefas e adequação do ritmo de trabalho.\nTREINAMENTO DE NR – NR-17"
    if "psicossocial" in str(risk.get("tipo_risco", "")).lower() or "sobrecarga" in risk_name or "assédio" in risk_name:
        return "Fortalecer canais de comunicação, acompanhar demandas e orientar liderança e trabalhadores.\nTREINAMENTO DE NR – NR-01 / NR-17"
    return "Manter acompanhamento das condições de trabalho e implantar medidas preventivas conforme necessidade."


def _aet_auto_sector_conclusion(sector_name: str, risks: list[Mapping[str, Any]], data: Mapping[str, Any]) -> str:
    manual = str(data.get("conclusao_setor", "") or "").strip()
    if manual:
        return manual
    has_psy = any(_normalize_option(r.get("tipo_risco")) == "ERGONÔMICO PSICOSSOCIAL" for r in risks)
    has_erg = any(_normalize_option(r.get("tipo_risco")) == "ERGONÔMICO" for r in risks)
    if has_psy and has_erg:
        return f"No setor {sector_name}, foram identificados fatores ergonômicos e psicossociais relacionados à organização e execução das atividades, recomendando-se acompanhamento contínuo, orientação dos trabalhadores e implantação das medidas preventivas descritas."
    if has_psy:
        return f"No setor {sector_name}, foram identificados fatores psicossociais associados à organização do trabalho, comunicação, demandas e relações laborais, recomendando-se acompanhamento gerencial e ações preventivas específicas."
    if has_erg:
        return f"No setor {sector_name}, foram identificados fatores ergonômicos compatíveis com a rotina laboral, recomendando-se adequações de postura, pausas, organização do posto e acompanhamento das queixas dos trabalhadores."
    return f"No setor {sector_name}, não foram selecionados fatores ergonômicos ou psicossociais específicos no momento da geração, recomendando-se manutenção do acompanhamento das condições de trabalho."



def _aet_sector_diagnostic(sector_name: str, risks: list[Mapping[str, Any]], data: Mapping[str, Any]) -> str:
    """Gera diagnóstico técnico sintético do setor para a AET."""
    postura = _aet_list_text(data.get("postura_predominante"), "postura não definida")
    exig_f = str(data.get("exigencia_fisica") or "a definir").lower()
    exig_c = str(data.get("exigencia_cognitiva") or "a definir").lower()
    atividade = str(data.get("tipo_atividade") or "atividade informada no cadastro de cargos").strip()
    fatores = _aet_list_text(data.get("fatores_organizacionais"), "fatores organizacionais não detalhados")
    has_psy = any(_normalize_option(r.get("tipo_risco")) == "ERGONÔMICO PSICOSSOCIAL" for r in risks)
    has_erg = any(_normalize_option(r.get("tipo_risco")) == "ERGONÔMICO" for r in risks)
    base = (
        f"O setor {sector_name} apresenta atividade predominante de {atividade}, com postura predominante: {postura}. "
        f"A exigência física foi classificada como {exig_f} e a exigência cognitiva como {exig_c}. "
        f"Foram considerados os seguintes fatores organizacionais: {fatores}."
    )
    if has_psy and has_erg:
        return base + " A análise indica necessidade de acompanhamento integrado dos fatores biomecânicos, organizacionais e psicossociais, com implantação das recomendações previstas nesta AET."
    if has_psy:
        return base + " A análise indica atenção aos aspectos de organização do trabalho, comunicação, autonomia, demandas e relações laborais, com acompanhamento preventivo dos fatores psicossociais."
    if has_erg:
        return base + " A análise indica necessidade de controle dos fatores ergonômicos relacionados à postura, ritmo, repetitividade, mobiliário, ferramentas e organização da atividade."
    return base + " Não foram selecionados fatores ergonômicos específicos, recomendando-se manutenção das boas práticas e reavaliação em caso de alterações do processo ou surgimento de queixas."

def _aet_priority_from_risks(risks: list[Mapping[str, Any]], manual: str = "") -> str:
    if manual:
        return manual
    levels = {_normalize_option(r.get("grau_nivel_risco")) for r in risks}
    if "MUITO ALTO" in levels or "ALTO" in levels:
        return "Alta"
    if "MODERADO" in levels:
        return "Média"
    return "Baixa"


def _aet_norm(value: Any) -> str:
    text = unicodedata.normalize("NFD", str(value or ""))
    text = "".join(ch for ch in text if unicodedata.category(ch) != "Mn")
    return re.sub(r"\s+", " ", text).strip().lower()


def _aet_company_profile(company_doc: Mapping[str, Any]) -> dict[str, Any]:
    haystack = _aet_norm(" ".join([
        str(company_doc.get("empresa", "")),
        str(company_doc.get("cnae", "")),
        str(company_doc.get("cnae_secundario", "")),
        str(company_doc.get("descricao_atividade", "")),
        str(company_doc.get("descricao_atividade_secundaria", "")),
    ]))
    for profile in AET_ACTIVITY_PROFILES:
        if any(_aet_norm(keyword) in haystack for keyword in profile.get("keywords", [])):
            return profile
    return {
        "label": "Modelo geral por atividade econômica",
        "contexto": "atividade econômica informada no cadastro da empresa, com rotinas administrativas, operacionais e/ou de atendimento conforme setores e cargos cadastrados no sistema.",
        "demanda": "avaliar as condições de trabalho, considerando as características reais dos setores, cargos, riscos ergonômicos, riscos psicossociais e informações complementares preenchidas na geração.",
        "tarefas": "execução das atividades descritas nos cargos cadastrados, observando a divisão setorial, a organização do trabalho e as exigências específicas de cada função.",
        "exigencias": "exigências físicas, cognitivas, organizacionais e psicossociais compatíveis com os setores e riscos selecionados.",
        "favoraveis": "existência de cadastro setorial, relação função x atividade e integração com os dados do PGR, PCMSO e LTCAT.",
        "recomendacoes": "manter avaliação periódica, orientar trabalhadores, acompanhar queixas e implantar as medidas preventivas/corretivas registradas nesta AET.",
    }


def _aet_default_sector_data(profile: Mapping[str, Any], risks: list[Mapping[str, Any]]) -> dict[str, Any]:
    label = _aet_norm(profile.get("label"))
    risk_names = _aet_norm(" ".join(str(r.get("risco", "")) for r in risks))

    if any(term in label for term in ("construcao", "agropecuaria")):
        postura = ["Em pé", "Deslocamento frequente", "Inclinação/flexão de tronco", "Agachamento eventual"]
        exig_fisica = "Elevada"
        carga = "Habitual"
    elif any(term in label for term in ("administrativo", "escritorio")):
        postura = ["Sentado", "Alternado"]
        exig_fisica = "Baixa"
        carga = "Não habitual"
    elif any(term in label for term in ("transporte", "taxi")):
        postura = ["Sentado", "Alternado"]
        exig_fisica = "Moderada"
        carga = "Eventual"
    else:
        postura = ["Em pé", "Sentado", "Alternado"]
        exig_fisica = "Moderada"
        carga = "Eventual"

    if any(term in label for term in ("funer", "educacao", "saude", "condominio", "transporte")):
        exig_cognitiva = "Elevada"
    else:
        exig_cognitiva = "Moderada"

    atendimento = "Frequente" if any(term in label for term in ("comercio", "funer", "saude", "educacao", "condominio", "alimentacao", "taxi")) else "Eventual"
    repetitividade = "Moderada frequência" if any(term in risk_names for term in ("repet", "digit", "caixa", "postura")) else "Baixa frequência"
    prioridade = _aet_priority_from_risks(risks)

    return {
        "postura_predominante": postura,
        "tipo_atividade": profile.get("tarefas") or "Atividades compatíveis com os cargos cadastrados",
        "exigencia_fisica": exig_fisica,
        "exigencia_cognitiva": exig_cognitiva,
        "levantamento_cargas": carga,
        "movimentos_repetitivos": repetitividade,
        "atencao_concentracao": exig_cognitiva,
        "atendimento_publico": atendimento,
        "autonomia": "Adequada" if prioridade == "Baixa" else "Parcialmente limitada",
        "metas_prioridades": "Demandas variáveis" if prioridade != "Baixa" else "Compatíveis",
        "comunicacao": "Necessita melhoria" if any(_normalize_option(r.get("tipo_risco")) == "ERGONÔMICO PSICOSSOCIAL" for r in risks) else "Adequada",
        "ritmo_trabalho": "Por demanda da atividade, fluxo operacional e prioridades do setor",
        "pausas": "Pausas conforme escala, intensidade da atividade e necessidade de recuperação física/mental",
        "mobiliario": "A avaliar/manter adequado à atividade, considerando assentos, bancadas, mesas, equipamentos e alcances",
        "ambiente": "A avaliar iluminação, ventilação, conforto térmico, ruído, circulação e organização do espaço",
        "organizacao": profile.get("demanda") or "Rotina organizada conforme demandas do setor e comunicação com a liderança",
        "equipamentos": "Equipamentos, ferramentas e materiais compatíveis com as atividades cadastradas",
        "fatores_organizacionais": ["Organização de demandas", "Comunicação", "Atenção contínua"],
        "medidas_recomendadas": ["Orientação NR-17", "Pausas breves", "Organização do posto", "Melhoria de comunicação"],
        "prioridade": prioridade,
        "prazo": "30 dias" if prioridade in {"Alta", "Média"} else "60 dias",
        "responsavel": "Empresa / Administração / SST",
    }


def _aet_merge_sector_data(defaults: Mapping[str, Any], manual: Mapping[str, Any]) -> dict[str, Any]:
    merged = dict(defaults)
    for key, value in (manual or {}).items():
        if isinstance(value, (list, tuple)):
            cleaned = [str(v).strip() for v in value if str(v).strip()]
            if cleaned:
                merged[key] = cleaned
        elif str(value or "").strip():
            merged[key] = value
    return merged


def _aet_sector_names(groups: list[Mapping[str, Any]]) -> str:
    names = []
    for group in groups:
        name = str((group.get("sector") or {}).get("setor", "")).strip()
        if name and name not in names:
            names.append(name)
    return ", ".join(names) if names else "setores selecionados"


def _aet_selected_risk_names(risks: list[Mapping[str, Any]], fallback: str = "riscos ergonômicos selecionados") -> str:
    names = []
    for risk in risks:
        name = str(risk.get("risco", "") or "").strip()
        if name and name not in names:
            names.append(name)
    if not names:
        return fallback
    if len(names) <= 6:
        return ", ".join(names)
    return ", ".join(names[:6]) + f", entre outros {len(names) - 6} fator(es)"


def _aet_critical_sector_text(groups: list[Mapping[str, Any]]) -> str:
    critical = []
    psychosocial = []
    for group in groups:
        sector_name = str((group.get("sector") or {}).get("setor", "") or "").strip()
        risks = group.get("risks", []) or []
        if any(_normalize_option(r.get("grau_nivel_risco")) in {"ALTO", "MUITO ALTO"} for r in risks):
            critical.append(sector_name)
        if any(_normalize_option(r.get("tipo_risco")) == "ERGONÔMICO PSICOSSOCIAL" for r in risks):
            psychosocial.append(sector_name)
    parts = []
    if critical:
        parts.append("setores com maior prioridade técnica: " + ", ".join(dict.fromkeys([s for s in critical if s])))
    if psychosocial:
        parts.append("setores com fatores psicossociais selecionados: " + ", ".join(dict.fromkeys([s for s in psychosocial if s])))
    return "; ".join(parts) if parts else "não foram indicados setores com criticidade elevada no momento da geração"


def generate_aet_docx(groups: list[Mapping[str, Any]], output_path: Path, empresa: str = "", cnpj: str = "", data_atual: str = "", data_final: str = "", company: Mapping[str, Any] | None = None) -> Path:
    """Gera AET completa com formulário detalhado, dados de PGR/PCMSO/LTCAT e padrão visual da clínica.

    Usa o modelo do PGR apenas como base de cabeçalho/rodapé e identidade visual,
    criando o corpo da AET de forma estruturada e editável.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.shared import Inches, Pt, RGBColor

    if not groups:
        raise ValueError("Selecione pelo menos um setor para gerar a AET.")
    company = company or {}
    company_doc = _company_dict_for_docs(empresa, cnpj, data_atual, data_final, company)
    activity_profile = _aet_company_profile(company_doc)
    aet_data = company.get("aet") or {}
    general = aet_data.get("general", {}) if isinstance(aet_data, Mapping) else {}
    by_sector = aet_data.get("by_sector", {}) if isinstance(aet_data, Mapping) else {}

    if TEMPLATE_AET_BASE_PATH.exists():
        doc = Document(str(TEMPLATE_AET_BASE_PATH))
    elif TEMPLATE_PGR_COMPLETO_PATH.exists():
        doc = Document(str(TEMPLATE_PGR_COMPLETO_PATH))
    else:
        doc = Document()
    _clear_document_body_keep_section(doc)
    _replace_doc_placeholders(doc, _company_replacements(company_doc))

    for section in doc.sections:
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)

    styles = doc.styles
    styles["Normal"].font.name = "Arial Narrow"
    styles["Normal"].font.size = Pt(10)

    def _run_font(run, size=10, bold=False, color=None):
        run.font.name = "Arial Narrow"
        run.font.size = Pt(size)
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor.from_string(color)

    def _set_paragraph_border(paragraph, position="bottom", color="6B1D1D", size="12"):
        p_pr = paragraph._p.get_or_add_pPr()
        p_bdr = p_pr.find(qn("w:pBdr"))
        if p_bdr is None:
            p_bdr = OxmlElement("w:pBdr")
            p_pr.append(p_bdr)
        el = p_bdr.find(qn(f"w:{position}"))
        if el is None:
            el = OxmlElement(f"w:{position}")
            p_bdr.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "1")
        el.set(qn("w:color"), color)

    def _set_table_borders(table, color="808080", size="4"):
        for row in table.rows:
            for cell in row.cells:
                tc_pr = cell._tc.get_or_add_tcPr()
                tc_borders = tc_pr.find(qn("w:tcBorders"))
                if tc_borders is None:
                    tc_borders = OxmlElement("w:tcBorders")
                    tc_pr.append(tc_borders)
                for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                    tag = qn(f"w:{edge}")
                    el = tc_borders.find(tag)
                    if el is None:
                        el = OxmlElement(f"w:{edge}")
                        tc_borders.append(el)
                    el.set(qn("w:val"), "single")
                    el.set(qn("w:sz"), size)
                    el.set(qn("w:space"), "0")
                    el.set(qn("w:color"), color)

    def _apply_aet_header_footer():
        for section in doc.sections:
            for part in (section.header, section.first_page_header, section.even_page_header):
                try:
                    for child in list(part._element):
                        part._element.remove(child)
                    p = part.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _run_font(p.add_run("AET – ANÁLISE ERGONÔMICA DO TRABALHO"), 11, True)
                    p2 = part.add_paragraph()
                    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _run_font(p2.add_run(company_doc.get("empresa", "")), 10, True)
                    _set_paragraph_border(p2, "bottom", "6B1D1D", "12")
                except Exception:
                    pass
            for part in (section.footer, section.first_page_footer, section.even_page_footer):
                try:
                    for child in list(part._element):
                        part._element.remove(child)
                    p = part.add_paragraph()
                    _set_paragraph_border(p, "top", "6B1D1D", "12")
                    _run_font(p.add_run("ELABORAÇÃO\nFONE: (91) 98354-0469 / 98354-0444 / 3349-6948\n(96) 3223-7946 / 3222-7682\nEMAIL: EDGESETORTECNICO@HOTMAIL.COM"), 8, True)
                except Exception:
                    pass

    def add_title(text: str, size: int = 18):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        _run_font(r, size=size, bold=True)
        return p

    def add_heading(text: str, level: int = 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8 if level == 1 else 5)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(text)
        _run_font(r, size=12 if level == 1 else 10, bold=True)
        return p

    def add_text(text: str):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(str(text or ""))
        _run_font(r, size=10)
        return p

    def set_table_font(table, header: bool = True):
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_table_borders(table)
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs:
                        _run_font(run, size=8 if len(table.columns) >= 5 else 9, bold=(header and row_idx == 0))

    def add_kv_table(items):
        tbl = doc.add_table(rows=0, cols=2)
        
        try:
            tbl.style = "Table Grid"
        except KeyError:
            pass
        for label, value in items:
            row = tbl.add_row().cells
            row[0].text = str(label)
            row[1].text = str(value or "")
        set_table_font(tbl)
        return tbl

    sectors = _unique_sectors_from_groups(groups)
    all_erg_risks = []
    for group in groups:
        for risk in group.get("risks", []) or []:
            if _normalize_option(risk.get("tipo_risco")) in {"ERGONÔMICO", "ERGONÔMICO PSICOSSOCIAL"}:
                all_erg_risks.append(risk)

    # CAPA
    add_title("AET – ANÁLISE ERGONÔMICA DO TRABALHO", 18)
    add_title("NR 17", 15)
    add_title(company_doc.get("empresa") or empresa or "EMPRESA", 14)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run_font(p.add_run(f"Data do documento: {company_doc.get('data_atual', '')}"), 11, True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run_font(p.add_run(f"Revisão periódica: {company_doc.get('data_atual', '')} à {company_doc.get('data_final', '')}"), 11, True)
    if company_doc.get("ajuste_psicossocial") == "1" and company_doc.get("data_da_revisao"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run_font(p.add_run(f"Revisão periódica de inclusão de Risco Psicossocial: {company_doc.get('data_da_revisao')}"), 11, True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run_font(p.add_run(f"Data de elaboração da AET: {company_doc.get('data_criacao_laudo') or company_doc.get('data_avaliacao') or company_doc.get('data_atual', '')}"), 10, True)

    doc.add_page_break()
    add_heading("IDENTIFICAÇÃO DA EMPRESA")
    add_kv_table([
        ("EMPRESA", company_doc.get("empresa")),
        ("ENDEREÇO", company_doc.get("endereco")),
        ("BAIRRO / CIDADE", company_doc.get("bairro_cidade")),
        ("CEP", company_doc.get("cep")),
        ("CNPJ", company_doc.get("cnpj")),
        ("CNAE", company_doc.get("cnae")),
        ("DESCRIÇÃO DA ATIVIDADE", company_doc.get("descricao_atividade")),
        ("GRAU DE RISCO", company_doc.get("grau_risco")),
        ("CNAE (SECUNDÁRIO)", company_doc.get("cnae_secundario")),
        ("DESCRIÇÃO DA ATIVIDADE (SECUNDÁRIA)", company_doc.get("descricao_atividade_secundaria")),
        ("GRAU DE RISCO (SECUNDÁRIO)", company_doc.get("grau_risco_secundario")),
        ("FUNCIONÁRIOS", company_doc.get("funcionarios") or _total_funcionarios_from_sectors(sectors)),
        ("VIGÊNCIA", f"{company_doc.get('data_atual', '')} – {company_doc.get('data_final', '')}"),
        ("EMAIL", company_doc.get("email")),
        ("FONE", company_doc.get("fone")),
    ])

    add_heading("RESPONSABILIDADE TÉCNICA")
    responsavel = str(general.get("responsavel_tecnico", "") or "Responsável técnico a definir pela clínica").strip()
    add_text(f"O presente documento foi elaborado com base nas informações fornecidas pela empresa, nos setores e cargos cadastrados, nos riscos ergonômicos e psicossociais selecionados no sistema e nos dados complementares preenchidos no formulário de AET. Responsável técnico: {responsavel}.")

    add_heading("INTRODUÇÃO")
    total_func = company_doc.get("funcionarios") or _total_funcionarios_from_sectors(sectors)
    add_text(
        f"A presente Análise Ergonômica do Trabalho foi desenvolvida com o objetivo de avaliar as condições reais de trabalho da empresa "
        f"{company_doc.get('empresa')}, considerando os aspectos físicos, cognitivos, organizacionais e psicossociais previstos na NR-17 e integrados ao gerenciamento de riscos da NR-01. "
        f"A empresa atua em {activity_profile.get('contexto')} O quadro analisado contempla {total_func or 'os'} trabalhador(es) distribuídos nos setores: {_aet_sector_names(groups)}."
    )
    add_text(
        "A elaboração da AET considera os dados específicos obtidos no processo de geração dos laudos, incluindo identificação da empresa, CNAE, relação função x atividade, "
        "riscos selecionados por setor, fontes/circunstâncias, possíveis agravos à saúde, ações preventivas/corretivas e informações preenchidas no formulário ergonômico."
    )

    add_heading("CARACTERIZAÇÃO DO DOCUMENTO")
    add_kv_table([
        ("Tipo de documento", general.get("tipo_documento") or general.get("tipo_aet") or "AET - Análise Ergonômica do Trabalho"),
        ("Modelo técnico aplicado", activity_profile.get("label")),
        ("Motivo da análise", general.get("motivo_analise") or activity_profile.get("demanda") or "Atendimento à NR-17 e integração com o PGR"),
        ("Condição ergonômica geral", general.get("condicao_ergonomica_geral") or "Adequada com recomendações"),
    ])

    add_heading("OBJETIVO")
    objetivo_extra = str(general.get("objetivo_complementar", "") or "").strip()
    add_text("A presente Análise Ergonômica do Trabalho tem por objetivo avaliar as condições de trabalho, considerando aspectos relacionados à organização do trabalho, exigências físicas, cognitivas, biomecânicas, ambientais e psicossociais, visando propor medidas de adequação, prevenção e melhoria das condições laborais, conforme diretrizes da NR-17." + (f" {objetivo_extra}" if objetivo_extra else ""))

    add_heading("OBJETIVOS ESPECÍFICOS")
    add_text(
        "Identificar e caracterizar os fatores ergonômicos presentes nos setores avaliados; verificar a compatibilidade entre as atividades realizadas, os postos de trabalho e a organização do trabalho; "
        "analisar riscos ergonômicos e psicossociais registrados no PGR; propor medidas preventivas e corretivas; e subsidiar o plano de ação da empresa com recomendações objetivas, responsáveis, prazos e indicadores de acompanhamento."
    )

    add_heading("METODOLOGIA")
    metodologias = general.get("metodologia") or []
    met_text = _aet_list_text(metodologias, "Levantamento documental, análise dos setores/cargos cadastrados, avaliação dos riscos selecionados e preenchimento do formulário ergonômico por setor")
    add_text(f"A análise foi realizada considerando: {met_text}. Foram integradas ao documento as informações já existentes nos módulos de PGR, PCMSO e LTCAT, especialmente identificação da empresa, relação função x atividade, riscos cadastrados, possíveis agravos, fontes/circunstâncias e medidas preventivas/corretivas.")
    origens = general.get("origem_dados") or []
    if origens:
        add_text(f"Fontes de dados consideradas: {_aet_list_text(origens)}.")
    criterios = str(general.get("criterios_analise", "") or "").strip()
    if criterios:
        add_text(f"Critérios complementares informados: {criterios}")
    limitacoes = str(general.get("limitacoes_analise", "") or "").strip()
    add_heading("LIMITAÇÕES DA ANÁLISE")
    add_text(limitacoes or "A análise foi realizada com base nas informações fornecidas pela empresa, documentos disponíveis, setores/cargos cadastrados e fatores de risco identificados no momento da elaboração. Recomenda-se atualização sempre que houver alteração relevante nas atividades, layout, organização do trabalho, jornada, mobiliário, equipamentos ou surgimento de novas queixas.")

    add_heading("ANÁLISE DA DEMANDA")
    add_text(
        f"A demanda para realização da AET está relacionada à necessidade de avaliar tecnicamente as condições ergonômicas da empresa considerando seu CNAE, sua atividade econômica e os setores efetivamente selecionados. "
        f"Para este perfil, a análise busca {str(activity_profile.get('demanda') or '').rstrip('.')}."
    )
    add_text(
        f"A leitura setorial indica {_aet_critical_sector_text(groups)}. Essa abordagem permite tratar a ergonomia de forma direcionada, evitando conclusões genéricas e aproximando o documento da realidade operacional da empresa."
    )

    add_heading("ANÁLISE DA TAREFA")
    add_heading("5.1 FUNÇÕES E DESCRIÇÃO DAS ATIVIDADES", 2)
    add_text(
        f"As funções avaliadas estão distribuídas nos setores {_aet_sector_names(groups)}. De forma geral, as atividades envolvem {activity_profile.get('tarefas')} "
        "A relação função x atividade apresentada a seguir detalha os cargos, CBOs, número de trabalhadores e descrições cadastradas no sistema."
    )
    add_heading("5.2 TAREFAS PRESCRITAS", 2)
    add_text(
        "As tarefas prescritas correspondem às atividades formais esperadas para cada setor/cargo, conforme cadastro realizado no sistema e informações documentais disponíveis. "
        f"No perfil analisado, predominam tarefas relacionadas a {activity_profile.get('tarefas')}"
    )
    add_heading("5.3 EXIGÊNCIAS DA TAREFA", 2)
    add_text(
        f"As exigências da tarefa abrangem {activity_profile.get('exigencias')} Além disso, são considerados os fatores de risco ergonômico selecionados: "
        f"{_aet_selected_risk_names(all_erg_risks)}."
    )

    doc.add_page_break()
    add_heading("RELAÇÃO FUNÇÃO X ATIVIDADE")
    rel = doc.add_table(rows=1, cols=5)
    
    try:
        rel.style = "Table Grid"
    except KeyError:
        pass
    headers = ["SETOR", "CARGO", "CBO", "Nº FUNC.", "DESCRIÇÃO DA ATIVIDADE"]
    for i, h in enumerate(headers):
        rel.rows[0].cells[i].text = h
    for sector in sectors:
        cargos = sector.get("cargos", []) or []
        if not cargos:
            row = rel.add_row().cells
            row[0].text = sector.get("setor", "")
        for cargo in cargos:
            row = rel.add_row().cells
            row[0].text = sector.get("setor", "")
            row[1].text = str(cargo.get("cargo", ""))
            row[2].text = str(cargo.get("cbo", ""))
            row[3].text = str(cargo.get("n_func", ""))
            row[4].text = str(cargo.get("descricao", ""))
    set_table_font(rel)

    add_heading("6. ANÁLISE DA ATIVIDADE REAL")
    add_text(
        "A análise da atividade real considera que a execução cotidiana das tarefas pode variar em relação ao trabalho prescrito, principalmente em função de demandas simultâneas, ritmo operacional, comunicação, atendimento ao público, disponibilidade de recursos, organização das prioridades e características individuais dos trabalhadores."
    )
    add_text(
        f"No contexto da empresa, observam-se exigências compatíveis com {activity_profile.get('contexto')} A atividade real deve ser acompanhada continuamente para verificar se as medidas propostas permanecem adequadas ao processo de trabalho."
    )

    add_heading("6.1 ANÁLISE DAS CONDIÇÕES DE TRABALHO")
    add_text(
        "As condições de trabalho foram analisadas considerando os setores selecionados, os cargos cadastrados, as descrições de atividade, os fatores de risco ergonômico e psicossocial e os dados preenchidos no formulário da AET. "
        "Quando algum campo não foi detalhado pelo usuário, o sistema utilizou preenchimento técnico conservador com base no CNAE e no tipo de atividade, mantendo a necessidade de validação pelo responsável técnico."
    )

    add_heading("7. MOBILIÁRIO, POSTO DE TRABALHO E ORGANIZAÇÃO")
    add_text(
        "O mobiliário, os postos de trabalho, equipamentos, ferramentas e a organização das rotinas devem ser compatíveis com a natureza das atividades desenvolvidas, possibilitando postura adequada, alcance seguro, redução de esforços desnecessários e organização das demandas."
    )
    add_text(
        "A organização do trabalho deve considerar ritmo, pausas, comunicação, autonomia, metas, prioridades, suporte da liderança e prevenção de conflitos, especialmente quando houver risco ergonômico psicossocial selecionado para o setor."
    )

    add_heading("8. CARGA DE TRABALHO")
    add_text(
        f"A carga de trabalho analisada envolve dimensões físicas, cognitivas, organizacionais e psicossociais. Para o perfil {activity_profile.get('label')}, destacam-se: {activity_profile.get('exigencias')}"
    )
    add_text(
        "A priorização das ações deve observar a gravidade dos riscos cadastrados, a possibilidade de ocorrência, a quantidade de trabalhadores expostos, a existência de queixas e a presença de fatores psicossociais."
    )

    add_heading("9. FATORES ERGONÔMICOS IDENTIFICADOS")
    add_text(
        f"Foram considerados os seguintes fatores principais: {_aet_selected_risk_names(all_erg_risks)}. "
        "Esses fatores são detalhados por setor nos quadros seguintes, com indicação de possíveis impactos, fontes/circunstâncias e recomendações."
    )

    add_heading("9.1 FATORES FAVORÁVEIS")
    add_text(
        f"Como fatores favoráveis, considera-se {activity_profile.get('favoraveis')} Esses elementos contribuem para a implementação de medidas de controle, desde que acompanhados por registros, orientação dos trabalhadores e revisão periódica."
    )

    add_heading("10. DIAGNÓSTICO ERGONÔMICO")
    add_text(
        f"O diagnóstico ergonômico aponta condição geral {general.get('condicao_ergonomica_geral') or 'adequada com recomendações'}, com necessidade de acompanhamento dos setores selecionados e implantação das medidas indicadas. "
        f"Resumo técnico da criticidade: {_aet_critical_sector_text(groups)}."
    )

    add_heading("11. RECOMENDAÇÕES ERGONÔMICAS")
    add_heading("11.1 IMEDIATAS", 2)
    add_text(activity_profile.get("recomendacoes") or "Implantar orientações ergonômicas, pausas, organização do posto e acompanhamento das condições de trabalho.")
    add_heading("11.2 MÉDIO PRAZO", 2)
    add_text("Revisar fluxos de trabalho, adequar mobiliário/equipamentos quando necessário, formalizar procedimentos, melhorar comunicação e acompanhar indicadores como queixas, afastamentos, rotatividade e registros de incidentes.")
    add_heading("11.3 ESTRUTURAIS", 2)
    add_text("Integrar a AET ao PGR, manter monitoramento periódico dos fatores ergonômicos e psicossociais, capacitar lideranças e trabalhadores e revisar a organização do trabalho sempre que houver mudança de processo, layout, jornada ou quadro funcional.")

    doc.add_page_break()
    add_heading("12. PLANO DE AÇÃO ERGONÔMICO / PSICOSSOCIAL POR SETOR")
    for group in groups:
        sector = group.get("sector", {}) or {}
        sector_id = sector.get("id", "")
        sector_name = sector.get("setor", "")
        risks = [risk for risk in group.get("risks", []) or [] if _normalize_option(risk.get("tipo_risco")) in {"ERGONÔMICO", "ERGONÔMICO PSICOSSOCIAL"}]
        manual_sector_data = by_sector.get(sector_id, {}) if isinstance(by_sector, Mapping) else {}
        sector_data = _aet_merge_sector_data(_aet_default_sector_data(activity_profile, risks), manual_sector_data)
        cargos = ", ".join([str(c.get("cargo", "")) for c in (sector.get("cargos", []) or []) if c.get("cargo")])
        descricoes = [str(c.get("descricao", "")) for c in (sector.get("cargos", []) or []) if c.get("descricao")]

        add_heading(f"SETOR: {sector_name}", 1)
        if cargos:
            add_text(f"Cargos abrangidos: {cargos}.")
        if descricoes:
            add_text("Características das atividades: " + " ".join(descricoes))

        info = [
            ("Tipo de atividade predominante", sector_data.get("tipo_atividade") or "A definir"),
            ("Postura predominante", _aet_list_text(sector_data.get("postura_predominante"), "A definir conforme observação da atividade")),
            ("Exigência física", sector_data.get("exigencia_fisica") or "A definir"),
            ("Exigência cognitiva", sector_data.get("exigencia_cognitiva") or "A definir"),
            ("Levantamento/transporte de cargas", sector_data.get("levantamento_cargas") or "A definir"),
            ("Movimentos repetitivos", sector_data.get("movimentos_repetitivos") or "A definir"),
            ("Atenção/concentração", sector_data.get("atencao_concentracao") or "A definir"),
            ("Atendimento ao público", sector_data.get("atendimento_publico") or "A definir"),
            ("Autonomia", sector_data.get("autonomia") or "A definir"),
            ("Metas/prioridades", sector_data.get("metas_prioridades") or "A definir"),
            ("Comunicação", sector_data.get("comunicacao") or "A definir"),
            ("Ritmo de trabalho", sector_data.get("ritmo_trabalho") or "Compatível com a rotina operacional informada"),
            ("Pausas/recuperação", sector_data.get("pausas") or "Pausas conforme organização interna e necessidade da atividade"),
            ("Mobiliário/posto", sector_data.get("mobiliario") or "A avaliar/manter adequado à atividade"),
            ("Condições ambientais", sector_data.get("ambiente") or "Condições ambientais devem ser mantidas em níveis adequados de conforto"),
            ("Organização do trabalho", sector_data.get("organizacao") or "Rotina organizada conforme demandas do setor"),
            ("Equipamentos/ferramentas", sector_data.get("equipamentos") or "Equipamentos compatíveis com as atividades cadastradas"),
            ("Fatores organizacionais observados", _aet_list_text(sector_data.get("fatores_organizacionais"), "A definir")),
            ("Medidas ergonômicas recomendadas", _aet_list_text(sector_data.get("medidas_recomendadas"), "A definir conforme riscos")),
            ("Queixas/observações", sector_data.get("queixas") or sector_data.get("observacoes") or "Não informado"),
            ("Prioridade", _aet_priority_from_risks(risks, sector_data.get("prioridade", ""))),
        ]
        add_kv_table(info)

        add_heading("Diagnóstico ergonômico do setor", 2)
        add_text(_aet_sector_diagnostic(sector_name, risks, sector_data))

        add_heading("Fatores ergonômicos e psicossociais identificados", 2)
        if risks:
            tbl = doc.add_table(rows=1, cols=5)
            try:
                tbl.style = "Table Grid"
            except KeyError:
                pass
            for i, h in enumerate(["FATOR/RISCO", "TIPO", "POSSÍVEIS IMPACTOS", "FONTES/CIRCUNSTÂNCIAS", "RECOMENDAÇÕES"]):
                tbl.rows[0].cells[i].text = h
            for risk in risks:
                row = tbl.add_row().cells
                row[0].text = str(risk.get("risco", ""))
                row[1].text = str(risk.get("tipo_risco", ""))
                row[2].text = str(risk.get("possiveis_lesoes", ""))
                row[3].text = str(risk.get("fontes_circunstancias", ""))
                row[4].text = _aet_risk_recommendation(risk)
            set_table_font(tbl)
        else:
            add_text("Não foram selecionados fatores ergonômicos ou psicossociais específicos para este setor. Recomenda-se manter acompanhamento das condições de trabalho e atualizar a AET quando houver alteração de atividade, layout, mobiliário, jornada, organização do trabalho ou surgimento de queixas.")

        add_heading("Recomendações e plano de ação ergonômico do setor", 2)
        rec_manual = str(sector_data.get("recomendacoes", "") or "").strip()
        if rec_manual:
            add_text(rec_manual)
        plan = doc.add_table(rows=1, cols=5)
        
        try:
            plan.style = "Table Grid"
        except KeyError:
            pass
        for i, h in enumerate(["ACHADO ERGONÔMICO", "MEDIDA RECOMENDADA", "PRAZO", "RESPONSÁVEL", "PRIORIDADE"]):
            plan.rows[0].cells[i].text = h
        if risks:
            for risk in risks:
                row = plan.add_row().cells
                row[0].text = str(risk.get("risco", ""))
                row[1].text = _aet_risk_recommendation(risk)
                row[2].text = str(sector_data.get("prazo") or ("30 DIAS" if _normalize_option(risk.get("tipo_risco")) == "ERGONÔMICO PSICOSSOCIAL" else "Conforme plano de ação"))
                row[3].text = str(sector_data.get("responsavel") or "Empresa")
                row[4].text = _aet_priority_from_risks([risk], sector_data.get("prioridade", ""))
        else:
            row = plan.add_row().cells
            row[0].text = "Acompanhamento preventivo"
            row[1].text = "Manter boas práticas ergonômicas, organização do posto e acompanhamento periódico."
            row[2].text = str(sector_data.get("prazo") or "Contínuo")
            row[3].text = str(sector_data.get("responsavel") or "Empresa")
            row[4].text = str(sector_data.get("prioridade") or "Baixa")
        set_table_font(plan)

        add_heading("Conclusão do setor", 2)
        add_text(_aet_auto_sector_conclusion(sector_name, risks, sector_data))

    add_heading("13. CONSIDERAÇÕES FINAIS")
    add_text(
        f"A presente AET permitiu consolidar a análise das condições de trabalho da empresa {company_doc.get('empresa')}, considerando o perfil de {activity_profile.get('label')}, os setores avaliados, a relação função x atividade e os riscos selecionados no sistema."
    )
    add_text(
        "A análise reforça que a ergonomia deve ser tratada de forma integrada ao gerenciamento de riscos ocupacionais, abrangendo não apenas aspectos físicos e biomecânicos, mas também organização do trabalho, comunicação, autonomia, ritmo, pausas e fatores psicossociais."
    )

    add_heading("14. INTEGRAÇÃO DA AET AO PGR (NR-01)")
    add_text(
        "Os fatores de risco identificados nesta AET devem ser integrados ao inventário de riscos e ao plano de ação do PGR, com definição de medidas preventivas/corretivas, responsáveis, prazos, indicadores de efetividade e reavaliação periódica."
    )
    add_text(
        "Quando houver riscos ergonômicos psicossociais, recomenda-se atenção especial à organização do trabalho, acompanhamento de queixas, comunicação, suporte da liderança, prevenção de conflitos e registro das ações adotadas."
    )

    add_heading("15. CONCLUSÃO")
    condicao_geral = str(general.get("condicao_ergonomica_geral") or "Adequada com recomendações").strip()
    add_text(f"Condição ergonômica geral classificada: {condicao_geral}.")
    manual_conclusion = str(general.get("conclusao_geral_manual", "") or "").strip()
    if manual_conclusion:
        add_text(manual_conclusion)
    else:
        if any(_normalize_option(r.get("tipo_risco")) == "ERGONÔMICO PSICOSSOCIAL" for r in all_erg_risks):
            add_text("Conclui-se que as atividades analisadas apresentam fatores ergonômicos e psicossociais que devem ser acompanhados pela empresa, com implantação das medidas preventivas/corretivas indicadas, fortalecimento da comunicação, organização das demandas, orientação de liderança e trabalhadores e revisão periódica das condições de trabalho.")
        elif all_erg_risks:
            add_text("Conclui-se que as atividades analisadas apresentam fatores ergonômicos compatíveis com a rotina operacional, sendo recomendada a implantação das medidas de adequação postural, organização do posto, pausas, orientação ergonômica e acompanhamento periódico das condições de trabalho.")
        else:
            add_text("Conclui-se que, com base nas informações cadastradas, não foram evidenciados fatores ergonômicos críticos específicos nos setores selecionados, recomendando-se a manutenção das medidas preventivas, acompanhamento periódico e atualização da AET sempre que houver alteração na organização do trabalho, layout, mobiliário, processo ou queixas dos trabalhadores.")
    add_text("Esta AET deve ser revisada sempre que houver alteração relevante nas atividades, processos, mobiliário, layout, jornada, organização do trabalho, número de trabalhadores, ocorrência de queixas recorrentes ou inclusão de novos riscos ergonômicos/psicossociais no PGR.")

    add_heading("16. FECHAMENTO TÉCNICO E JURÍDICO")
    add_text(
        "Sob a ótica técnica, a AET constitui documento de apoio à gestão de segurança e saúde no trabalho, devendo ser utilizada para orientar medidas de prevenção, controle e melhoria contínua das condições laborais. A ausência de acompanhamento das medidas propostas pode comprometer a efetividade do gerenciamento de riscos."
    )
    add_text(
        "A empresa deve manter registros das ações implantadas, treinamentos, orientações, avaliações, revisões e evidências de acompanhamento, especialmente quando os fatores avaliados puderem impactar a saúde física ou mental dos trabalhadores."
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run_font(p.add_run("___________________________________________"), 10, True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run_font(p.add_run(responsavel), 10, True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run_font(p.add_run("Responsável Técnico"), 10, True)

    _replace_doc_placeholders(doc, _company_replacements(company_doc))
    _apply_aet_header_footer()

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    save_docx_with_default_font(doc, output_path)
    return output_path
