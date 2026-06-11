from __future__ import annotations

import json
import os
import re
import shutil
import tempfile
import uuid
import zipfile
from datetime import datetime, timedelta
from html.parser import HTMLParser
from pathlib import Path
from typing import Any

from flask import Flask, flash, jsonify, redirect, render_template, request, send_file, url_for
from werkzeug.utils import secure_filename
from werkzeug.exceptions import RequestEntityTooLarge
from flask_sqlalchemy import SQLAlchemy
from sqlalchemy import Text, inspect, text

try:
    from openpyxl import load_workbook
except Exception:  # pragma: no cover - usado apenas se a dependência não estiver instalada
    load_workbook = None

from word_generator import (
    NIVEL_RISCO_COLORS,
    POSSIBILIDADE_COLORS,
    SEVERIDADE_COLORS,
    TIPO_RISCO_COLORS,
    generate_action_plan_docx,
    generate_complete_pgr_docx,
    generate_complete_pcmso_docx,
    generate_complete_ltcat_docx,
    generate_aet_docx,
    generate_descritivo_setor_docx,
    generate_pcmso_docx,
    generate_riscos_pcmso_docx,
    generate_pgr_docx,
    generate_relacao_funcao_atividade_docx,
)

BASE_DIR = Path(__file__).resolve().parent
DATA_FILE = BASE_DIR / "data" / "riscos.json"
SETORES_FILE = BASE_DIR / "data" / "setores.json"
EXAMES_FILE = BASE_DIR / "data" / "exames.json"
RISK_IMPORT_TEMPLATE = BASE_DIR / "modelos" / "modelo_importacao_riscos.xlsx"
SECTOR_IMPORT_TEMPLATE = BASE_DIR / "modelos" / "modelo_importacao_setores.xlsx"
LINK_PGR_AET_TEMPLATE = BASE_DIR / "modelos" / "link_pgr_para_aet.docx"
LINK_AET_PSICOSSOCIAL_TEMPLATE = BASE_DIR / "modelos" / "link_aet_para_psicossocial.docx"
OUTPUT_DIR = BASE_DIR / "outputs"
INSTANCE_DIR = BASE_DIR / "instance"


def _database_uri() -> str:
    """Usa PostgreSQL no Render via DATABASE_URL e SQLite local como fallback."""
    url = os.environ.get("DATABASE_URL", "").strip()
    if url.startswith("postgres://"):
        url = url.replace("postgres://", "postgresql+psycopg://", 1)
    elif url.startswith("postgresql://"):
        url = url.replace("postgresql://", "postgresql+psycopg://", 1)
    if url:
        return url
    INSTANCE_DIR.mkdir(parents=True, exist_ok=True)
    return f"sqlite:///{INSTANCE_DIR / 'sst_riscos.db'}"


app = Flask(__name__)
app.secret_key = os.environ.get("SECRET_KEY", "troque-esta-chave-em-producao")
app.config["SQLALCHEMY_DATABASE_URI"] = _database_uri()
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
app.config["SQLALCHEMY_ENGINE_OPTIONS"] = {"pool_pre_ping": True}
# Evita uploads gigantes travarem o serviço no Render. Ajuste por variável de ambiente se precisar.
app.config["MAX_CONTENT_LENGTH"] = int(os.environ.get("MAX_UPLOAD_MB", "1024")) * 1024 * 1024
# A tela Gerar Laudos envia muitos campos quando há muitos setores, riscos, exames e dados de AET.
# Sem estes limites maiores, o Werkzeug pode retornar: Request Entity Too Large.
app.config["MAX_FORM_MEMORY_SIZE"] = int(os.environ.get("MAX_FORM_MEMORY_MB", "512")) * 1024 * 1024
app.config["MAX_FORM_PARTS"] = int(os.environ.get("MAX_FORM_PARTS", "500000"))

# Configuração da conversão visual do PDF psicossocial para Word.
# Usamos JPEG em DPI moderado para preservar o visual sem travar por arquivos enormes.
PSICOSSOCIAL_RENDER_DPI = max(90, min(180, int(os.environ.get("PSICOSSOCIAL_RENDER_DPI", "135"))))
PSICOSSOCIAL_JPEG_QUALITY = max(60, min(92, int(os.environ.get("PSICOSSOCIAL_JPEG_QUALITY", "82"))))
PSICOSSOCIAL_MAX_PAGES = int(os.environ.get("PSICOSSOCIAL_MAX_PAGES", "80"))

db = SQLAlchemy(app)

@app.errorhandler(RequestEntityTooLarge)
def handle_request_entity_too_large(error):
    flash(
        "O formulário ou arquivo enviado ficou maior que o limite configurado. "
        "Tente gerar novamente. Se estiver no Render, adicione ou aumente MAX_UPLOAD_MB=1024, MAX_FORM_MEMORY_MB=512 e MAX_FORM_PARTS=500000 nas variáveis de ambiente.",
        "error",
    )
    return redirect(request.referrer or url_for("generate"))


risk_group_items = db.Table(
    "risk_group_items",
    db.Column("risk_group_id", db.String(32), db.ForeignKey("risk_groups.id"), primary_key=True),
    db.Column("risk_id", db.String(32), db.ForeignKey("risks.id"), primary_key=True),
)


class Risk(db.Model):
    __tablename__ = "risks"

    id = db.Column(db.String(32), primary_key=True, default=lambda: uuid.uuid4().hex)
    risco = db.Column(db.String(255), nullable=False)
    acoes = db.Column(Text, nullable=False)
    indicador = db.Column(Text, nullable=False)
    tipo_risco = db.Column(db.String(80), nullable=False)
    descricao_agente = db.Column(Text, default="")
    possiveis_lesoes = db.Column(Text, nullable=False)
    fontes_circunstancias = db.Column(Text, nullable=False, default="Durante o processo de trabalho.")
    epis = db.Column(Text, nullable=False)
    epcs = db.Column(Text, nullable=False)
    ltcat_meio_propagacao = db.Column(Text, default="")
    ltcat_insalubridade = db.Column(Text, default="Não")
    ltcat_grau_insalubridade = db.Column(Text, default="Não aplicável")
    ltcat_aposentadoria_especial = db.Column(Text, default="Não")
    ltcat_enquadramento_tecnico = db.Column(Text, default="")
    ltcat_parecer_previdenciario = db.Column(Text, default="")
    ltcat_periodicidade_jornada = db.Column(Text, default="Mensal (<= 4 horas < 10% jornada)")
    grau_severidade = db.Column(db.String(80), nullable=False)
    grau_possibilidade = db.Column(db.String(80), nullable=False)
    grau_nivel_risco = db.Column(db.String(80), nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow, nullable=False)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow, nullable=False)

    def to_dict(self) -> dict[str, Any]:
        return {
            "id": self.id,
            "risco": self.risco,
            "acoes": self.acoes,
            "prazo_implantacao": "",
            "prazo_reavaliacao": "",
            "indicador": self.indicador,
            "tipo_risco": self.tipo_risco,
            "descricao_agente": self.descricao_agente or "",
            "possiveis_lesoes": self.possiveis_lesoes,
            "fontes_circunstancias": self.fontes_circunstancias or "Durante o processo de trabalho.",
            "epis": self.epis,
            "epcs": self.epcs,
            "ltcat_meio_propagacao": self.ltcat_meio_propagacao or "",
            "ltcat_insalubridade": self.ltcat_insalubridade or "Não",
            "ltcat_grau_insalubridade": self.ltcat_grau_insalubridade or "Não aplicável",
            "ltcat_aposentadoria_especial": self.ltcat_aposentadoria_especial or "Não",
            "ltcat_enquadramento_tecnico": self.ltcat_enquadramento_tecnico or "",
            "ltcat_parecer_previdenciario": self.ltcat_parecer_previdenciario or "",
            "ltcat_periodicidade_jornada": self.ltcat_periodicidade_jornada or "Mensal (<= 4 horas < 10% jornada)",
            "grau_severidade": self.grau_severidade,
            "grau_possibilidade": self.grau_possibilidade,
            "grau_nivel_risco": self.grau_nivel_risco,
            "created_at": self.created_at.strftime("%Y-%m-%d %H:%M:%S") if self.created_at else "",
            "updated_at": self.updated_at.strftime("%Y-%m-%d %H:%M:%S") if self.updated_at else "",
        }


class RiskGroup(db.Model):
    __tablename__ = "risk_groups"

    id = db.Column(db.String(32), primary_key=True, default=lambda: uuid.uuid4().hex)
    nome = db.Column(db.String(255), nullable=False, unique=True)
    descricao = db.Column(Text, default="")
    risks = db.relationship("Risk", secondary=risk_group_items, lazy="select")
    created_at = db.Column(db.DateTime, default=datetime.utcnow, nullable=False)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow, nullable=False)

    def to_dict(self) -> dict[str, Any]:
        ordered_risks = sorted(self.risks or [], key=lambda item: (item.risco or "").lower())
        return {
            "id": self.id,
            "nome": self.nome,
            "descricao": self.descricao or "",
            "risk_ids": [risk.id for risk in ordered_risks],
            "risk_names": [risk.risco for risk in ordered_risks],
            "risk_count": len(ordered_risks),
            "created_at": self.created_at.strftime("%Y-%m-%d %H:%M:%S") if self.created_at else "",
            "updated_at": self.updated_at.strftime("%Y-%m-%d %H:%M:%S") if self.updated_at else "",
        }



class SectorGroup(db.Model):
    __tablename__ = "sector_groups"

    id = db.Column(db.String(32), primary_key=True, default=lambda: uuid.uuid4().hex)
    nome = db.Column(db.String(255), nullable=False, unique=True)
    is_temporary = db.Column(db.Boolean, default=False, nullable=False)
    temporary_expires_at = db.Column(db.DateTime, nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow, nullable=False)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow, nullable=False)

    def to_dict(self) -> dict[str, Any]:
        return {
            "id": self.id,
            "nome": self.nome,
            "is_temporary": bool(self.is_temporary),
            "temporary_expires_at": self.temporary_expires_at.strftime("%Y-%m-%d %H:%M:%S") if self.temporary_expires_at else "",
            "created_at": self.created_at.strftime("%Y-%m-%d %H:%M:%S") if self.created_at else "",
            "updated_at": self.updated_at.strftime("%Y-%m-%d %H:%M:%S") if self.updated_at else "",
        }


class Sector(db.Model):
    __tablename__ = "sectors"

    id = db.Column(db.String(32), primary_key=True, default=lambda: uuid.uuid4().hex)
    setor = db.Column(db.String(255), nullable=False)
    group_id = db.Column(db.String(32), nullable=True)
    cargos = db.Column(db.JSON, nullable=False, default=list)
    created_at = db.Column(db.DateTime, default=datetime.utcnow, nullable=False)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow, nullable=False)

    def to_dict(self) -> dict[str, Any]:
        cargos = self.cargos or []
        group = db.session.get(SectorGroup, self.group_id) if self.group_id else None
        return {
            "id": self.id,
            "setor": self.setor,
            "grupo_id": self.group_id or "",
            "grupo_nome": group.nome if group else "Sem grupo",
            "cargos": cargos if isinstance(cargos, list) else [],
            "created_at": self.created_at.strftime("%Y-%m-%d %H:%M:%S") if self.created_at else "",
            "updated_at": self.updated_at.strftime("%Y-%m-%d %H:%M:%S") if self.updated_at else "",
        }


class Exam(db.Model):
    __tablename__ = "exams"

    id = db.Column(db.String(32), primary_key=True, default=lambda: uuid.uuid4().hex)
    exame = db.Column(db.String(255), nullable=False)
    periodicidade = db.Column(db.String(120), default="")
    admissional = db.Column(db.String(120), default="")
    periodico = db.Column(db.String(120), default="")
    retorno = db.Column(db.String(120), default="")
    mudanca = db.Column(db.String(120), default="")
    demissional = db.Column(db.String(120), default="")
    created_at = db.Column(db.DateTime, default=datetime.utcnow, nullable=False)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow, nullable=False)

    def to_dict(self) -> dict[str, Any]:
        return {
            "id": self.id,
            "exame": self.exame,
            "periodicidade": self.periodicidade or "",
            "admissional": self.admissional or "",
            "periodico": self.periodico or "",
            "retorno": self.retorno or "",
            "mudanca": self.mudanca or "",
            "demissional": self.demissional or "",
            "created_at": self.created_at.strftime("%Y-%m-%d %H:%M:%S") if self.created_at else "",
            "updated_at": self.updated_at.strftime("%Y-%m-%d %H:%M:%S") if self.updated_at else "",
        }


class Company(db.Model):
    __tablename__ = "companies"

    id = db.Column(db.String(32), primary_key=True, default=lambda: uuid.uuid4().hex)
    nome = db.Column(Text, nullable=False)
    cnpj = db.Column(Text, nullable=False)
    endereco = db.Column(Text, default="")
    bairro_cidade = db.Column(Text, default="")
    cep = db.Column(Text, default="")
    cnae1 = db.Column(Text, default="")
    descricao1 = db.Column(Text, default="")
    grau1 = db.Column(Text, default="")
    cnae2 = db.Column(Text, default="")
    descricao2 = db.Column(Text, default="")
    grau2 = db.Column(Text, default="")
    funcionarios = db.Column(Text, default="")
    data_atual = db.Column(Text, default="")
    data_final = db.Column(Text, default="")
    email = db.Column(Text, default="")
    fone = db.Column(Text, default="")
    data_avaliacao = db.Column(Text, default="")
    created_at = db.Column(db.DateTime, default=datetime.utcnow, nullable=False)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow, nullable=False)

    def to_dict(self) -> dict[str, Any]:
        return {
            "id": self.id,
            "nome": self.nome,
            "empresa": self.nome,
            "cnpj": self.cnpj,
            "endereco": self.endereco or "",
            "bairro_cidade": self.bairro_cidade or "",
            "cep": self.cep or "",
            "cnae1": self.cnae1 or "",
            "descricao1": self.descricao1 or "",
            "grau1": self.grau1 or "",
            "cnae2": self.cnae2 or "",
            "descricao2": self.descricao2 or "",
            "grau2": self.grau2 or "",
            "funcionarios": self.funcionarios or "",
            "data_atual": self.data_atual or "",
            "data_final": self.data_final or "",
            "email": self.email or "",
            "fone": self.fone or "",
            "data_avaliacao": self.data_avaliacao or "",
            # aliases usados pelos geradores antigos
            "cnae": self.cnae1 or "",
            "descricao_atividade": self.descricao1 or "",
            "grau_risco": self.grau1 or "",
            "cnae_secundario": self.cnae2 or "",
            "descricao_atividade_secundaria": self.descricao2 or "",
            "grau_risco_secundario": self.grau2 or "",
            "created_at": self.created_at.strftime("%Y-%m-%d %H:%M:%S") if self.created_at else "",
            "updated_at": self.updated_at.strftime("%Y-%m-%d %H:%M:%S") if self.updated_at else "",
        }


class ReportProfile(db.Model):
    """Configuração salva da tela Gerar laudos para uma empresa.

    Guarda setores, riscos, grupos de riscos, exames e datas usadas na finalização,
    permitindo regenerar os laudos depois sem refazer toda a seleção.
    """

    __tablename__ = "report_profiles"

    id = db.Column(db.String(32), primary_key=True, default=lambda: uuid.uuid4().hex)
    company_id = db.Column(db.String(32), nullable=False, index=True)
    nome = db.Column(db.String(255), nullable=False)
    data_criacao_laudo = db.Column(Text, default="")
    ajuste_psicossocial = db.Column(db.String(1), default="")
    data_da_revisao = db.Column(Text, default="")
    state = db.Column(db.JSON, nullable=False, default=dict)
    created_at = db.Column(db.DateTime, default=datetime.utcnow, nullable=False)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow, nullable=False)

    def to_dict(self) -> dict[str, Any]:
        company = db.session.get(Company, self.company_id) if self.company_id else None
        return {
            "id": self.id,
            "company_id": self.company_id,
            "company_nome": company.nome if company else "",
            "nome": self.nome,
            "data_criacao_laudo": self.data_criacao_laudo or "",
            "ajuste_psicossocial": self.ajuste_psicossocial or "",
            "data_da_revisao": self.data_da_revisao or "",
            "state": self.state or {},
            "created_at": self.created_at.strftime("%Y-%m-%d %H:%M:%S") if self.created_at else "",
            "updated_at": self.updated_at.strftime("%Y-%m-%d %H:%M:%S") if self.updated_at else "",
        }


class ImportedLaudoTemplate(db.Model):
    """Modelo reutilizável extraído de um laudo antigo.

    Diferente da configuração salva de uma empresa, este modelo não fica preso
    à empresa original. Ele guarda setores/cargos/riscos/exames extraídos para
    serem aplicados depois em qualquer empresa cadastrada.
    """

    __tablename__ = "imported_laudo_templates"

    id = db.Column(db.String(32), primary_key=True, default=lambda: uuid.uuid4().hex)
    nome = db.Column(db.String(255), nullable=False)
    source_company = db.Column(Text, default="")
    source_cnpj = db.Column(Text, default="")
    state = db.Column(db.JSON, nullable=False, default=dict)
    created_at = db.Column(db.DateTime, default=datetime.utcnow, nullable=False)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow, nullable=False)

    def to_dict(self) -> dict[str, Any]:
        state = self.state or {}
        return {
            "id": self.id,
            "nome": self.nome,
            "source_company": self.source_company or "",
            "source_cnpj": self.source_cnpj or "",
            "sector_count": len(state.get("setores_cargos") or state.get("setores") or []),
            "risk_count": len(state.get("riscos_detalhados") or state.get("riscos") or []),
            "linked_risk_count": sum(len(v) for v in (state.get("sector_risks") or {}).values() if isinstance(v, list)),
            "exam_count": len(state.get("exames") or []),
            "created_at": self.created_at.strftime("%Y-%m-%d %H:%M:%S") if self.created_at else "",
            "updated_at": self.updated_at.strftime("%Y-%m-%d %H:%M:%S") if self.updated_at else "",
        }


class ImportedLaudoDraft(db.Model):
    """Rascunho de importação salvo no banco.

    Evita enviar JSON gigante em campos ocultos no formulário de revisão.
    O upload lê o laudo, grava tudo aqui e a tela de revisão envia somente o ID.
    """

    __tablename__ = "imported_laudo_drafts"

    id = db.Column(db.String(32), primary_key=True, default=lambda: uuid.uuid4().hex)
    state = db.Column(db.JSON, nullable=False, default=dict)
    expires_at = db.Column(db.DateTime, nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow, nullable=False)

    def to_dict(self) -> dict[str, Any]:
        return {
            "id": self.id,
            "state": self.state or {},
            "expires_at": self.expires_at.strftime("%Y-%m-%d %H:%M:%S") if self.expires_at else "",
            "created_at": self.created_at.strftime("%Y-%m-%d %H:%M:%S") if self.created_at else "",
        }


FORM_OPTIONS = {
    "tipos_risco": list(TIPO_RISCO_COLORS.keys()),
    "severidades": list(SEVERIDADE_COLORS.keys()),
    "possibilidades": list(POSSIBILIDADE_COLORS.keys()),
    "niveis_risco": list(NIVEL_RISCO_COLORS.keys()),
}

TIPOS_RISCO_LTCAT = {"FÍSICO", "QUÍMICO", "BIOLÓGICO"}

# Presets técnicos de AET por CNAE/atividade.
# Usados para pré-marcar o formulário da AET conforme CNAE/atividade da empresa.
# Tudo continua editável antes da geração.
AET_CNAE_PRESETS = [
    {
        "keywords": ["47.81", "4781", "vestuário", "vestuario", "comércio varejista", "comercio varejista", "loja", "vendas", "artigos do vestuário", "acessórios"],
        "label": "Comércio varejista / vestuário",
        "categoria": "Comércio e atendimento",
        "cnaes_referencia": ["47.81-4-00", "47.82-2-01", "47.82-2-02"],
        "setores_sugeridos": ["VENDAS", "ATENDIMENTO", "CAIXA", "ESTOQUE/DEPÓSITO", "SUPERVISÃO"],
        "riscos_sugeridos": ["Postura em pé por longos períodos", "Frequente deslocamento a pé", "Trabalho em posturas incômodas", "Baixa autonomia", "Atendimento ao público"],
        "exames_sugeridos": ["Exame clínico", "Acuidade visual", "Anamnese ocupacional"],
        "general": {"tipo_documento": "AET completa com formulário ergonômico", "motivo_analise": "Atendimento à NR-17 e integração com o PGR", "condicao_ergonomica_geral": "Adequada com recomendações"},
        "sector": {"postura": ["Em pé", "Deslocamento frequente", "Alternado"], "tipo_atividade": "Atendimento, organização de mercadorias e apoio operacional", "exigencia_fisica": "Moderada", "exigencia_cognitiva": "Moderada", "ritmo_trabalho": "Por demanda e fluxo de clientes", "pausas": "Intervalos legais e pausas breves conforme organização da loja", "mobiliario": "Parcialmente adequado", "ambiente": "A avaliar conforto térmico, iluminação e circulação", "organizacao": "Atendimento ao público, organização de prioridades e comunicação com liderança", "equipamentos": "Balcão, prateleiras, araras, computador/sistema e materiais de apoio", "fatores": ["Atendimento ao público", "Deslocamentos no setor", "Postura em pé", "Organização de demandas"], "medidas": ["Alternância postural", "Pausas breves", "Organização do posto", "Orientação NR-17"], "prioridade": "Média", "prazo": "30 dias", "responsavel": "Empresa / Administração"},
    },
    {
        "keywords": ["47.11", "4711", "mercado", "supermercado", "minimercado", "mercearia", "atacado", "varejista de mercadorias"],
        "label": "Mercado / supermercado / atacarejo",
        "categoria": "Comércio e atendimento",
        "cnaes_referencia": ["47.11-3-01", "47.11-3-02", "47.12-1-00"],
        "setores_sugeridos": ["CAIXA", "REPOSIÇÃO", "AÇOUGUE", "PADARIA", "HORTIFRUTI", "DEPÓSITO", "ADMINISTRATIVO"],
        "riscos_sugeridos": ["Postura em pé por longos períodos", "Levantamento e transporte manual de cargas", "Movimentos repetitivos", "Ritmo intenso de trabalho", "Atendimento ao público"],
        "exames_sugeridos": ["Exame clínico", "Acuidade visual", "Audiometria quando houver ruído", "Espirometria quando houver poeiras/fumos"],
        "general": {"tipo_documento": "AET completa com formulário ergonômico", "motivo_analise": "Atendimento à NR-17 e avaliação das exigências físicas e organizacionais", "condicao_ergonomica_geral": "Adequada com recomendações"},
        "sector": {"postura": ["Em pé", "Deslocamento frequente", "Alternado"], "tipo_atividade": "Atendimento, reposição, organização de mercadorias e operação de caixa", "exigencia_fisica": "Moderada", "exigencia_cognitiva": "Moderada", "ritmo_trabalho": "Por demanda de clientes, recebimento e reposição de produtos", "pausas": "Intervalos legais e pausas breves conforme escala operacional", "mobiliario": "A avaliar bancadas, caixas, assentos e alturas de apoio", "ambiente": "A avaliar circulação, ruído, temperatura, iluminação e organização de cargas", "organizacao": "Fluxo operacional variável, atendimento ao público e necessidade de definição de prioridades", "equipamentos": "Checkouts, carrinhos, gôndolas, pallets, computadores e equipamentos de apoio", "fatores": ["Atendimento ao público", "Postura em pé", "Manuseio de materiais", "Deslocamentos no setor"], "medidas": ["Alternância postural", "Pausas breves", "Rodízio de tarefas", "Ferramentas adequadas", "Orientação NR-17"], "prioridade": "Média", "prazo": "30 dias", "responsavel": "Empresa / Administração"},
    },
    {
        "keywords": ["82.19", "8219", "administrativo", "escritório", "escritorio", "preparação de documentos", "apoio administrativo", "serviços administrativos", "69.", "70."],
        "label": "Administrativo / escritório",
        "categoria": "Administrativo",
        "cnaes_referencia": ["82.19-9-99", "82.11-3-00", "69.20-6-01", "70.20-4-00"],
        "setores_sugeridos": ["ADMINISTRATIVO", "RECEPÇÃO", "FINANCEIRO", "RH", "SUPORTE"],
        "riscos_sugeridos": ["Postura sentado por longos períodos", "Presença de reflexos em telas", "Movimentos repetitivos", "Baixa autonomia", "Sobrecarga de trabalho"],
        "exames_sugeridos": ["Exame clínico", "Acuidade visual", "Anamnese ocupacional", "Anamnese psicossocial quando aplicável"],
        "general": {"tipo_documento": "AET documental com análise por setor", "motivo_analise": "Atendimento à NR-17 e análise do posto administrativo", "condicao_ergonomica_geral": "Adequada com recomendações"},
        "sector": {"postura": ["Sentado", "Alternado"], "tipo_atividade": "Atividades administrativas, digitação, atendimento e organização documental", "exigencia_fisica": "Baixa", "exigencia_cognitiva": "Moderada", "ritmo_trabalho": "Por demanda administrativa", "pausas": "Pausas breves e alternância postural recomendadas", "mobiliario": "A avaliar cadeira, mesa, monitor e acessórios", "ambiente": "A avaliar iluminação, reflexos, ventilação e conforto acústico", "organizacao": "Rotina administrativa com atenção a prazos, prioridades e comunicação", "equipamentos": "Computador, monitor, teclado, mouse, telefone e documentos", "fatores": ["Trabalho sentado", "Digitação", "Atenção contínua", "Organização de prioridades"], "medidas": ["Ajuste de mobiliário", "Alternância postural", "Pausas visuais", "Orientação NR-17"], "prioridade": "Baixa", "prazo": "60 dias", "responsavel": "Empresa / Administração"},
    },
    {
        "keywords": ["81.21", "81.22", "81.29", "limpeza", "serviços gerais", "servicos gerais", "conservação", "conservacao", "serviços combinados"],
        "label": "Limpeza / serviços gerais",
        "categoria": "Serviços operacionais",
        "cnaes_referencia": ["81.21-4-00", "81.22-2-00", "81.29-0-00"],
        "setores_sugeridos": ["SERVIÇOS GERAIS", "LIMPEZA", "APOIO", "ÁREAS COMUNS"],
        "riscos_sugeridos": ["Posturas incômodas", "Esforço físico", "Frequente deslocamento a pé", "Levantamento e transporte manual de cargas", "Produtos saneantes"],
        "exames_sugeridos": ["Exame clínico", "Acuidade visual", "Hemograma quando houver exposição química", "Espirometria quando houver aerodispersoides"],
        "general": {"tipo_documento": "AET completa com formulário ergonômico", "motivo_analise": "Atendimento à NR-17 e avaliação das exigências físicas das atividades", "condicao_ergonomica_geral": "Adequada com recomendações"},
        "sector": {"postura": ["Em pé", "Deslocamento frequente", "Inclinação/flexão de tronco", "Agachamento eventual"], "tipo_atividade": "Limpeza, conservação, deslocamento e manuseio de materiais", "exigencia_fisica": "Moderada", "exigencia_cognitiva": "Baixa", "ritmo_trabalho": "Por rotina e demandas do ambiente", "pausas": "Pausas e alternância conforme intensidade da atividade", "mobiliario": "Não aplicável diretamente; avaliar equipamentos e ferramentas manuais", "ambiente": "A avaliar piso, circulação, ventilação e disponibilidade de local para descanso", "organizacao": "Rotina operacional com definição clara de prioridades", "equipamentos": "Vassouras, rodos, baldes, panos, produtos saneantes e carrinhos quando disponíveis", "fatores": ["Esforço físico", "Posturas incômodas", "Deslocamentos frequentes", "Manuseio de materiais"], "medidas": ["Alternância de tarefas", "Pausas breves", "Ferramentas adequadas", "Orientação NR-17"], "prioridade": "Média", "prazo": "30 dias", "responsavel": "Empresa / Administração"},
    },
    {
        "keywords": ["81.12", "8112", "condomínio", "condominio", "portaria", "porteiro", "zelador", "vigilância", "vigilancia"],
        "label": "Condomínio / portaria",
        "categoria": "Serviços prediais",
        "cnaes_referencia": ["81.12-5-00", "80.11-1-01"],
        "setores_sugeridos": ["PORTARIA", "ADMINISTRAÇÃO", "LIMPEZA", "MANUTENÇÃO"],
        "riscos_sugeridos": ["Postura sentado e/ou em pé", "Atenção contínua", "Atendimento ao público", "Comunicação hostil", "Deslocamento frequente"],
        "exames_sugeridos": ["Exame clínico", "Acuidade visual", "Anamnese psicossocial quando aplicável"],
        "general": {"tipo_documento": "AET completa com formulário ergonômico", "motivo_analise": "Atendimento à NR-17 e análise de atividade com atenção contínua", "condicao_ergonomica_geral": "Adequada com recomendações"},
        "sector": {"postura": ["Sentado", "Em pé", "Alternado"], "tipo_atividade": "Controle de acesso, atendimento, monitoramento e comunicação", "exigencia_fisica": "Baixa", "exigencia_cognitiva": "Elevada", "ritmo_trabalho": "Contínuo, com atenção permanente e demandas variáveis", "pausas": "Pausas conforme escala, mantendo cobertura operacional", "mobiliario": "A avaliar cadeira, bancada, campo visual e acesso aos controles", "ambiente": "A avaliar ventilação, iluminação, conforto térmico e ruído", "organizacao": "Atenção constante, comunicação e controle de prioridades", "equipamentos": "Portão, interfone, rádio/telefone, computador, câmeras e controles", "fatores": ["Atenção contínua", "Atendimento ao público", "Comunicação", "Postura alternada"], "medidas": ["Ajuste do posto", "Pausas breves", "Definição de procedimentos", "Orientação NR-17"], "prioridade": "Média", "prazo": "30 dias", "responsavel": "Empresa / Administração"},
    },
    {
        "keywords": ["56.11", "5611", "restaurante", "lanchonete", "bar", "alimentação", "alimentacao", "cozinha", "refeição", "refeicao"],
        "label": "Restaurante / lanchonete / cozinha",
        "categoria": "Alimentação",
        "cnaes_referencia": ["56.11-2-01", "56.11-2-03", "56.20-1-04"],
        "setores_sugeridos": ["COZINHA", "ATENDIMENTO", "CAIXA", "ESTOQUE", "SALÃO"],
        "riscos_sugeridos": ["Postura em pé por longos períodos", "Trabalho em posturas incômodas", "Movimentos repetitivos", "Calor", "Levantamento e transporte de cargas"],
        "exames_sugeridos": ["Exame clínico", "Acuidade visual", "Hemograma quando aplicável"],
        "general": {"tipo_documento": "AET completa com formulário ergonômico", "motivo_analise": "Atendimento à NR-17 e avaliação de atividades de cozinha e atendimento", "condicao_ergonomica_geral": "Adequada com recomendações"},
        "sector": {"postura": ["Em pé", "Deslocamento frequente", "Inclinação/flexão de tronco"], "tipo_atividade": "Preparo de alimentos, atendimento, organização e limpeza de utensílios/ambientes", "exigencia_fisica": "Moderada", "exigencia_cognitiva": "Moderada", "ritmo_trabalho": "Por demanda de preparo e atendimento", "pausas": "Pausas conforme escala e alternância de atividades", "mobiliario": "A avaliar bancadas, alturas de apoio e circulação", "ambiente": "A avaliar calor, ventilação, piso, iluminação e organização do fluxo", "organizacao": "Demandas variáveis, atenção ao preparo e atendimento ao público", "equipamentos": "Fogão, balcões, utensílios, refrigeradores, caixas e equipamentos de apoio", "fatores": ["Postura em pé", "Deslocamentos no setor", "Manuseio de materiais", "Atenção contínua"], "medidas": ["Alternância postural", "Pausas breves", "Organização do posto", "Ferramentas adequadas", "Orientação NR-17"], "prioridade": "Média", "prazo": "30 dias", "responsavel": "Empresa / Administração"},
    },
    {
        "keywords": ["86.30", "8630", "86.40", "8640", "clínica", "clinica", "consultório", "consultorio", "atividade médica", "exames complementares", "laboratório", "laboratorio"],
        "label": "Clínica / consultório / laboratório",
        "categoria": "Saúde",
        "cnaes_referencia": ["86.30-5-02", "86.40-2-02", "86.40-2-08"],
        "setores_sugeridos": ["RECEPÇÃO", "ATENDIMENTO", "COLETA/EXAMES", "ADMINISTRATIVO", "LIMPEZA"],
        "riscos_sugeridos": ["Postura sentada", "Atenção contínua", "Atendimento ao público", "Movimentos repetitivos", "Comunicação com pacientes"],
        "exames_sugeridos": ["Exame clínico", "Acuidade visual", "Hemograma quando houver exposição biológica", "Cartão de vacina quando aplicável"],
        "general": {"tipo_documento": "AET completa com formulário ergonômico", "motivo_analise": "Atendimento à NR-17 e análise de atividades administrativas/assistenciais", "condicao_ergonomica_geral": "Adequada com recomendações"},
        "sector": {"postura": ["Sentado", "Em pé", "Alternado"], "tipo_atividade": "Atendimento, registro, apoio assistencial e organização documental", "exigencia_fisica": "Baixa", "exigencia_cognitiva": "Moderada", "ritmo_trabalho": "Por agendamento, demanda de pacientes e rotinas internas", "pausas": "Pausas breves e alternância postural recomendadas", "mobiliario": "A avaliar cadeiras, bancadas, macas, computadores e campo visual", "ambiente": "A avaliar iluminação, circulação, ruído, ventilação e privacidade", "organizacao": "Atendimento ao público, atenção contínua e comunicação com equipe/pacientes", "equipamentos": "Computadores, impressoras, equipamentos de exames, telefone e materiais de apoio", "fatores": ["Atendimento ao público", "Atenção contínua", "Trabalho sentado", "Comunicação"], "medidas": ["Ajuste de mobiliário", "Alternância postural", "Pausas visuais", "Melhoria de comunicação", "Orientação NR-17"], "prioridade": "Baixa", "prazo": "60 dias", "responsavel": "Empresa / Administração"},
    },
    {
        "keywords": ["45.20", "4520", "oficina", "mecânica", "mecanica", "manutenção de veículos", "manutencao de veiculos", "autopeças", "autopecas"],
        "label": "Oficina mecânica / manutenção automotiva",
        "categoria": "Manutenção",
        "cnaes_referencia": ["45.20-0-01", "45.20-0-02", "45.30-7-03"],
        "setores_sugeridos": ["OFICINA", "MANUTENÇÃO", "ADMINISTRATIVO", "ESTOQUE"],
        "riscos_sugeridos": ["Posturas incômodas", "Esforço físico", "Levantamento de cargas", "Movimentos repetitivos", "Uso de ferramentas manuais"],
        "exames_sugeridos": ["Exame clínico", "Acuidade visual", "Audiometria quando houver ruído", "Hemograma/TGO/TGP quando houver químicos"],
        "general": {"tipo_documento": "AET completa com formulário ergonômico", "motivo_analise": "Atendimento à NR-17 e avaliação de atividades de manutenção", "condicao_ergonomica_geral": "Adequada com recomendações"},
        "sector": {"postura": ["Em pé", "Inclinação/flexão de tronco", "Agachamento eventual", "Deslocamento frequente"], "tipo_atividade": "Manutenção, inspeção, uso de ferramentas e organização de peças/equipamentos", "exigencia_fisica": "Moderada", "exigencia_cognitiva": "Moderada", "ritmo_trabalho": "Por ordens de serviço e demandas de manutenção", "pausas": "Pausas conforme intensidade e alternância de tarefas", "mobiliario": "Não aplicável diretamente; avaliar bancadas, alturas e ferramentas", "ambiente": "A avaliar circulação, ruído, iluminação, ventilação e organização de ferramentas", "organizacao": "Demandas variáveis, necessidade de procedimentos e comunicação de segurança", "equipamentos": "Ferramentas manuais, bancadas, elevadores, compressores e equipamentos de apoio", "fatores": ["Posturas incômodas", "Manuseio de materiais", "Atenção contínua", "Deslocamentos no setor"], "medidas": ["Ferramentas adequadas", "Alternância de tarefas", "Pausas breves", "Organização do posto", "Orientação NR-17"], "prioridade": "Média", "prazo": "30 dias", "responsavel": "Empresa / Administração"},
    },
    {
        "keywords": ["43.", "41.", "42.", "construção", "construcao", "obra", "engenharia", "serviços de construção", "edificação", "edificacao"],
        "label": "Construção civil / obras",
        "categoria": "Construção",
        "cnaes_referencia": ["41.20-4-00", "43.99-1-99", "42.99-5-99"],
        "setores_sugeridos": ["OBRA", "OPERACIONAL", "ALMOXARIFADO", "ADMINISTRATIVO", "ENGENHARIA"],
        "riscos_sugeridos": ["Esforço físico", "Levantamento e transporte manual de cargas", "Posturas incômodas", "Trabalho em altura", "Ritmo de trabalho"],
        "exames_sugeridos": ["Exame clínico", "Acuidade visual", "Audiometria", "Espirometria quando houver poeiras", "ECG quando aplicável"],
        "general": {"tipo_documento": "AET completa com formulário ergonômico", "motivo_analise": "Atendimento à NR-17 e análise das exigências físicas de atividades operacionais", "condicao_ergonomica_geral": "Parcialmente adequada"},
        "sector": {"postura": ["Em pé", "Deslocamento frequente", "Inclinação/flexão de tronco", "Agachamento eventual"], "tipo_atividade": "Execução, apoio e acompanhamento de atividades operacionais de obra", "exigencia_fisica": "Elevada", "exigencia_cognitiva": "Moderada", "ritmo_trabalho": "Por etapas da obra, cronograma e demandas operacionais", "pausas": "Pausas conforme intensidade física, calor e organização da frente de serviço", "mobiliario": "Não aplicável diretamente; avaliar ferramentas, bancadas e áreas de apoio", "ambiente": "A avaliar calor, ruído, poeira, circulação, iluminação e áreas de descanso", "organizacao": "Atividades com prazos, frentes de serviço e necessidade de comunicação contínua", "equipamentos": "Ferramentas manuais, escadas, andaimes, máquinas e materiais de construção", "fatores": ["Esforço físico", "Posturas incômodas", "Manuseio de materiais", "Deslocamentos no setor"], "medidas": ["Rodízio de tarefas", "Pausas breves", "Ferramentas adequadas", "Organização do posto", "Orientação NR-17"], "prioridade": "Alta", "prazo": "30 dias", "responsavel": "Empresa / Engenharia / Segurança do Trabalho"},
    },
    {
        "keywords": ["49.", "52.", "transporte", "logística", "logistica", "armazenagem", "depósito", "deposito", "entrega", "motorista", "carga"],
        "label": "Transporte / logística / depósito",
        "categoria": "Logística",
        "cnaes_referencia": ["49.30-2-02", "52.11-7-99", "52.50-8-04"],
        "setores_sugeridos": ["OPERACIONAL", "DEPÓSITO", "EXPEDIÇÃO", "GARAGEM", "ADMINISTRATIVO"],
        "riscos_sugeridos": ["Levantamento e transporte manual de cargas", "Postura sentada por longos períodos", "Vibração de corpo inteiro", "Deslocamento frequente", "Ritmo de entregas"],
        "exames_sugeridos": ["Exame clínico", "Acuidade visual", "Audiometria quando houver ruído", "Avaliação osteomuscular quando aplicável"],
        "general": {"tipo_documento": "AET completa com formulário ergonômico", "motivo_analise": "Atendimento à NR-17 e avaliação das exigências de transporte, armazenamento e movimentação", "condicao_ergonomica_geral": "Adequada com recomendações"},
        "sector": {"postura": ["Sentado", "Em pé", "Deslocamento frequente", "Alternado"], "tipo_atividade": "Transporte, movimentação, conferência, organização e expedição de materiais", "exigencia_fisica": "Moderada", "exigencia_cognitiva": "Moderada", "ritmo_trabalho": "Por rotas, demanda de carga/descarga e prazos operacionais", "pausas": "Pausas conforme escala, trajetos e intensidade da movimentação", "mobiliario": "A avaliar assentos, cabines, bancadas e áreas de apoio", "ambiente": "A avaliar circulação, piso, iluminação, ruído, temperatura e organização de cargas", "organizacao": "Demandas variáveis, prazos, rotas e comunicação operacional", "equipamentos": "Veículos, carrinhos, paleteiras, prateleiras, caixas e sistemas de controle", "fatores": ["Manuseio de materiais", "Deslocamentos no setor", "Trabalho sentado", "Atenção contínua"], "medidas": ["Alternância postural", "Pausas breves", "Ferramentas adequadas", "Organização do posto", "Orientação NR-17"], "prioridade": "Média", "prazo": "30 dias", "responsavel": "Empresa / Administração"},
    },
    {
        "keywords": ["85.", "educação", "educacao", "escola", "curso", "ensino", "professor", "creche"],
        "label": "Escola / educação",
        "categoria": "Educação",
        "cnaes_referencia": ["85.13-9-00", "85.20-1-00", "85.99-6-99"],
        "setores_sugeridos": ["SALA DE AULA", "ADMINISTRATIVO", "COORDENAÇÃO", "LIMPEZA", "APOIO"],
        "riscos_sugeridos": ["Postura em pé", "Uso contínuo da voz", "Atenção contínua", "Sobrecarga de trabalho", "Atendimento ao público"],
        "exames_sugeridos": ["Exame clínico", "Acuidade visual", "Avaliação vocal quando aplicável", "Anamnese psicossocial quando aplicável"],
        "general": {"tipo_documento": "AET documental com análise por setor", "motivo_analise": "Atendimento à NR-17 e avaliação de atividades pedagógicas/administrativas", "condicao_ergonomica_geral": "Adequada com recomendações"},
        "sector": {"postura": ["Em pé", "Sentado", "Alternado"], "tipo_atividade": "Atividades pedagógicas, atendimento, orientação, registros e acompanhamento de alunos", "exigencia_fisica": "Baixa", "exigencia_cognitiva": "Elevada", "ritmo_trabalho": "Por calendário, horários de aula e demandas de atendimento", "pausas": "Pausas conforme jornada, intervalos e organização pedagógica", "mobiliario": "A avaliar cadeiras, mesas, quadro, computador e campo visual", "ambiente": "A avaliar iluminação, ruído, ventilação, conforto térmico e circulação", "organizacao": "Atenção contínua, comunicação, demandas pedagógicas e relacionamento interpessoal", "equipamentos": "Quadro, computador, mesas, cadeiras, materiais didáticos e telefone", "fatores": ["Atenção contínua", "Comunicação", "Postura alternada", "Atendimento ao público"], "medidas": ["Pausas breves", "Organização do posto", "Melhoria de comunicação", "Definição de prioridades", "Orientação NR-17"], "prioridade": "Média", "prazo": "60 dias", "responsavel": "Empresa / Coordenação"},
    },
    {
        "keywords": ["96.02", "9602", "salão", "salao", "beleza", "barbearia", "cabeleireiro", "estética", "estetica", "manicure"],
        "label": "Salão de beleza / estética",
        "categoria": "Serviços pessoais",
        "cnaes_referencia": ["96.02-5-01", "96.02-5-02"],
        "setores_sugeridos": ["ATENDIMENTO", "PROCEDIMENTOS", "RECEPÇÃO", "LIMPEZA"],
        "riscos_sugeridos": ["Postura em pé", "Movimentos repetitivos", "Posturas incômodas", "Atenção visual", "Atendimento ao público"],
        "exames_sugeridos": ["Exame clínico", "Acuidade visual", "Hemograma/TGO/TGP quando houver químicos"],
        "general": {"tipo_documento": "AET completa com formulário ergonômico", "motivo_analise": "Atendimento à NR-17 e análise de atividades com movimentos repetitivos e atendimento", "condicao_ergonomica_geral": "Adequada com recomendações"},
        "sector": {"postura": ["Em pé", "Sentado", "Inclinação/flexão de tronco"], "tipo_atividade": "Atendimento, procedimentos estéticos, organização do posto e uso de materiais", "exigencia_fisica": "Moderada", "exigencia_cognitiva": "Moderada", "ritmo_trabalho": "Por agenda e demanda de clientes", "pausas": "Pausas breves entre atendimentos e alternância postural recomendadas", "mobiliario": "A avaliar cadeiras, bancadas, espelhos e alturas de trabalho", "ambiente": "A avaliar iluminação, ventilação, odores, conforto térmico e circulação", "organizacao": "Atendimento ao público, agenda, comunicação e organização de procedimentos", "equipamentos": "Cadeiras, bancadas, instrumentos, secadores, produtos e materiais de apoio", "fatores": ["Atendimento ao público", "Posturas incômodas", "Movimentos repetitivos", "Atenção contínua"], "medidas": ["Ajuste de mobiliário", "Pausas breves", "Alternância postural", "Organização do posto", "Orientação NR-17"], "prioridade": "Média", "prazo": "30 dias", "responsavel": "Empresa / Administração"},
    },
    {
        "keywords": ["47.3", "473", "posto", "combustível", "combustivel", "gasolina", "lubrificante", "conveniência", "conveniencia"],
        "label": "Posto de combustível / conveniência",
        "categoria": "Comércio e combustíveis",
        "cnaes_referencia": ["47.31-8-00", "47.32-6-00", "47.29-6-02"],
        "setores_sugeridos": ["PISTA", "CAIXA", "CONVENIÊNCIA", "ADMINISTRATIVO"],
        "riscos_sugeridos": ["Postura em pé", "Atendimento ao público", "Deslocamento frequente", "Exigência de atenção", "Comunicação com clientes"],
        "exames_sugeridos": ["Exame clínico", "Acuidade visual", "Hemograma/TGO/TGP quando houver químicos"],
        "general": {"tipo_documento": "AET completa com formulário ergonômico", "motivo_analise": "Atendimento à NR-17 e avaliação das atividades de atendimento e operação", "condicao_ergonomica_geral": "Adequada com recomendações"},
        "sector": {"postura": ["Em pé", "Deslocamento frequente", "Alternado"], "tipo_atividade": "Atendimento, abastecimento, conferência, caixa e organização do ambiente", "exigencia_fisica": "Moderada", "exigencia_cognitiva": "Moderada", "ritmo_trabalho": "Por fluxo de clientes e demanda operacional", "pausas": "Pausas conforme escala, mantendo cobertura operacional", "mobiliario": "A avaliar balcões, caixas, assentos e áreas de apoio", "ambiente": "A avaliar circulação, conforto térmico, exposição externa e organização do posto", "organizacao": "Atendimento contínuo, comunicação e controle de demandas", "equipamentos": "Bombas, sistemas de caixa, balcões, prateleiras e equipamentos de apoio", "fatores": ["Atendimento ao público", "Postura em pé", "Deslocamentos no setor", "Atenção contínua"], "medidas": ["Alternância postural", "Pausas breves", "Organização do posto", "Melhoria de comunicação", "Orientação NR-17"], "prioridade": "Média", "prazo": "30 dias", "responsavel": "Empresa / Administração"},
    },
    {
        "keywords": ["01.", "agro", "agricultura", "pecuária", "pecuaria", "rural", "fazenda", "plantio", "criação de animais"],
        "label": "Atividade rural / agropecuária",
        "categoria": "Rural",
        "cnaes_referencia": ["01.11-3-01", "01.51-2-01", "01.62-8-01"],
        "setores_sugeridos": ["OPERACIONAL", "CAMPO", "MANUTENÇÃO", "ADMINISTRATIVO"],
        "riscos_sugeridos": ["Esforço físico", "Posturas incômodas", "Levantamento de cargas", "Deslocamento frequente", "Exposição ao calor"],
        "exames_sugeridos": ["Exame clínico", "Acuidade visual", "Audiometria quando houver máquinas", "Espirometria quando houver poeira"],
        "general": {"tipo_documento": "AET completa com formulário ergonômico", "motivo_analise": "Atendimento à NR-17 e análise das exigências físicas de atividades rurais", "condicao_ergonomica_geral": "Parcialmente adequada"},
        "sector": {"postura": ["Em pé", "Deslocamento frequente", "Inclinação/flexão de tronco", "Agachamento eventual"], "tipo_atividade": "Atividades operacionais rurais, manejo, apoio, manutenção e deslocamentos", "exigencia_fisica": "Elevada", "exigencia_cognitiva": "Moderada", "ritmo_trabalho": "Por rotina operacional, sazonalidade e condições climáticas", "pausas": "Pausas conforme intensidade física, calor e organização da jornada", "mobiliario": "Não aplicável diretamente; avaliar ferramentas, veículos e áreas de apoio", "ambiente": "A avaliar calor, iluminação natural, terreno, poeira, ruído e áreas de descanso", "organizacao": "Demandas variáveis, deslocamentos e necessidade de comunicação operacional", "equipamentos": "Ferramentas manuais, máquinas, veículos, materiais e equipamentos de apoio", "fatores": ["Esforço físico", "Posturas incômodas", "Deslocamentos no setor", "Manuseio de materiais"], "medidas": ["Pausas breves", "Alternância de tarefas", "Ferramentas adequadas", "Organização do posto", "Orientação NR-17"], "prioridade": "Alta", "prazo": "30 dias", "responsavel": "Empresa / Administração"},
    },
    {
        "keywords": ["33.", "manutenção", "manutencao", "refrigeração", "refrigeracao", "instalação", "instalacao", "máquinas", "maquinas", "equipamentos"],
        "label": "Instalação / manutenção técnica / refrigeração",
        "categoria": "Manutenção",
        "cnaes_referencia": ["33.14-7-07", "43.22-3-02", "33.21-0-00"],
        "setores_sugeridos": ["OPERACIONAL", "MANUTENÇÃO", "TÉCNICO", "ADMINISTRATIVO"],
        "riscos_sugeridos": ["Posturas incômodas", "Esforço físico", "Trabalho em altura", "Atenção contínua", "Levantamento de cargas"],
        "exames_sugeridos": ["Exame clínico", "Acuidade visual", "Audiometria quando houver ruído", "ECG quando aplicável"],
        "general": {"tipo_documento": "AET completa com formulário ergonômico", "motivo_analise": "Atendimento à NR-17 e avaliação de atividades técnicas e de manutenção", "condicao_ergonomica_geral": "Adequada com recomendações"},
        "sector": {"postura": ["Em pé", "Deslocamento frequente", "Inclinação/flexão de tronco", "Agachamento eventual"], "tipo_atividade": "Instalação, manutenção, inspeção, uso de ferramentas e deslocamentos técnicos", "exigencia_fisica": "Moderada", "exigencia_cognitiva": "Moderada", "ritmo_trabalho": "Por ordens de serviço, chamados e demandas técnicas", "pausas": "Pausas conforme intensidade e alternância de tarefas", "mobiliario": "Não aplicável diretamente; avaliar ferramentas, bancadas e acesso aos equipamentos", "ambiente": "A avaliar circulação, ruído, iluminação, calor/frio e espaço de trabalho", "organizacao": "Demandas variáveis, deslocamento, prazos e necessidade de procedimentos", "equipamentos": "Ferramentas manuais, escadas, equipamentos técnicos, instrumentos e materiais", "fatores": ["Posturas incômodas", "Atenção contínua", "Manuseio de materiais", "Deslocamentos no setor"], "medidas": ["Ferramentas adequadas", "Rodízio de tarefas", "Pausas breves", "Organização do posto", "Orientação NR-17"], "prioridade": "Média", "prazo": "30 dias", "responsavel": "Empresa / Administração"},
    },
]

# Catálogo resumido de CNAEs usados na tela de consulta dos modelos de AET.
AET_CNAE_CATALOG = [
    {"cnae": cnae, "modelo": item["label"], "categoria": item.get("categoria", ""), "setores": item.get("setores_sugeridos", []), "riscos": item.get("riscos_sugeridos", []), "exames": item.get("exames_sugeridos", [])}
    for item in AET_CNAE_PRESETS for cnae in item.get("cnaes_referencia", [])
]

# Motor simples de regras técnicas para sugestão de exames por risco.
# O sistema não trava a geração; ele apenas marca/sugere exames já cadastrados
# quando o nome do exame combina com palavras-chave.
EXAM_RULES = [
    {"keywords": ["ruído", "ruido", "audição", "audicao", "nível de pressão sonora", "pressao sonora"], "exams": ["Audiometria", "Exame clínico"]},
    {"keywords": ["poeira", "poeiras", "sílica", "silica", "fumos", "névoa", "nevoa", "respirável", "respiravel", "gases", "vapores"], "exams": ["Espirometria", "Raio X de tórax", "Exame clínico"]},
    {"keywords": ["químico", "quimico", "produto químico", "saneante", "solvente", "hidrocarboneto", "gasolina", "diesel", "óleo", "oleo", "graxa"], "exams": ["Exame clínico", "Hemograma", "TGO", "TGP"]},
    {"keywords": ["biológico", "biologico", "sangue", "vírus", "virus", "bactéria", "bacteria", "fungo", "parasita", "resíduo", "residuo", "infectocontagioso"], "exams": ["Exame clínico", "Hemograma", "Cartão de vacina"]},
    {"keywords": ["altura", "queda", "nível", "nivel", "telhado", "escada"], "exams": ["Exame clínico", "Acuidade visual", "Eletrocardiograma"]},
    {"keywords": ["eletricidade", "elétrico", "eletrico", "choque"], "exams": ["Exame clínico", "Eletrocardiograma"]},
    {"keywords": ["calor", "temperatura elevada", "frio", "câmara frigorífica", "camara frigorifica"], "exams": ["Exame clínico"]},
    {"keywords": ["ergonômico", "ergonomico", "postura", "repetitivo", "carga", "esforço", "esforco", "levantamento", "transporte manual"], "exams": ["Exame clínico", "Anamnese ocupacional"]},
    {"keywords": ["psicossocial", "assédio", "assedio", "estresse", "sobrecarga", "conflito", "comunicação hostil", "comunicacao hostil", "baixa autonomia"], "exams": ["Anamnese psicossocial ocupacional", "SRQ-20", "Exame clínico"]},
    {"keywords": ["trânsito", "transito", "motorista", "direção", "direcao", "veículo", "veiculo"], "exams": ["Exame clínico", "Acuidade visual"]},
]

# Exames complementares padrão usados pelo motor de regras.
# São cadastrados automaticamente se ainda não existirem, para que o botão
# “Sugerir exames pelas regras” funcione mesmo em bases novas.
STANDARD_COMPLEMENTARY_EXAMS = [
    {"exame": "Exame clínico", "periodicidade": "Conforme PCMSO", "admissional": "Sim", "periodico": "Sim", "retorno": "Quando aplicável", "mudanca": "Quando aplicável", "demissional": "Sim"},
    {"exame": "Audiometria", "periodicidade": "Conforme exposição a ruído", "admissional": "Sim, quando exposto", "periodico": "Conforme PCMSO", "retorno": "Quando aplicável", "mudanca": "Quando aplicável", "demissional": "Quando aplicável"},
    {"exame": "Espirometria", "periodicidade": "Conforme exposição respiratória", "admissional": "Sim, quando exposto", "periodico": "Conforme PCMSO", "retorno": "Quando aplicável", "mudanca": "Quando aplicável", "demissional": "Quando aplicável"},
    {"exame": "Acuidade visual", "periodicidade": "Conforme função e risco", "admissional": "Sim, quando aplicável", "periodico": "Conforme PCMSO", "retorno": "Quando aplicável", "mudanca": "Quando aplicável", "demissional": "Quando aplicável"},
    {"exame": "Eletrocardiograma", "periodicidade": "Conforme avaliação médica e risco", "admissional": "Quando aplicável", "periodico": "Conforme PCMSO", "retorno": "Quando aplicável", "mudanca": "Quando aplicável", "demissional": "Quando aplicável"},
    {"exame": "Hemograma", "periodicidade": "Conforme exposição e critério médico", "admissional": "Quando aplicável", "periodico": "Conforme PCMSO", "retorno": "Quando aplicável", "mudanca": "Quando aplicável", "demissional": "Quando aplicável"},
    {"exame": "TGO", "periodicidade": "Conforme exposição química e critério médico", "admissional": "Quando aplicável", "periodico": "Conforme PCMSO", "retorno": "Quando aplicável", "mudanca": "Quando aplicável", "demissional": "Quando aplicável"},
    {"exame": "TGP", "periodicidade": "Conforme exposição química e critério médico", "admissional": "Quando aplicável", "periodico": "Conforme PCMSO", "retorno": "Quando aplicável", "mudanca": "Quando aplicável", "demissional": "Quando aplicável"},
    {"exame": "Raio X de tórax", "periodicidade": "Conforme exposição respiratória e critério médico", "admissional": "Quando aplicável", "periodico": "Conforme PCMSO", "retorno": "Quando aplicável", "mudanca": "Quando aplicável", "demissional": "Quando aplicável"},
    {"exame": "Anamnese ocupacional", "periodicidade": "Conforme PCMSO", "admissional": "Sim", "periodico": "Sim", "retorno": "Quando aplicável", "mudanca": "Quando aplicável", "demissional": "Sim"},
    {"exame": "Anamnese psicossocial ocupacional", "periodicidade": "Conforme avaliação psicossocial ocupacional", "admissional": "Quando aplicável", "periodico": "Conforme PCMSO", "retorno": "Quando aplicável", "mudanca": "Quando aplicável", "demissional": "Quando aplicável"},
    {"exame": "SRQ-20", "periodicidade": "Conforme rastreamento psicossocial ocupacional", "admissional": "Quando aplicável", "periodico": "Conforme PCMSO", "retorno": "Quando aplicável", "mudanca": "Quando aplicável", "demissional": "Quando aplicável"},
    {"exame": "Cartão de vacina", "periodicidade": "Conforme risco biológico e orientação médica", "admissional": "Quando aplicável", "periodico": "Conforme PCMSO", "retorno": "Quando aplicável", "mudanca": "Quando aplicável", "demissional": "Quando aplicável"},
]

def _read_json_list(path: Path) -> list[dict[str, Any]]:
    if not path.exists():
        return []
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
        return data if isinstance(data, list) else []
    except json.JSONDecodeError:
        return []


def _migrate_json_files_if_needed() -> None:
    """Importa cadastros das versões antigas em JSON somente se o banco estiver vazio."""
    if Risk.query.count() == 0:
        for item in _read_json_list(DATA_FILE):
            risk = Risk(
                id=item.get("id") or uuid.uuid4().hex,
                risco=str(item.get("risco", "")).strip(),
                acoes=str(item.get("acoes", "")).strip(),
                indicador=str(item.get("indicador", "")).strip(),
                tipo_risco=str(item.get("tipo_risco", "")).strip(),
                descricao_agente=str(item.get("descricao_agente", "")).strip(),
                possiveis_lesoes=str(item.get("possiveis_lesoes", "")).strip(),
                fontes_circunstancias=str(item.get("fontes_circunstancias", "Durante o processo de trabalho.")).strip(),
                epis=str(item.get("epis", "")).strip(),
                epcs=str(item.get("epcs", "")).strip(),
                grau_severidade=str(item.get("grau_severidade", "")).strip(),
                grau_possibilidade=str(item.get("grau_possibilidade", "")).strip(),
                grau_nivel_risco=str(item.get("grau_nivel_risco", "")).strip(),
            )
            if risk.risco:
                db.session.merge(risk)

    if Sector.query.count() == 0:
        for item in _read_json_list(SETORES_FILE):
            sector = Sector(
                id=item.get("id") or uuid.uuid4().hex,
                setor=str(item.get("setor", "")).strip(),
                cargos=item.get("cargos") if isinstance(item.get("cargos"), list) else [],
            )
            if sector.setor:
                db.session.merge(sector)

    if Exam.query.count() == 0:
        for item in _read_json_list(EXAMES_FILE):
            exam = Exam(
                id=item.get("id") or uuid.uuid4().hex,
                exame=str(item.get("exame", "")).strip(),
                periodicidade=str(item.get("periodicidade", "")).strip(),
                admissional=str(item.get("admissional", "")).strip(),
                periodico=str(item.get("periodico", "")).strip(),
                retorno=str(item.get("retorno", "")).strip(),
                mudanca=str(item.get("mudanca", "")).strip(),
                demissional=str(item.get("demissional", "")).strip(),
            )
            if exam.exame:
                db.session.merge(exam)

    db.session.commit()


def _ensure_schema_columns() -> None:
    """Adiciona colunas novas em bancos já existentes sem apagar dados."""
    inspector = inspect(db.engine)
    tables = inspector.get_table_names()

    def has_column(table_name: str, column_name: str) -> bool:
        if table_name not in tables:
            return False
        return column_name in {col["name"] for col in inspector.get_columns(table_name)}

    dialect = db.engine.dialect.name

    def add_column(table_name: str, column_sqlite: str, column_pg: str | None = None) -> None:
        column_name = column_sqlite.split()[0]
        if has_column(table_name, column_name):
            return
        if dialect == "postgresql":
            stmt = f"ALTER TABLE {table_name} ADD COLUMN IF NOT EXISTS {column_pg or column_sqlite}"
        else:
            stmt = f"ALTER TABLE {table_name} ADD COLUMN {column_sqlite}"
        db.session.execute(text(stmt))

    def alter_company_columns_to_text() -> None:
        # Bancos já criados no Render podem ter campos curtos (VARCHAR).
        # Como alguns campos recebem descrições completas de CNAE/atividade,
        # ampliamos para TEXT sem apagar nenhum dado.
        if dialect != "postgresql":
            return
        text_columns = [
            "nome", "cnpj", "endereco", "bairro_cidade", "cep",
            "cnae1", "descricao1", "grau1", "cnae2", "descricao2", "grau2",
            "funcionarios", "data_atual", "data_final", "email", "fone", "data_avaliacao",
        ]
        for column_name in text_columns:
            if has_column("companies", column_name):
                db.session.execute(text(f"ALTER TABLE companies ALTER COLUMN {column_name} TYPE TEXT"))

    def alter_risk_columns_to_text() -> None:
        # A importação em lote usa textos técnicos longos em vários campos.
        # Em versões antigas alguns campos foram criados como VARCHAR(80/120),
        # o que causava erro no PostgreSQL ao importar planilhas completas.
        if dialect != "postgresql":
            return
        text_columns = [
            "acoes", "indicador", "descricao_agente", "possiveis_lesoes",
            "fontes_circunstancias", "epis", "epcs", "ltcat_meio_propagacao",
            "ltcat_insalubridade", "ltcat_grau_insalubridade",
            "ltcat_aposentadoria_especial", "ltcat_enquadramento_tecnico",
            "ltcat_parecer_previdenciario", "ltcat_periodicidade_jornada",
        ]
        for column_name in text_columns:
            if has_column("risks", column_name):
                db.session.execute(text(f"ALTER TABLE risks ALTER COLUMN {column_name} TYPE TEXT"))

    add_column("sectors", "group_id VARCHAR(32)")
    add_column("sector_groups", "is_temporary BOOLEAN DEFAULT FALSE NOT NULL")
    add_column("sector_groups", "temporary_expires_at TIMESTAMP")
    add_column("risks", "ltcat_meio_propagacao TEXT")
    add_column("risks", "ltcat_insalubridade TEXT")
    add_column("risks", "ltcat_grau_insalubridade TEXT")
    add_column("risks", "ltcat_aposentadoria_especial TEXT")
    add_column("risks", "ltcat_enquadramento_tecnico TEXT")
    add_column("risks", "ltcat_parecer_previdenciario TEXT")
    add_column("risks", "ltcat_periodicidade_jornada TEXT")
    # Evolução segura do banco para o cadastro completo de empresas.
    for col in [
        "nome TEXT", "cnpj TEXT", "endereco TEXT", "bairro_cidade TEXT",
        "cep TEXT", "cnae1 TEXT", "descricao1 TEXT", "grau1 TEXT",
        "cnae2 TEXT", "descricao2 TEXT", "grau2 TEXT", "funcionarios TEXT",
        "data_atual TEXT", "data_final TEXT", "email TEXT", "fone TEXT",
        "data_avaliacao TEXT",
    ]:
        add_column("companies", col)
    alter_company_columns_to_text()
    alter_risk_columns_to_text()
    db.session.commit()


def _ensure_standard_complementary_exams() -> None:
    """Garante uma base mínima de exames complementares para o motor de regras."""
    created = False
    for item in STANDARD_COMPLEMENTARY_EXAMS:
        name = str(item.get("exame", "")).strip()
        if not name:
            continue
        existing = Exam.query.filter(db.func.lower(Exam.exame) == name.lower()).first()
        if existing:
            # Não sobrescreve cadastros editados pelo usuário. Só completa campos vazios.
            for key in ["periodicidade", "admissional", "periodico", "retorno", "mudanca", "demissional"]:
                if not getattr(existing, key, None) and item.get(key):
                    setattr(existing, key, item.get(key))
                    created = True
            continue
        exam = Exam(
            exame=name,
            periodicidade=item.get("periodicidade", "Conforme PCMSO"),
            admissional=item.get("admissional", ""),
            periodico=item.get("periodico", ""),
            retorno=item.get("retorno", ""),
            mudanca=item.get("mudanca", ""),
            demissional=item.get("demissional", ""),
        )
        db.session.add(exam)
        created = True
    if created:
        db.session.commit()


def _cleanup_expired_temp_import_groups() -> None:
    """Remove grupos/setores temporários antigos criados ao aplicar modelos importados."""
    now = datetime.utcnow()
    expired = SectorGroup.query.filter(
        SectorGroup.is_temporary == True,  # noqa: E712
        SectorGroup.temporary_expires_at != None,  # noqa: E711
        SectorGroup.temporary_expires_at < now,
    ).all()
    if not expired:
        return
    for group in expired:
        Sector.query.filter_by(group_id=group.id).delete()
        db.session.delete(group)
    db.session.commit()




def _cleanup_old_import_drafts() -> None:
    """Remove rascunhos antigos de importação para não acumular JSON no banco."""
    now = datetime.utcnow()
    old_drafts = ImportedLaudoDraft.query.filter(
        ImportedLaudoDraft.expires_at != None,  # noqa: E711
        ImportedLaudoDraft.expires_at < now,
    ).all()
    if not old_drafts:
        return
    for draft in old_drafts:
        db.session.delete(draft)
    db.session.commit()


def _create_import_draft(extracted: dict[str, Any]) -> ImportedLaudoDraft:
    draft = ImportedLaudoDraft(
        state=extracted or {},
        expires_at=datetime.utcnow() + timedelta(hours=12),
    )
    db.session.add(draft)
    db.session.commit()
    return draft


def _get_import_draft(draft_id: str | None) -> ImportedLaudoDraft | None:
    if not draft_id:
        return None
    draft = db.session.get(ImportedLaudoDraft, draft_id)
    if draft and draft.expires_at and draft.expires_at < datetime.utcnow():
        db.session.delete(draft)
        db.session.commit()
        return None
    return draft

def _create_temp_import_group(company: Company, template: ImportedLaudoTemplate) -> SectorGroup:
    base = f"TEMP - {company.nome or 'EMPRESA'} - {template.nome or 'MODELO'}"
    short = re.sub(r"\s+", " ", base).strip()[:180]
    suffix = datetime.now().strftime("%d%m%y%H%M")
    name = f"{short} - {suffix}"
    group = SectorGroup(nome=name, is_temporary=True, temporary_expires_at=datetime.utcnow() + timedelta(hours=24))
    db.session.add(group)
    db.session.flush()
    return group


def init_database() -> None:
    db.create_all()
    _ensure_schema_columns()
    _migrate_json_files_if_needed()
    _ensure_standard_complementary_exams()
    _cleanup_expired_temp_import_groups()
    _cleanup_old_import_drafts()


def _field(name: str) -> str:
    return request.form.get(name, "").strip()


def _form_to_risk(existing_id: str | None = None) -> dict[str, Any]:
    return {
        "id": existing_id or uuid.uuid4().hex,
        "risco": _field("risco"),
        "acoes": _field("acoes"),
        "prazo_implantacao": "",
        "prazo_reavaliacao": "",
        "indicador": _field("indicador"),
        "tipo_risco": _field("tipo_risco"),
        "descricao_agente": _field("descricao_agente"),
        "possiveis_lesoes": _field("possiveis_lesoes"),
        "fontes_circunstancias": _field("fontes_circunstancias") or "Durante o processo de trabalho.",
        "epis": _field("epis"),
        "epcs": _field("epcs"),
        "ltcat_meio_propagacao": _field("ltcat_meio_propagacao"),
        "ltcat_insalubridade": _field("ltcat_insalubridade") or "Não",
        "ltcat_grau_insalubridade": _field("ltcat_grau_insalubridade") or "Não aplicável",
        "ltcat_aposentadoria_especial": _field("ltcat_aposentadoria_especial") or "Não",
        "ltcat_enquadramento_tecnico": _field("ltcat_enquadramento_tecnico"),
        "ltcat_parecer_previdenciario": _field("ltcat_parecer_previdenciario"),
        "ltcat_periodicidade_jornada": _field("ltcat_periodicidade_jornada") or "Mensal (<= 4 horas < 10% jornada)",
        "grau_severidade": _field("grau_severidade"),
        "grau_possibilidade": _field("grau_possibilidade"),
        "grau_nivel_risco": _field("grau_nivel_risco"),
    }


def _form_to_sector(existing_id: str | None = None) -> dict[str, Any]:
    cargos = []
    cargo_names = request.form.getlist("cargo_nome[]")
    cargo_cbos = request.form.getlist("cargo_cbo[]")
    cargo_nfuncs = request.form.getlist("cargo_nfunc[]")
    cargo_descricoes = request.form.getlist("cargo_descricao[]")

    total = max(len(cargo_names), len(cargo_cbos), len(cargo_nfuncs), len(cargo_descricoes), 1)
    for index in range(total):
        cargo = {
            "id": uuid.uuid4().hex,
            "cargo": cargo_names[index].strip() if index < len(cargo_names) else "",
            "cbo": cargo_cbos[index].strip() if index < len(cargo_cbos) else "",
            "n_func": cargo_nfuncs[index].strip() if index < len(cargo_nfuncs) else "",
            "descricao": cargo_descricoes[index].strip() if index < len(cargo_descricoes) else "",
        }
        if any(cargo[key] for key in ["cargo", "cbo", "n_func", "descricao"]):
            cargos.append(cargo)

    return {
        "id": existing_id or uuid.uuid4().hex,
        "setor": _field("setor"),
        "grupo_id": _field("grupo_id"),
        "cargos": cargos,
    }


def _form_to_exam(existing_id: str | None = None) -> dict[str, Any]:
    return {
        "id": existing_id or uuid.uuid4().hex,
        "exame": _field("exame"),
        "periodicidade": _field("periodicidade"),
        "admissional": _field("admissional"),
        "periodico": _field("periodico"),
        "retorno": _field("retorno"),
        "mudanca": _field("mudanca"),
        "demissional": _field("demissional"),
    }


def _risk_from_dict(data: dict[str, Any], risk: Risk | None = None) -> Risk:
    risk = risk or Risk(id=data["id"])
    risk.risco = data["risco"]
    risk.acoes = data["acoes"]
    risk.indicador = data["indicador"]
    risk.tipo_risco = data["tipo_risco"]
    risk.descricao_agente = data.get("descricao_agente", "")
    risk.possiveis_lesoes = data["possiveis_lesoes"]
    risk.fontes_circunstancias = data.get("fontes_circunstancias") or "Durante o processo de trabalho."
    risk.epis = data["epis"]
    risk.epcs = data["epcs"]
    risk.ltcat_meio_propagacao = data.get("ltcat_meio_propagacao", "")
    risk.ltcat_insalubridade = data.get("ltcat_insalubridade") or "Não"
    risk.ltcat_grau_insalubridade = data.get("ltcat_grau_insalubridade") or "Não aplicável"
    risk.ltcat_aposentadoria_especial = data.get("ltcat_aposentadoria_especial") or "Não"
    risk.ltcat_enquadramento_tecnico = data.get("ltcat_enquadramento_tecnico", "")
    risk.ltcat_parecer_previdenciario = data.get("ltcat_parecer_previdenciario", "")
    risk.ltcat_periodicidade_jornada = data.get("ltcat_periodicidade_jornada") or "Mensal (<= 4 horas < 10% jornada)"
    risk.grau_severidade = data["grau_severidade"]
    risk.grau_possibilidade = data["grau_possibilidade"]
    risk.grau_nivel_risco = data["grau_nivel_risco"]
    risk.updated_at = datetime.utcnow()
    return risk


def _sector_from_dict(data: dict[str, Any], sector: Sector | None = None) -> Sector:
    sector = sector or Sector(id=data["id"])
    sector.setor = data["setor"]
    sector.group_id = data.get("grupo_id") or None
    sector.cargos = data.get("cargos", [])
    sector.updated_at = datetime.utcnow()
    return sector




def _form_to_company(existing_id: str | None = None) -> dict[str, Any]:
    return {
        "id": existing_id or uuid.uuid4().hex,
        "nome": _field("nome") or _field("empresa"),
        "empresa": _field("nome") or _field("empresa"),
        "cnpj": _field("cnpj"),
        "endereco": _field("endereco"),
        "bairro_cidade": _field("bairro_cidade"),
        "cep": _field("cep"),
        "cnae1": _field("cnae1") or _field("cnae"),
        "descricao1": _field("descricao1") or _field("descricao_atividade"),
        "grau1": _field("grau1") or _field("grau_risco"),
        "cnae2": _field("cnae2") or _field("cnae_secundario"),
        "descricao2": _field("descricao2") or _field("descricao_atividade_secundaria"),
        "grau2": _field("grau2") or _field("grau_risco_secundario"),
        "funcionarios": _field("funcionarios"),
        "data_atual": _field("data_atual"),
        "data_final": _field("data_final"),
        "email": _field("email"),
        "fone": _field("fone"),
        "data_avaliacao": _field("data_avaliacao"),
    }


def _company_from_dict(data: dict[str, Any], company: Company | None = None) -> Company:
    company = company or Company(id=data["id"])
    company.nome = data.get("nome") or data.get("empresa", "")
    company.cnpj = data.get("cnpj", "")
    company.endereco = data.get("endereco", "")
    company.bairro_cidade = data.get("bairro_cidade", "")
    company.cep = data.get("cep", "")
    company.cnae1 = data.get("cnae1", "")
    company.descricao1 = data.get("descricao1", "")
    company.grau1 = data.get("grau1", "")
    company.cnae2 = data.get("cnae2", "")
    company.descricao2 = data.get("descricao2", "")
    company.grau2 = data.get("grau2", "")
    company.funcionarios = data.get("funcionarios", "")
    company.data_atual = data.get("data_atual", "")
    company.data_final = data.get("data_final", "")
    company.email = data.get("email", "")
    company.fone = data.get("fone", "")
    company.data_avaliacao = data.get("data_avaliacao", "")
    company.updated_at = datetime.utcnow()
    return company


def _is_numeric_text(value: str) -> bool:
    value = str(value or "").strip()
    return value == "" or value.isdigit()


def _validate_company(company: dict[str, Any]) -> list[str]:
    errors: list[str] = []
    if not company.get("nome"):
        errors.append("Preencha o nome da empresa.")
    if not company.get("cnpj"):
        errors.append("Preencha o CNPJ da empresa.")

    # O grau de risco é sempre número. Essa validação evita o erro técnico do banco
    # quando a descrição da atividade é colada no campo de grau por engano.
    if not _is_numeric_text(company.get("grau1", "")):
        errors.append("O Grau de risco principal deve conter somente número, como 1, 2, 3 ou 4. Verifique se a descrição da atividade foi colocada no campo correto.")
    if not _is_numeric_text(company.get("grau2", "")):
        errors.append("O Grau de risco secundário deve conter somente número, como 1, 2, 3 ou 4. Verifique se a descrição da atividade secundária foi colocada no campo correto.")
    return errors

def _exam_from_dict(data: dict[str, Any], exam: Exam | None = None) -> Exam:
    exam = exam or Exam(id=data["id"])
    exam.exame = data["exame"]
    exam.periodicidade = data.get("periodicidade", "")
    exam.admissional = data.get("admissional", "")
    exam.periodico = data.get("periodico", "")
    exam.retorno = data.get("retorno", "")
    exam.mudanca = data.get("mudanca", "")
    exam.demissional = data.get("demissional", "")
    exam.updated_at = datetime.utcnow()
    return exam


def _validate_exam(exam: dict[str, Any]) -> list[str]:
    errors: list[str] = []
    if not exam.get("exame"):
        errors.append("Preencha o nome do exame.")
    return errors


def _validate_risk(risk: dict[str, Any]) -> list[str]:
    errors: list[str] = []
    required = {
        "risco": "Perigo / Risco",
        "acoes": "Ações Preventiva / Corretiva",
        "indicador": "Indicador de Efetividade",
        "tipo_risco": "Tipo de Risco",
        "possiveis_lesoes": "Possíveis lesões ou agravos à saúde",
        "fontes_circunstancias": "Fontes ou circunstâncias",
        "epis": "EPI",
        "epcs": "EPC",
        "grau_severidade": "Grau de Severidade",
        "grau_possibilidade": "Grau de Possibilidade",
        "grau_nivel_risco": "Grau de Nível de Risco",
    }
    for key, label in required.items():
        if not risk.get(key):
            errors.append(f"Preencha o campo: {label}.")

    option_checks = {
        "tipo_risco": (FORM_OPTIONS["tipos_risco"], "Tipo de Risco"),
        "grau_severidade": (FORM_OPTIONS["severidades"], "Grau de Severidade"),
        "grau_possibilidade": (FORM_OPTIONS["possibilidades"], "Grau de Possibilidade"),
        "grau_nivel_risco": (FORM_OPTIONS["niveis_risco"], "Grau de Nível de Risco"),
    }
    for key, (valid_options, label) in option_checks.items():
        if risk.get(key) and risk[key] not in valid_options:
            errors.append(f"Selecione uma opção válida em: {label}.")
    return errors


def _validate_sector(sector: dict[str, Any]) -> list[str]:
    errors: list[str] = []
    if not sector.get("setor"):
        errors.append("Preencha o nome do setor.")
    if not sector.get("cargos"):
        errors.append("Adicione pelo menos um cargo ao setor.")
        return errors

    for index, cargo in enumerate(sector.get("cargos", []), start=1):
        if not cargo.get("cargo"):
            errors.append(f"Preencha o nome do cargo na linha {index}.")
        if not cargo.get("cbo"):
            errors.append(f"Preencha o CBO do cargo na linha {index}.")
        if not cargo.get("n_func"):
            errors.append(f"Preencha o número de funcionários na linha {index}.")
        if not cargo.get("descricao"):
            errors.append(f"Preencha a descrição da atividade na linha {index}.")
    return errors


def _sorted_risks() -> list[dict[str, Any]]:
    return [risk.to_dict() for risk in Risk.query.order_by(Risk.risco.asc()).all()]


def _sorted_risk_groups() -> list[dict[str, Any]]:
    return [group.to_dict() for group in RiskGroup.query.order_by(RiskGroup.nome.asc()).all()]


def _sorted_sectors() -> list[dict[str, Any]]:
    return [sector.to_dict() for sector in Sector.query.order_by(Sector.setor.asc()).all()]


def _sorted_exams() -> list[dict[str, Any]]:
    return [exam.to_dict() for exam in Exam.query.order_by(Exam.exame.asc()).all()]


def _sorted_groups() -> list[dict[str, Any]]:
    return [group.to_dict() for group in SectorGroup.query.order_by(SectorGroup.nome.asc()).all()]




def _sorted_companies() -> list[dict[str, Any]]:
    return [company.to_dict() for company in Company.query.order_by(Company.nome.asc()).all()]


def _simple_norm(value: str) -> str:
    import unicodedata
    value = unicodedata.normalize("NFKD", str(value or "")).lower()
    return "".join(ch for ch in value if not unicodedata.combining(ch))




def _cell_text(cell) -> str:
    return re.sub(r"\s+", " ", (cell.text or "")).strip()


def _row_cells_text(row) -> list[str]:
    return [_cell_text(cell) for cell in row.cells]


def _label_key(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", " ", _simple_norm(value)).strip()


def _clean_imported_cargo_and_cbo(raw_cargo: str, raw_cbo: str = "") -> tuple[str, str]:
    """Limpa cargo/CBO extraído de laudos antigos.

    Alguns modelos antigos trazem o CBO dentro da célula do cargo como
    "REPOSITOR (A) – CBO: CBO: 5211-25". Se isso for salvo sem limpeza, o
    Word novo fica como "REPOSITOR (A) – CBO: CBO: 5211-25 – CBO: A DEFINIR".
    Esta rotina separa o nome do cargo e captura o último CBO válido encontrado.
    """
    raw_cargo = re.sub(r"\s+", " ", str(raw_cargo or "").strip())
    raw_cbo = re.sub(r"\s+", " ", str(raw_cbo or "").strip())
    combined = f"{raw_cargo} {raw_cbo}"
    cbo_matches = re.findall(r"\b\d{4}-\d{2}\b", combined)
    cbo = cbo_matches[-1] if cbo_matches else ""

    # Remove tudo a partir do primeiro marcador de CBO no texto do cargo.
    cargo = re.split(r"\s*[–—-]\s*CBO\s*:|\bCBO\s*:", raw_cargo, maxsplit=1, flags=re.I)[0]
    cargo = re.sub(r"\s*[–—-]\s*$", "", cargo).strip()
    cargo = re.sub(r"\s+", " ", cargo)

    if not cbo:
        cbo_from_field = re.search(r"\b\d{4}-\d{2}\b", raw_cbo)
        cbo = cbo_from_field.group(0) if cbo_from_field else "A DEFINIR"
    if not cargo:
        cargo = "A DEFINIR"
    return cargo.upper(), cbo


def _is_noise_import_risk_name(value: str) -> bool:
    """Evita salvar como risco frases fixas de rodapé/controle dos inventários."""
    key = _label_key(value)
    noise_patterns = [
        "controles existentes no ges e sua eficacia",
        "monitoramento da saude do trabalhador atraves de exames ocupacionais",
        "nenhum fator de risco psicossocial foi identificado",
        "especificacao dos perigos fatores de risco",
        "funcoes do grupo de exposicao similar",
        "classificacao",
        "aceitavel",
    ]
    return any(pattern in key for pattern in noise_patterns)


def _first_value_after_label(cells: list[str], label: str) -> str:
    """Retorna o primeiro valor útil depois de uma coluna/rótulo."""
    label_norm = _label_key(label)
    for idx, cell in enumerate(cells):
        if _label_key(cell) == label_norm:
            for nxt in cells[idx + 1:]:
                cleaned = re.sub(r"\s+", " ", nxt or "").strip()
                if cleaned and _label_key(cleaned) != label_norm:
                    return cleaned
    return ""


def _extract_field_from_docx_tables(doc) -> dict[str, str]:
    fields = {
        "empresa": "", "cnpj": "", "endereco": "", "bairro_cidade": "", "cep": "",
        "cnae1": "", "descricao1": "", "grau1": "", "cnae2": "", "descricao2": "",
        "grau2": "", "funcionarios": "", "data_atual": "", "data_final": "", "email": "", "fone": "",
    }
    labels = {
        "EMPRESA": "empresa", "CNPJ": "cnpj", "ENDEREÇO": "endereco", "ENDERECO": "endereco",
        "BAIRRO / CIDADE": "bairro_cidade", "BAIRRO/CIDADE": "bairro_cidade", "CEP": "cep",
        "CNAE": "cnae1", "DESCRIÇÃO DA ATIVIDADE": "descricao1", "DESCRICAO DA ATIVIDADE": "descricao1",
        "GRAU DE RISCO": "grau1", "CNAE (SECUNDÁRIO)": "cnae2", "CNAE (SECUNDARIO)": "cnae2",
        "GRAU DE RISCO (SECUNDÁRIO)": "grau2", "GRAU DE RISCO (SECUNDARIO)": "grau2",
        "FUNCIONÁRIOS": "funcionarios", "FUNCIONARIOS": "funcionarios", "EMAIL": "email", "FONE": "fone",
    }
    # Tabelas de identificação costumam ter duas colunas: rótulo | valor. Mantém ordem para diferenciar descrições primária/secundária.
    descricao_hits = 0
    for table in doc.tables:
        for row in table.rows:
            cells = _row_cells_text(row)
            non_empty = [c for c in cells if c]
            if len(non_empty) < 2:
                continue
            label_raw = non_empty[0].upper().strip()
            value = non_empty[1].strip()
            label_norm = re.sub(r"\s+", " ", label_raw)
            if not value or value.upper() == label_norm:
                continue
            if label_norm in ("DESCRIÇÃO DA ATIVIDADE", "DESCRICAO DA ATIVIDADE"):
                descricao_hits += 1
                key = "descricao1" if descricao_hits == 1 else "descricao2"
            elif label_norm == "CNAE" and fields.get("cnae1"):
                key = "cnae2"
            elif label_norm == "GRAU DE RISCO" and fields.get("grau1"):
                key = "grau2"
            else:
                key = labels.get(label_norm)
            if key and not fields.get(key):
                fields[key] = value
            if label_norm == "VIGÊNCIA" or label_norm == "VIGENCIA":
                dates = re.findall(r"\d{2}/\d{4}|\d{2}/\d{2}/\d{4}|\d{4}", value)
                if dates:
                    fields["data_atual"] = dates[0]
                if len(dates) > 1:
                    fields["data_final"] = dates[1]
    return fields


def _extract_sectors_from_docx_tables(doc) -> list[dict[str, Any]]:
    sectors: list[dict[str, Any]] = []
    for table in doc.tables:
        if len(table.rows) < 3 or len(table.columns) < 3:
            continue
        first = _row_cells_text(table.rows[0])
        sector_name = ""
        if any(_label_key(c) == "nome do setor" for c in first):
            sector_name = _first_value_after_label(first, "Nome do setor")
        if not sector_name:
            continue
        sector_name = sector_name.strip().upper()
        cargos: list[dict[str, str]] = []
        for row in table.rows[2:]:
            cells = _row_cells_text(row)
            if len(cells) < 3:
                continue
            cargo_cell = cells[0]
            cargo_label = _label_key(cargo_cell)
            if (
                not cargo_cell
                or cargo_label in {"funcoes no setor", "funcionarios", "descricao da atividade"}
                or "funcoes no setor" in cargo_label
                or "funcionarios" in cargo_label and "descricao da atividade" in cargo_label
            ):
                continue
            if any(skip in cargo_cell.upper() for skip in ["NOME DO SETOR", "VIGÊNCIA", "VIGENCIA"]):
                continue
            cargo, cbo = _clean_imported_cargo_and_cbo(cargo_cell)
            n_func = cells[1] if len(cells) > 1 else "1"
            desc = cells[2] if len(cells) > 2 else ""
            if cargo and len(cargo) <= 120:
                cargos.append({"cargo": cargo, "cbo": cbo, "n_func": n_func or "1", "descricao": desc or "Atividade importada de laudo antigo; revisar conforme função."})
        if cargos:
            sectors.append({"setor": sector_name, "cargos": cargos})
    # Remove duplicados preservando cargos
    merged: dict[str, dict[str, Any]] = {}
    for item in sectors:
        key = _simple_norm(item["setor"])
        if key not in merged:
            merged[key] = {"setor": item["setor"], "cargos": []}
        existing_cargos = {_simple_norm(c.get("cargo", "")) + "|" + _simple_norm(c.get("cbo", "")) for c in merged[key]["cargos"]}
        for cargo in item.get("cargos", []):
            ckey = _simple_norm(cargo.get("cargo", "")) + "|" + _simple_norm(cargo.get("cbo", ""))
            if ckey not in existing_cargos:
                merged[key]["cargos"].append(cargo)
                existing_cargos.add(ckey)
    return list(merged.values())


def _extract_risks_from_docx_tables(doc) -> list[dict[str, str]]:
    """Extrai riscos técnicos do inventário do PGR/PCMSO/LTCAT antigo.

    O formato dos laudos antigos usa tabelas do Word com células mescladas,
    então a mesma informação aparece repetida em várias colunas. A extração
    abaixo sempre procura o primeiro valor útil depois do rótulo da linha.
    """
    all_by_sector = _extract_sector_risks_from_docx_tables(doc)
    flattened: list[dict[str, str]] = []
    seen: set[str] = set()
    for sector_risks in all_by_sector.values():
        for risk in sector_risks:
            name = re.sub(r"\s+", " ", risk.get("risco", "")).strip()
            if not name or len(name) < 4 or len(name) > 220 or _is_noise_import_risk_name(name):
                continue
            key = _simple_norm(name)
            if key in seen:
                continue
            seen.add(key)
            cleaned = dict(risk)
            cleaned["risco"] = name
            flattened.append(cleaned)
    return flattened


def _is_probable_inventory_table(rows: list[list[str]]) -> bool:
    joined = "\n".join(" | ".join(row) for row in rows[:8])
    norm = _label_key(joined)
    return "perigo fator de risco" in norm and ("descricao do agente" in norm or "especificacao dos perigos" in norm)


def _table_first_meaningful_value(cells: list[str], skip_values: set[str] | None = None) -> str:
    skip_values = skip_values or set()
    for value in cells:
        cleaned = re.sub(r"\s+", " ", str(value or "")).strip()
        if not cleaned:
            continue
        key = _label_key(cleaned)
        if key in skip_values:
            continue
        return cleaned
    return ""


def _row_value_after_repeated_label(cells: list[str], label_norm: str) -> str:
    """Retorna o primeiro valor real de uma linha com rótulos duplicados.

    Ex.: ['Descrição do agente', 'Descrição do agente', 'Descrição do agente', 'Queda...'] -> 'Queda...'
    """
    skip = {label_norm, "prevencao e controle", "exposicao", "classificacao"}
    return _table_first_meaningful_value(cells[1:], skip)


def _extract_sector_name_from_inventory_table(rows: list[list[str]]) -> str:
    for row in rows[:3]:
        values = [re.sub(r"\s+", " ", c).strip() for c in row if re.sub(r"\s+", " ", c).strip()]
        for value in values:
            up = value.upper()
            if 2 <= len(value) <= 80 and not any(skip in up for skip in [
                "FUNÇÕES", "FUNCOES", "GRUPO", "EXPOSIÇÃO", "EXPOSICAO", "PERIGO", "FATOR", "VIGÊNCIA", "VIGENCIA"
            ]):
                return up
    return "SETOR IMPORTADO"


def _extract_sector_risks_from_docx_tables(doc) -> dict[str, list[dict[str, str]]]:
    """Extrai riscos por setor a partir das tabelas do inventário.

    Essa é a parte mais importante da importação reutilizável: o modelo antigo
    deve guardar quais riscos pertenciam a cada setor, para aplicar isso em uma
    empresa nova sem misturar todos os riscos em todos os setores.
    """
    sector_risks: dict[str, list[dict[str, str]]] = {}
    for table in doc.tables:
        if len(table.rows) < 6 or len(table.columns) < 4:
            continue
        rows = [_row_cells_text(row) for row in table.rows]
        if not _is_probable_inventory_table(rows):
            continue
        sector_name = _extract_sector_name_from_inventory_table(rows)
        current: dict[str, str] | None = None
        for cells in rows:
            if not cells:
                continue
            label_norm = _label_key(cells[0])
            # Em alguns modelos o rótulo vem na 2ª/3ª célula por causa de mesclagens.
            if label_norm not in {
                "perigo fator de risco", "descricao do agente", "possiveis lesoes ou agravos a saude",
                "fontes ou circunstancias", "epi", "epc", "medidas administrativas", "severidade"
            }:
                for c in cells[:3]:
                    candidate = _label_key(c)
                    if candidate in {
                        "perigo fator de risco", "descricao do agente", "possiveis lesoes ou agravos a saude",
                        "fontes ou circunstancias", "epi", "epc", "medidas administrativas", "severidade"
                    }:
                        label_norm = candidate
                        break
            if label_norm == "perigo fator de risco":
                if current and current.get("risco"):
                    sector_risks.setdefault(sector_name, []).append(current)
                tipo = _row_value_after_repeated_label(cells, label_norm) or "ERGONÔMICO"
                current = {"tipo_risco": tipo.strip().upper(), "setor_origem": sector_name}
            elif current is not None and label_norm == "descricao do agente":
                current["risco"] = _row_value_after_repeated_label(cells, label_norm)
                current["descricao_agente"] = current.get("risco", "")
            elif current is not None and label_norm == "possiveis lesoes ou agravos a saude":
                current["possiveis_lesoes"] = _row_value_after_repeated_label(cells, label_norm)
            elif current is not None and label_norm == "fontes ou circunstancias":
                current["fontes_circunstancias"] = _row_value_after_repeated_label(cells, label_norm)
            elif current is not None and label_norm == "epi":
                current["epis"] = _row_value_after_repeated_label(cells, label_norm)
            elif current is not None and label_norm == "epc":
                current["epcs"] = _row_value_after_repeated_label(cells, label_norm)
            elif current is not None and label_norm == "medidas administrativas":
                current["acoes"] = _row_value_after_repeated_label(cells, label_norm)
            elif current is not None and label_norm == "severidade":
                # Linha comum: Severidade | INSIGNIFICANTE | Probabilidade | POSSÍVEL | Nível de risco | BAIXO
                current["grau_severidade"] = cells[1].strip().upper() if len(cells) > 1 and cells[1].strip() else "MÉDIO"
                for idx, c in enumerate(cells):
                    cn = _label_key(c)
                    if cn == "probabilidade" and idx + 1 < len(cells):
                        current["grau_possibilidade"] = cells[idx + 1].strip().upper()
                    if cn == "nivel de risco" and idx + 1 < len(cells):
                        current["grau_nivel_risco"] = cells[idx + 1].strip().upper()
        if current and current.get("risco"):
            sector_risks.setdefault(sector_name, []).append(current)

    # Limpa duplicados por setor mantendo os detalhes do primeiro registro.
    cleaned: dict[str, list[dict[str, str]]] = {}
    for sector, risks in sector_risks.items():
        seen: set[str] = set()
        for risk in risks:
            name = re.sub(r"\s+", " ", risk.get("risco", "")).strip()
            if not name or len(name) < 4 or len(name) > 220 or _is_noise_import_risk_name(name):
                continue
            key = _simple_norm(name)
            if key in seen:
                continue
            seen.add(key)
            risk["risco"] = name
            risk.setdefault("tipo_risco", "ERGONÔMICO")
            risk.setdefault("grau_severidade", "MÉDIO")
            risk.setdefault("grau_possibilidade", "POSSÍVEL")
            risk.setdefault("grau_nivel_risco", "MODERADO")
            cleaned.setdefault(sector, []).append(risk)
    return cleaned


def _extract_docx_import_payload(path: Path) -> dict[str, Any]:
    from docx import Document
    doc = Document(str(path))
    company_fields = _extract_field_from_docx_tables(doc)
    sectors_cargos = _extract_sectors_from_docx_tables(doc)
    sector_risks = _extract_sector_risks_from_docx_tables(doc)
    risks = _extract_risks_from_docx_tables(doc)
    return {
        "company_fields": company_fields,
        "sectors_cargos": sectors_cargos,
        "risks_detailed": risks,
        "sector_risks": sector_risks,
    }


def _extract_text_from_upload(file_storage) -> str:
    """Extrai texto básico de DOCX/PDF para importação inteligente de laudos antigos."""
    if not file_storage or not getattr(file_storage, "filename", ""):
        raise ValueError("Envie pelo menos um arquivo antigo em DOCX ou PDF.")
    filename = secure_filename(file_storage.filename)
    ext = Path(filename).suffix.lower()
    with tempfile.TemporaryDirectory() as tmp:
        tmpdir = Path(tmp)
        path = tmpdir / f"{uuid.uuid4().hex}_{filename}"
        file_storage.save(path)
        if ext == ".docx":
            from docx import Document
            doc = Document(str(path))
            payload = _extract_docx_import_payload(path)
            parts: list[str] = []
            parts.append("###IMPORT_JSON###")
            parts.append(json.dumps(payload, ensure_ascii=False))
            parts.append("###END_IMPORT_JSON###")
            parts.extend([p.text for p in doc.paragraphs if (p.text or "").strip()])
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            return "\n".join(parts)
        if ext == ".pdf":
            import fitz
            doc = fitz.open(str(path))
            try:
                return "\n".join(page.get_text("text") for page in doc)
            finally:
                doc.close()
        raise ValueError("Envie apenas arquivos .docx ou .pdf para a importação inteligente.")


def _unique_clean_lines(values: list[str]) -> list[str]:
    result: list[str] = []
    seen: set[str] = set()
    for value in values:
        cleaned = re.sub(r"\s+", " ", str(value or "")).strip(" :-|\t")
        if not cleaned or len(cleaned) < 2:
            continue
        key = cleaned.lower()
        if key not in seen:
            seen.add(key)
            result.append(cleaned)
    return result


def _smart_extract_laudo_data(text_value: str) -> dict[str, Any]:
    """Extrai empresa/CNPJ/setores/cargos/riscos/exames com heurísticas para laudos antigos."""
    text_value = text_value or ""
    imported_payload: dict[str, Any] = {}
    marker = re.search(r"###IMPORT_JSON###\s*(\{.*?\})\s*###END_IMPORT_JSON###", text_value, flags=re.S)
    if marker:
        try:
            imported_payload = json.loads(marker.group(1))
        except Exception:
            imported_payload = {}
        text_value = re.sub(r"###IMPORT_JSON###.*?###END_IMPORT_JSON###", "", text_value, flags=re.S)

    lines = [re.sub(r"\s+", " ", line).strip() for line in text_value.splitlines() if line.strip()]
    company_fields = imported_payload.get("company_fields") or {}
    cnpj = company_fields.get("cnpj") or ""
    if not cnpj:
        cnpj_match = re.search(r"\b\d{2}\.?\d{3}\.?\d{3}/?\d{4}-?\d{2}\b", text_value)
        cnpj = cnpj_match.group(0) if cnpj_match else ""

    empresa = company_fields.get("empresa") or ""
    if not empresa:
        # Prioriza linhas/tabelas de identificação. Evita capturar títulos como RESPONSABILIDADE TÉCNICA.
        for pattern in [
            r"(?:^|\n)EMPRESA\s*\|\s*([^\n|]+)",
            r"(?:^|\n)EMPRESA\s+([A-Z0-9ÁÉÍÓÚÂÊÔÃÕÇ .,&/\-]+(?:LTDA|ME|EPP|EIRELI|S/A)[^\n|]*)",
            r"Raz[aã]o Social[:\s]+(.+)",
        ]:
            match = re.search(pattern, text_value, flags=re.I)
            if match:
                candidate = match.group(1).strip().split("|")[0].strip()
                if not any(skip in candidate.upper() for skip in ["RESPONSABILIDADE", "IDENTIFICAÇÃO", "CONTATO"]):
                    empresa = candidate
                    break
    if not empresa:
        for line in lines[:35]:
            up = line.upper()
            if cnpj and cnpj in line:
                continue
            if any(skip in up for skip in ["PROGRAMA", "PCMSO", "PGR", "LTCAT", "RELATÓRIO", "LAUDO", "DATA", "FONE", "EMAIL", "ELABORAÇÃO", "RESPONSABILIDADE"]):
                continue
            if len(line) >= 5 and ("LTDA" in up or "ME" in up or "EPP" in up or "EIRELI" in up):
                empresa = line
                break

    # Setores/cargos: prioriza tabelas estruturadas do Word antigo.
    setores_cargos = imported_payload.get("sectors_cargos") or []
    setores = [item.get("setor", "") for item in setores_cargos if item.get("setor")]
    if not setores:
        for match in re.finditer(r"(?:Nome do setor|SETOR(?:ES)?|DEPARTAMENTO(?:S)?|GES)\s*[:\-]?\s*([^\n;|]+)", text_value, flags=re.I):
            value = match.group(1).strip()
            value = re.split(r"(?:\s{2,}|\||Total|Categoria|Risco|Cargo|Vigência|VIGÊNCIA)", value)[0].strip()
            for item in re.split(r",|/", value):
                item = item.strip()
                if 2 <= len(item) <= 60 and not any(skip in item.upper() for skip in ["GRUPO", "HOMOGÊNEO", "EXPOSIÇÃO", "TRABALHADOR"]):
                    setores.append(item)

    # Riscos: prioriza blocos estruturados do inventário. Depois cruza com riscos já cadastrados.
    riscos_detalhados = imported_payload.get("risks_detailed") or []
    sector_risks = imported_payload.get("sector_risks") or {}
    if isinstance(sector_risks, dict):
        # Garante que riscos encontrados dentro dos setores também entrem no cadastro global do modelo.
        existing_keys = {_simple_norm(str(item.get("risco", ""))) for item in riscos_detalhados if item.get("risco")}
        for _sector_name, _risk_items in sector_risks.items():
            if not isinstance(_risk_items, list):
                continue
            for _risk in _risk_items:
                if not isinstance(_risk, dict):
                    continue
                _key = _simple_norm(str(_risk.get("risco", "")))
                if _key and _key not in existing_keys:
                    riscos_detalhados.append(_risk)
                    existing_keys.add(_key)
    else:
        sector_risks = {}
    riscos = [r.get("risco", "") for r in riscos_detalhados if r.get("risco")]
    known_risks = [risk.risco for risk in Risk.query.order_by(Risk.risco.asc()).all()]
    norm_text = _simple_norm(text_value)
    for risk_name in known_risks:
        if risk_name and _simple_norm(risk_name) in norm_text:
            riscos.append(risk_name)
    # Captura padrões soltos, mas ignora textos legais e cabeçalhos.
    for match in re.finditer(r"(?:Risco|Perigo|Fator de risco)\s*[:\-]\s*([^\n|;]+)", text_value, flags=re.I):
        val = match.group(1).strip()
        if 4 <= len(val) <= 140 and not any(skip in val.upper() for skip in ["OCUPACIONAL", "GRUPO", "HOMOGÊNEO", "COMBINAÇÃO"]):
            riscos.append(val)

    exames: list[str] = []
    known_exams = [exam.exame for exam in Exam.query.order_by(Exam.exame.asc()).all()]
    for exam_name in known_exams:
        if exam_name and _simple_norm(exam_name) in norm_text:
            exames.append(exam_name)

    return {
        "empresa": empresa,
        "cnpj": cnpj,
        "endereco": company_fields.get("endereco", ""),
        "bairro_cidade": company_fields.get("bairro_cidade", ""),
        "cep": company_fields.get("cep", ""),
        "cnae1": company_fields.get("cnae1", ""),
        "descricao1": company_fields.get("descricao1", ""),
        "grau1": company_fields.get("grau1", ""),
        "cnae2": company_fields.get("cnae2", ""),
        "descricao2": company_fields.get("descricao2", ""),
        "grau2": company_fields.get("grau2", ""),
        "funcionarios": company_fields.get("funcionarios", ""),
        "data_atual": company_fields.get("data_atual", ""),
        "data_final": company_fields.get("data_final", ""),
        "email": company_fields.get("email", ""),
        "fone": company_fields.get("fone", ""),
        "setores": _unique_clean_lines(setores)[:120],
        "setores_cargos": setores_cargos[:120],
        "setores_json": json.dumps(setores_cargos[:120], ensure_ascii=False),
        "riscos": _unique_clean_lines(riscos)[:240],
        "riscos_detalhados": riscos_detalhados[:400],
        "riscos_json": json.dumps(riscos_detalhados[:400], ensure_ascii=False),
        "sector_risks": sector_risks,
        "sector_risks_json": json.dumps(sector_risks, ensure_ascii=False),
        "sector_risk_summary": {k: len(v) for k, v in sector_risks.items() if isinstance(v, list)},
        "exames": _unique_clean_lines(exames)[:100],
        "texto_preview": text_value[:6000],
    }


def _exam_rule_suggestions_for_groups(groups: list[dict[str, Any]]) -> dict[str, list[str]]:
    exams = _sorted_exams()
    suggestions: dict[str, list[str]] = {}
    for group in groups:
        sector = group.get("sector") or {}
        sector_id = sector.get("id", "")
        combined = " ".join([
            str(risk.get("risco", "")) + " " + str(risk.get("tipo_risco", "")) + " " + str(risk.get("descricao_agente", ""))
            for risk in group.get("risks", [])
        ])
        norm_combined = _simple_norm(combined)
        wanted_terms: list[str] = []
        for rule in EXAM_RULES:
            if any(_simple_norm(keyword) in norm_combined for keyword in rule.get("keywords", [])):
                wanted_terms.extend(rule.get("exams", []))
        wanted_terms = _dedupe_preserve_order(wanted_terms)
        matched_ids: list[str] = []
        for exam in exams:
            exam_norm = _simple_norm(exam.get("exame", ""))
            if any(_simple_norm(term) in exam_norm or exam_norm in _simple_norm(term) for term in wanted_terms if term):
                matched_ids.append(exam["id"])
        suggestions[sector_id] = _dedupe_preserve_order(matched_ids)
    return suggestions


def _build_generation_preview(groups: list[dict[str, Any]], company: dict[str, str]) -> dict[str, Any]:
    total_risks = sum(len(group.get("risks", [])) for group in groups)
    total_exams = sum(len(group.get("exams", [])) for group in groups)
    psychosocial = sum(1 for group in groups for risk in group.get("risks", []) if str(risk.get("tipo_risco", "")).strip().upper() == "ERGONÔMICO PSICOSSOCIAL")
    environmental = sum(1 for group in groups for risk in group.get("risks", []) if str(risk.get("tipo_risco", "")).strip().upper() in TIPOS_RISCO_LTCAT)
    warnings = []
    if not groups:
        warnings.append("Nenhum setor selecionado.")
    for group in groups:
        setor = (group.get("sector") or {}).get("setor", "")
        if not group.get("risks"):
            warnings.append(f"Setor {setor}: sem riscos selecionados.")
        if not group.get("exams"):
            warnings.append(f"Setor {setor}: sem exames selecionados para PCMSO.")
    if environmental == 0:
        warnings.append("Nenhum risco ambiental selecionado para LTCAT; os setores sairão como ausência de riscos ambientais.")
    if not company.get("data_criacao_laudo"):
        warnings.append("Data de criação do laudo não preenchida.")
    return {"setores": len(groups), "riscos": total_risks, "exames": total_exams, "psicossociais": psychosocial, "ambientais": environmental, "avisos": warnings}


def _normalize_filename(value: str) -> str:
    value = re.sub(r"[^A-Za-z0-9._-]+", "_", str(value or "").strip())
    return value.strip("_") or "empresa"


def _send_zip_with_cleanup(path: Path, download_name: str):
    response = send_file(path, as_attachment=True, download_name=download_name, mimetype="application/zip")
    @response.call_on_close
    def _cleanup() -> None:
        try:
            path.unlink(missing_ok=True)
        except Exception:
            pass
    return response



def _send_pdf_with_cleanup(path: Path, download_name: str):
    response = send_file(path, as_attachment=True, download_name=download_name, mimetype="application/pdf")
    @response.call_on_close
    def _cleanup() -> None:
        try:
            path.unlink(missing_ok=True)
        except Exception:
            pass
    return response


class _HtmlTableParser(HTMLParser):
    """Leitor simples para arquivos .xls exportados como HTML pelo sistema."""
    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.in_table = 0
        self.in_row = False
        self.in_cell = False
        self.current_row: list[str] = []
        self.current_cell: list[str] = []
        self.rows: list[list[str]] = []

    def handle_starttag(self, tag: str, attrs) -> None:  # noqa: ANN001
        tag = tag.lower()
        if tag == "table":
            self.in_table += 1
        elif self.in_table and tag == "tr":
            self.in_row = True
            self.current_row = []
        elif self.in_table and self.in_row and tag in {"td", "th"}:
            self.in_cell = True
            self.current_cell = []
        elif self.in_cell and tag in {"br", "p"}:
            self.current_cell.append(" ")

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()
        if tag in {"td", "th"} and self.in_cell:
            value = " ".join("".join(self.current_cell).replace("\xa0", " ").split())
            self.current_row.append(value)
            self.current_cell = []
            self.in_cell = False
        elif tag == "tr" and self.in_row:
            if any(str(cell).strip() for cell in self.current_row):
                self.rows.append(self.current_row)
            self.current_row = []
            self.in_row = False
        elif tag == "table" and self.in_table:
            self.in_table -= 1

    def handle_data(self, data: str) -> None:
        if self.in_cell:
            self.current_cell.append(data)


def _receipt_norm_header(value: Any) -> str:
    import unicodedata
    text = unicodedata.normalize("NFKD", str(value or ""))
    text = "".join(ch for ch in text if not unicodedata.combining(ch))
    return re.sub(r"[^a-z0-9]+", "", text.lower())


RECEIPT_PDF_COLUMNS = [
    ("evento", "EVENTO"),
    ("empresa", "empresa"),
    ("nome", "NOME"),
    ("cpf", "CPF"),
    ("tipo", "TIPO"),
    ("status", "STATUS"),
    ("data", "DATA"),
    ("reciboesocial", "Recibo\neSocial"),
    ("recibosefaz", "Recibo\nSefaz"),
]


def _parse_html_table_rows(path: Path) -> list[list[str]]:
    raw = path.read_bytes()
    try:
        html = raw.decode("utf-8")
    except UnicodeDecodeError:
        html = raw.decode("latin-1", errors="ignore")
    parser = _HtmlTableParser()
    parser.feed(html)
    return parser.rows


def _parse_xlsx_table_rows(path: Path) -> list[list[str]]:
    if load_workbook is None:
        raise RuntimeError("A biblioteca openpyxl não está instalada. Verifique o requirements.txt no Render.")
    wb = load_workbook(path, data_only=True, read_only=True)
    ws = wb.active
    rows: list[list[str]] = []
    for row in ws.iter_rows(values_only=True):
        values: list[str] = []
        for value in row:
            if value is None:
                values.append("")
            elif isinstance(value, datetime):
                values.append(value.strftime("%d/%m/%Y"))
            else:
                values.append(str(value).strip())
        if any(values):
            rows.append(values)
    return rows


def _read_receipt_spreadsheet(path: Path) -> list[dict[str, str]]:
    """Lê RELFUNCGERAL (.xls HTML ou .xlsx) e retorna somente as colunas usadas no PDF."""
    raw_head = path.read_bytes()[:512].lstrip().lower()
    ext = path.suffix.lower()
    if raw_head.startswith(b"<!doctype") or raw_head.startswith(b"<html") or b"<table" in raw_head:
        rows = _parse_html_table_rows(path)
    elif ext == ".xlsx":
        rows = _parse_xlsx_table_rows(path)
    else:
        raise ValueError("Arquivo não reconhecido. Envie a planilha RELFUNCGERAL em .xls exportado do sistema ou .xlsx.")

    if not rows:
        raise ValueError("Não encontrei nenhuma tabela na planilha enviada.")

    header_index = None
    header_map: dict[str, int] = {}
    required = {key for key, _ in RECEIPT_PDF_COLUMNS}
    for idx, row in enumerate(rows[:25]):
        normed = [_receipt_norm_header(cell) for cell in row]
        found = {name for name in normed if name in required}
        if len(found) >= 5:
            header_index = idx
            for col_idx, name in enumerate(normed):
                if name in required and name not in header_map:
                    header_map[name] = col_idx
            break
    if header_index is None:
        raise ValueError("Não localizei o cabeçalho da planilha. Confirme se ela possui as colunas EVENTO, empresa, NOME, CPF, TIPO, STATUS, DATA, Recibo eSocial e Recibo Sefaz.")

    data_rows: list[dict[str, str]] = []
    for row in rows[header_index + 1:]:
        item: dict[str, str] = {}
        for key, _label in RECEIPT_PDF_COLUMNS:
            col_idx = header_map.get(key)
            item[key] = str(row[col_idx]).strip() if col_idx is not None and col_idx < len(row) else ""
        if any(item.values()):
            data_rows.append(item)
    if not data_rows:
        raise ValueError("A planilha foi lida, mas não encontrei linhas de recibos para converter.")
    return data_rows


def _wrap_pdf_text(page, text: str, max_width: float, fontname: str, fontsize: float) -> list[str]:  # noqa: ANN001
    import fitz
    raw = str(text or "").replace("\r", "\n")
    explicit_parts = [part.strip() for part in raw.split("\n")]
    lines: list[str] = []

    def append_long_token(token: str) -> None:
        chunk = token
        while chunk:
            cut = 1
            for i in range(1, len(chunk) + 1):
                if fitz.get_text_length(chunk[:i], fontname=fontname, fontsize=fontsize) > max_width:
                    break
                cut = i
            lines.append(chunk[:cut].strip())
            chunk = chunk[cut:]

    for part in explicit_parts:
        part = " ".join(part.split())
        if not part:
            lines.append("")
            continue
        current = ""
        for word in part.split(" "):
            candidate = f"{current} {word}".strip() if current else word
            if fitz.get_text_length(candidate, fontname=fontname, fontsize=fontsize) <= max_width:
                current = candidate
                continue
            if current:
                lines.append(current)
                current = ""
            # Quebra trechos longos, como recibos/UUID, preservando palavras normais com espaços.
            if fitz.get_text_length(word, fontname=fontname, fontsize=fontsize) <= max_width:
                current = word
            else:
                append_long_token(word)
        if current:
            lines.append(current)
    return lines or [""]


def _draw_centered_cell(page, rect, text: str, fontname: str, fontsize: float, color=(0, 0, 0)) -> None:  # noqa: ANN001
    lines = _wrap_pdf_text(page, text, rect.width - 8, fontname, fontsize)
    line_height = fontsize * 1.18
    total_height = line_height * len(lines)
    y = rect.y0 + max(3, (rect.height - total_height) / 2)
    for line in lines:
        text_width = page.get_text_length(line, fontname=fontname, fontsize=fontsize)
        x = rect.x0 + max(3, (rect.width - text_width) / 2)
        page.insert_text((x, y + fontsize), line, fontname=fontname, fontsize=fontsize, color=color)
        y += line_height


def _build_receipts_pdf(rows: list[dict[str, str]], output_path: Path) -> Path:
    import fitz
    if not rows:
        raise ValueError("Não há recibos para gerar o PDF.")
    page_width, page_height = 841.89, 595.28  # A4 paisagem em pontos
    margin_x = 12
    margin_top = 14
    margin_bottom = 14
    inner_width = page_width - (margin_x * 2)
    # Proporções ajustadas para todas as colunas caberem na mesma folha em paisagem.
    proportions = [0.073, 0.115, 0.170, 0.112, 0.100, 0.105, 0.092, 0.105, 0.128]
    widths = [inner_width * p for p in proportions]
    header_height = 38
    body_font = 11
    header_font = 12
    doc = fitz.open()

    def new_page():
        page = doc.new_page(width=page_width, height=page_height)
        x = margin_x
        y = margin_top
        for idx, (_key, label) in enumerate(RECEIPT_PDF_COLUMNS):
            rect = fitz.Rect(x, y, x + widths[idx], y + header_height)
            page.draw_rect(rect, color=(0, 0, 0), fill=(0.36, 0.36, 0.36), width=0.9)
            _draw_centered_cell(page, rect, label, "hebo", header_font, color=(1, 1, 1))
            x += widths[idx]
        return page, margin_top + header_height

    page, y = new_page()
    for item in rows:
        line_counts = []
        for idx, (key, _label) in enumerate(RECEIPT_PDF_COLUMNS):
            line_counts.append(len(_wrap_pdf_text(page, item.get(key, ""), widths[idx] - 8, "helv", body_font)))
        row_height = max(52, max(line_counts) * body_font * 1.28 + 14)
        if y + row_height > page_height - margin_bottom:
            page, y = new_page()
        x = margin_x
        for idx, (key, _label) in enumerate(RECEIPT_PDF_COLUMNS):
            rect = fitz.Rect(x, y, x + widths[idx], y + row_height)
            page.draw_rect(rect, color=(0, 0, 0), width=0.9)
            _draw_centered_cell(page, rect, item.get(key, ""), "helv", body_font, color=(0, 0, 0))
            x += widths[idx]
        y += row_height
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path), deflate=True, garbage=4)
    doc.close()
    return output_path

def _normalize_import_header(value: Any) -> str:
    import re
    import unicodedata

    text_value = str(value or "").strip().lower()
    text_value = unicodedata.normalize("NFKD", text_value)
    text_value = "".join(ch for ch in text_value if not unicodedata.combining(ch))
    text_value = re.sub(r"[^a-z0-9]+", "_", text_value).strip("_")
    return text_value


def _xlsx_rows_from_upload(file_storage) -> tuple[list[dict[str, str]], list[str]]:
    """Lê a primeira aba da planilha enviada e retorna linhas por cabeçalho."""
    errors: list[str] = []
    if load_workbook is None:
        return [], ["A biblioteca openpyxl não está instalada. Rode: pip install -r requirements.txt"]
    if not file_storage or not getattr(file_storage, "filename", ""):
        return [], ["Selecione uma planilha Excel para importar."]
    if not file_storage.filename.lower().endswith(".xlsx"):
        return [], ["Envie uma planilha no formato .xlsx."]

    try:
        wb = load_workbook(file_storage.stream, data_only=True)
    except Exception as exc:
        return [], [f"Não foi possível ler a planilha. Verifique se o arquivo é .xlsx válido. Detalhe: {exc}"]

    ws = wb.active
    header_values = [cell.value for cell in ws[1]]
    headers = [_normalize_import_header(value) for value in header_values]
    if not any(headers):
        return [], ["A primeira linha da planilha precisa conter os cabeçalhos."]

    rows: list[dict[str, str]] = []
    for row_index, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
        if not any(value not in (None, "") for value in row):
            continue
        item = {headers[i]: str(row[i] if row[i] is not None else "").strip() for i in range(min(len(headers), len(row))) if headers[i]}
        item["__linha"] = str(row_index)
        rows.append(item)
    return rows, errors


def _row_get(row: dict[str, str], *aliases: str) -> str:
    for alias in aliases:
        value = row.get(_normalize_import_header(alias), "")
        if value:
            return value.strip()
    return ""


def _truthy_import_value(value: str) -> str:
    """Normaliza marcações simples de planilha para campos textuais."""
    value = str(value or "").strip()
    if value.lower() in {"sim", "s", "x", "1", "true", "verdadeiro"}:
        return "X"
    return value


def _dedupe_preserve_order(items: list[str]) -> list[str]:
    seen = set()
    out: list[str] = []
    for item in items:
        key = item.strip().lower()
        if item and key not in seen:
            seen.add(key)
            out.append(item)
    return out


def _grouped_sectors() -> list[dict[str, Any]]:
    sectors = _sorted_sectors()
    grouped: dict[str, dict[str, Any]] = {}
    for sector in sectors:
        gid = sector.get("grupo_id") or "__sem_grupo__"
        if gid not in grouped:
            grouped[gid] = {
                "id": gid,
                "nome": sector.get("grupo_nome") or "Sem grupo",
                "sectors": [],
            }
        grouped[gid]["sectors"].append(sector)
    return sorted(grouped.values(), key=lambda item: (item["nome"] == "Sem grupo", item["nome"].lower()))


def _generation_form_extras() -> dict[str, str]:
    """Dados específicos da finalização do laudo, preenchidos na tela Gerar laudos."""
    ajuste = request.form.get("ajuste_psicossocial") == "1"
    return {
        "data_criacao_laudo": _field("data_criacao_laudo"),
        "datacriacaolaudo": _field("data_criacao_laudo"),
        "ajuste_psicossocial": "1" if ajuste else "",
        "data_da_revisao": _field("data_da_revisao") if ajuste else "",
        "data_revisao_ajuste": _field("data_da_revisao") if ajuste else "",
    }


def _company_payload_from_form() -> dict[str, str]:
    company_id = _field("company_id")
    if company_id:
        company = db.session.get(Company, company_id)
        if company:
            data = company.to_dict()
            # Permite sobrescrever a vigência se um campo de formulário existir no futuro.
            data["data_atual"] = _field("data_atual") or data.get("data_atual", "")
            data["data_final"] = _field("data_final") or data.get("data_final", "")
            data.update(_generation_form_extras())
            data["aet"] = _aet_form_data_from_request()
            return data
    data = _form_to_company()
    data["empresa"] = data.get("nome", "")
    data["cnae"] = data.get("cnae1", "")
    data["descricao_atividade"] = data.get("descricao1", "")
    data["grau_risco"] = data.get("grau1", "")
    data["cnae_secundario"] = data.get("cnae2", "")
    data["descricao_atividade_secundaria"] = data.get("descricao2", "")
    data["grau_risco_secundario"] = data.get("grau2", "")
    data.update(_generation_form_extras())
    data["aet"] = _aet_form_data_from_request()
    return data

def _selected_risks() -> list[dict[str, Any]]:
    selected_ids = request.form.getlist("risk_ids")
    if not selected_ids:
        return []
    risks = Risk.query.filter(Risk.id.in_(selected_ids)).all()
    order = {risk_id: index for index, risk_id in enumerate(selected_ids)}
    risks.sort(key=lambda item: order.get(item.id, 999999))
    return [risk.to_dict() for risk in risks]


def _selected_sectors() -> list[dict[str, Any]]:
    selected_ids = request.form.getlist("sector_ids") or request.form.getlist("pgr_sector_ids")
    if not selected_ids:
        return []
    sectors = Sector.query.filter(Sector.id.in_(selected_ids)).all()
    order = {sector_id: index for index, sector_id in enumerate(selected_ids)}
    sectors.sort(key=lambda item: order.get(item.id, 999999))
    return [sector.to_dict() for sector in sectors]


def _risk_ids_from_group_ids(group_ids: list[str]) -> list[str]:
    if not group_ids:
        return []
    groups = RiskGroup.query.filter(RiskGroup.id.in_(group_ids)).all()
    risk_ids: list[str] = []
    for group in groups:
        risk_ids.extend([risk.id for risk in group.risks or []])
    return _dedupe_preserve_order(risk_ids)


def _selected_sector_risk_groups() -> tuple[list[dict[str, Any]], list[str]]:
    selected_sector_ids = request.form.getlist("pgr_sector_ids")
    risks = {risk.id: risk.to_dict() for risk in Risk.query.all()}
    sectors = {sector.id: sector.to_dict() for sector in Sector.query.all()}
    exams = {exam.id: exam.to_dict() for exam in Exam.query.all()}

    groups: list[dict[str, Any]] = []
    errors: list[str] = []

    if not selected_sector_ids:
        errors.append("Selecione pelo menos um setor para gerar o PGR/PCMSO.")
        return groups, errors

    for sector_id in selected_sector_ids:
        sector = sectors.get(sector_id)
        if not sector:
            continue
        risk_ids = request.form.getlist(f"sector_risk_ids_{sector_id}")
        group_ids = request.form.getlist(f"sector_risk_group_ids_{sector_id}")
        risk_ids = _dedupe_preserve_order(risk_ids + _risk_ids_from_group_ids(group_ids))
        exam_ids = request.form.getlist(f"sector_exam_ids_{sector_id}")
        selected_risks = [risks[risk_id] for risk_id in risk_ids if risk_id in risks]
        selected_exams = [exams[exam_id] for exam_id in exam_ids if exam_id in exams]
        if not selected_risks:
            errors.append(f"Selecione pelo menos um risco para o setor: {sector.get('setor', '')}.")
        else:
            groups.append({"sector": sector, "risks": selected_risks, "exams": selected_exams})

    return groups, errors


@app.route("/")
def index():
    return redirect(url_for("cadastro"))


@app.route("/cadastro")
def cadastro():
    return render_template("cadastro.html", risks=_sorted_risks(), risk_groups=_sorted_risk_groups(), options=FORM_OPTIONS)


@app.get("/modelo-importacao-riscos")
def download_risk_import_template():
    return send_file(
        RISK_IMPORT_TEMPLATE,
        as_attachment=True,
        download_name="modelo_importacao_riscos.xlsx",
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


@app.post("/importar-riscos")
def import_risks():
    rows, read_errors = _xlsx_rows_from_upload(request.files.get("arquivo"))
    if read_errors:
        for error in read_errors:
            flash(error, "error")
        return redirect(url_for("cadastro"))

    existing_names = {risk.risco.strip().lower() for risk in Risk.query.all()}
    imported = 0
    skipped = 0
    errors: list[str] = []

    for row in rows:
        line = row.get("__linha", "?")
        risk_data = {
            "id": uuid.uuid4().hex,
            "risco": _row_get(row, "risco", "perigo/fator de risco", "perigo", "nome do risco"),
            "acoes": _row_get(row, "ações preventiva / corretiva", "acoes preventiva corretiva", "ações", "acoes"),
            "indicador": _row_get(row, "indicador de efetividade", "indicador"),
            "tipo_risco": _row_get(row, "tipo de risco", "tipo"),
            "descricao_agente": _row_get(row, "descrição do agente", "descricao do agente"),
            "possiveis_lesoes": _row_get(row, "possíveis lesões ou agravos à saúde", "possiveis lesoes ou agravos a saude", "lesões", "lesoes"),
            "fontes_circunstancias": _row_get(row, "fontes ou circunstâncias", "fontes ou circunstancias") or "Durante o processo de trabalho.",
            "epis": _row_get(row, "epi", "epis"),
            "epcs": _row_get(row, "epc", "epcs"),
            "grau_severidade": _row_get(row, "grau de severidade", "severidade"),
            "grau_possibilidade": _row_get(row, "grau de possibilidade", "possibilidade", "probabilidade"),
            "grau_nivel_risco": _row_get(row, "grau de nível de risco", "grau de nivel de risco", "nível de risco", "nivel de risco"),
            "ltcat_meio_propagacao": _row_get(row, "ltcat meio de propagação / via de exposição", "meio de propagação", "via de exposição", "ltcat meio propagacao"),
            "ltcat_periodicidade_jornada": _row_get(row, "ltcat periodicidade / jornada de exposição", "periodicidade / jornada de exposição", "jornada de exposição"),
            "ltcat_insalubridade": _row_get(row, "ltcat insalubridade", "insalubridade") or "Não",
            "ltcat_grau_insalubridade": _row_get(row, "ltcat grau de insalubridade", "grau de insalubridade") or "Não aplicável",
            "ltcat_aposentadoria_especial": _row_get(row, "ltcat aposentadoria especial", "aposentadoria especial") or "Não",
            "ltcat_enquadramento_tecnico": _row_get(row, "ltcat enquadramento técnico", "enquadramento técnico", "enquadramento tecnico"),
            "ltcat_parecer_previdenciario": _row_get(row, "ltcat parecer previdenciário", "parecer previdenciário", "parecer previdenciario"),
        }

        if risk_data["risco"].strip().lower() in existing_names:
            skipped += 1
            continue

        validation_errors = _validate_risk(risk_data)
        if validation_errors:
            skipped += 1
            errors.extend([f"Linha {line}: {error}" for error in validation_errors])
            continue

        db.session.add(_risk_from_dict(risk_data))
        existing_names.add(risk_data["risco"].strip().lower())
        imported += 1

    db.session.commit()
    if imported:
        flash(f"{imported} risco(s) importado(s) com sucesso.", "success")
    if skipped:
        flash(f"{skipped} linha(s) não foram importadas por erro ou duplicidade.", "error")
    for error in errors[:10]:
        flash(error, "error")
    if len(errors) > 10:
        flash(f"Há mais {len(errors) - 10} erro(s) não exibidos. Corrija a planilha e importe novamente.", "error")
    return redirect(url_for("cadastro"))


@app.get("/modelo-importacao-setores")
def download_sector_import_template():
    return send_file(
        SECTOR_IMPORT_TEMPLATE,
        as_attachment=True,
        download_name="modelo_importacao_setores.xlsx",
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


@app.post("/importar-setores")
def import_sectors():
    group_id = _field("grupo_id_importacao")
    group = db.session.get(SectorGroup, group_id) if group_id else None
    if not group:
        flash("Selecione o grupo onde os setores serão cadastrados antes de importar a planilha.", "error")
        return redirect(url_for("setores"))

    rows, read_errors = _xlsx_rows_from_upload(request.files.get("arquivo"))
    if read_errors:
        for error in read_errors:
            flash(error, "error")
        return redirect(url_for("setores"))

    grouped: dict[str, list[dict[str, str]]] = {}
    errors: list[str] = []
    for row in rows:
        line = row.get("__linha", "?")
        setor = _row_get(row, "setor", "nome do setor")
        cargo = _row_get(row, "cargo", "função", "funcao")
        cbo = _row_get(row, "cbo")
        n_func = _row_get(row, "número de funcionários", "numero de funcionarios", "nº funcionários", "n_func", "funcionários")
        descricao = _row_get(row, "descrição da atividade", "descricao da atividade", "atividade")
        if not setor:
            errors.append(f"Linha {line}: preencha o Setor.")
            continue
        if not cargo:
            errors.append(f"Linha {line}: preencha o Cargo.")
            continue
        if not cbo:
            errors.append(f"Linha {line}: preencha o CBO.")
            continue
        if not n_func:
            errors.append(f"Linha {line}: preencha o número de funcionários.")
            continue
        if not descricao:
            errors.append(f"Linha {line}: preencha a descrição da atividade.")
            continue
        grouped.setdefault(setor.strip(), []).append({
            "id": uuid.uuid4().hex,
            "cargo": cargo.strip(),
            "cbo": cbo.strip(),
            "n_func": n_func.strip(),
            "descricao": descricao.strip(),
        })

    imported_sectors = 0
    imported_cargos = 0
    for setor_nome, cargos in grouped.items():
        sector = Sector.query.filter(
            db.func.lower(Sector.setor) == setor_nome.lower(),
            Sector.group_id == group.id,
        ).first()
        if not sector:
            sector = Sector(id=uuid.uuid4().hex, setor=setor_nome, group_id=group.id, cargos=[])
            db.session.add(sector)
            imported_sectors += 1

        existing_cargos = list(sector.cargos) if isinstance(sector.cargos, list) else []
        existing_keys = {(str(c.get("cargo", "")).strip().lower(), str(c.get("cbo", "")).strip().lower()) for c in existing_cargos}
        for cargo in cargos:
            key = (cargo["cargo"].strip().lower(), cargo["cbo"].strip().lower())
            if key not in existing_keys:
                existing_cargos.append(cargo)
                existing_keys.add(key)
                imported_cargos += 1
        sector.cargos = existing_cargos
        sector.updated_at = datetime.utcnow()

    db.session.commit()
    if imported_sectors or imported_cargos:
        flash(f"Importação concluída: {imported_sectors} setor(es) novo(s) e {imported_cargos} cargo(s) cadastrado(s) no grupo {group.nome}.", "success")
    if errors:
        flash(f"{len(errors)} linha(s) não foram importadas por erro.", "error")
        for error in errors[:10]:
            flash(error, "error")
        if len(errors) > 10:
            flash(f"Há mais {len(errors) - 10} erro(s) não exibidos.", "error")
    return redirect(url_for("setores"))


@app.route("/setores")
def setores():
    return render_template("setores.html", sectors=_sorted_sectors(), groups=_sorted_groups(), grouped_sectors=_grouped_sectors())


@app.post("/grupo/novo")
def create_group():
    nome = _field("nome_grupo")
    if not nome:
        flash("Informe o nome do grupo.", "error")
        return redirect(url_for("setores"))
    existing = SectorGroup.query.filter(db.func.lower(SectorGroup.nome) == nome.lower()).first()
    if existing:
        flash("Esse grupo já existe.", "error")
        return redirect(url_for("setores"))
    db.session.add(SectorGroup(nome=nome))
    db.session.commit()
    flash("Grupo cadastrado com sucesso.", "success")
    return redirect(url_for("setores"))


@app.post("/grupo/<group_id>/excluir")
def delete_group(group_id: str):
    group = db.session.get(SectorGroup, group_id)
    if group:
        Sector.query.filter_by(group_id=group_id).update({"group_id": None})
        db.session.delete(group)
        db.session.commit()
        flash("Grupo excluído. Os setores foram mantidos sem grupo.", "success")
    return redirect(url_for("setores"))




@app.route("/empresas")
def empresas():
    return render_template("empresas.html", companies=_sorted_companies())


@app.route("/empresa/nova", methods=["GET", "POST"])
def create_company():
    if request.method == "POST":
        company = _form_to_company()
        errors = _validate_company(company)
        if errors:
            for error in errors:
                flash(error, "error")
            return render_template("empresa_form.html", company=company, title="Nova empresa")
        db.session.add(_company_from_dict(company))
        db.session.commit()
        flash("Empresa cadastrada com sucesso.", "success")
        return redirect(url_for("empresas"))
    today = datetime.now().strftime("%m/%Y")
    next_year = datetime.now().replace(year=datetime.now().year + 1).strftime("%m/%Y")
    return render_template("empresa_form.html", company={"data_atual": today, "data_final": next_year}, title="Nova empresa")


@app.route("/empresa/<company_id>/editar", methods=["GET", "POST"])
def edit_company(company_id: str):
    company_model = db.session.get(Company, company_id)
    if not company_model:
        flash("Empresa não encontrada.", "error")
        return redirect(url_for("empresas"))
    if request.method == "POST":
        updated = _form_to_company(existing_id=company_id)
        errors = _validate_company(updated)
        if errors:
            for error in errors:
                flash(error, "error")
            return render_template("empresa_form.html", company={**company_model.to_dict(), **updated}, title="Editar empresa")
        _company_from_dict(updated, company_model)
        db.session.commit()
        flash("Empresa atualizada com sucesso.", "success")
        return redirect(url_for("empresas"))
    return render_template("empresa_form.html", company=company_model.to_dict(), title="Editar empresa")


@app.post("/empresa/<company_id>/excluir")
def delete_company(company_id: str):
    company_model = db.session.get(Company, company_id)
    if company_model:
        db.session.delete(company_model)
        db.session.commit()
        flash("Empresa excluída.", "success")
    return redirect(url_for("empresas"))

@app.route("/exames")
def exames():
    return render_template("exames.html", exams=_sorted_exams())


@app.route("/exame/novo", methods=["GET", "POST"])
def create_exam():
    if request.method == "POST":
        new_exam = _form_to_exam()
        errors = _validate_exam(new_exam)
        if errors:
            for error in errors:
                flash(error, "error")
            return render_template("exame_form.html", exam=new_exam, title="Novo exame")
        db.session.add(_exam_from_dict(new_exam))
        db.session.commit()
        flash("Exame cadastrado com sucesso.", "success")
        return redirect(url_for("exames"))
    return render_template("exame_form.html", exam={}, title="Novo exame")


@app.route("/exame/<exam_id>/editar", methods=["GET", "POST"])
def edit_exam(exam_id: str):
    exam_model = db.session.get(Exam, exam_id)
    if not exam_model:
        flash("Exame não encontrado.", "error")
        return redirect(url_for("exames"))
    if request.method == "POST":
        updated = _form_to_exam(existing_id=exam_id)
        errors = _validate_exam(updated)
        if errors:
            for error in errors:
                flash(error, "error")
            return render_template("exame_form.html", exam={**exam_model.to_dict(), **updated}, title="Editar exame")
        _exam_from_dict(updated, exam_model)
        db.session.commit()
        flash("Exame atualizado com sucesso.", "success")
        return redirect(url_for("exames"))
    return render_template("exame_form.html", exam=exam_model.to_dict(), title="Editar exame")


@app.post("/exame/<exam_id>/excluir")
def delete_exam(exam_id: str):
    exam_model = db.session.get(Exam, exam_id)
    if exam_model:
        db.session.delete(exam_model)
        db.session.commit()
        flash("Exame excluído.", "success")
    return redirect(url_for("exames"))


def _aet_form_data_from_request() -> dict[str, Any]:
    """Dados detalhados da AET informados no wizard de geração.

    Fica salvo junto da configuração da empresa para permitir regeneração sem
    perder análise ergonômica, observações e conclusões por setor.
    """
    sector_ids = request.form.getlist("pgr_sector_ids")
    general = {
        "tipo_aet": _field("aet_tipo") or "AET completa com formulário ergonômico",
        "tipo_documento": _field("aet_tipo_documento") or _field("aet_tipo") or "AET - Análise Ergonômica do Trabalho",
        "motivo_analise": _field("aet_motivo_analise") or "Atendimento à NR-17 e integração com o PGR",
        "responsavel_tecnico": _field("aet_responsavel_tecnico"),
        "metodologia": request.form.getlist("aet_metodologia"),
        "origem_dados": request.form.getlist("aet_origem_dados"),
        "objetivo_complementar": _field("aet_objetivo_complementar"),
        "criterios_analise": _field("aet_criterios_analise"),
        "limitacoes_analise": _field("aet_limitacoes_analise"),
        "condicao_ergonomica_geral": _field("aet_condicao_ergonomica_geral"),
        "conclusao_geral_manual": _field("aet_conclusao_geral"),
    }
    by_sector: dict[str, Any] = {}
    for sector_id in sector_ids:
        by_sector[sector_id] = {
            "postura_predominante": request.form.getlist(f"aet_postura_{sector_id}"),
            "tipo_atividade": _field(f"aet_tipo_atividade_{sector_id}"),
            "exigencia_fisica": _field(f"aet_exigencia_fisica_{sector_id}"),
            "exigencia_cognitiva": _field(f"aet_exigencia_cognitiva_{sector_id}"),
            "levantamento_cargas": _field(f"aet_levantamento_cargas_{sector_id}"),
            "movimentos_repetitivos": _field(f"aet_movimentos_repetitivos_{sector_id}"),
            "atencao_concentracao": _field(f"aet_atencao_concentracao_{sector_id}"),
            "atendimento_publico": _field(f"aet_atendimento_publico_{sector_id}"),
            "autonomia": _field(f"aet_autonomia_{sector_id}"),
            "metas_prioridades": _field(f"aet_metas_prioridades_{sector_id}"),
            "comunicacao": _field(f"aet_comunicacao_{sector_id}"),
            "ritmo_trabalho": _field(f"aet_ritmo_trabalho_{sector_id}"),
            "pausas": _field(f"aet_pausas_{sector_id}"),
            "mobiliario": _field(f"aet_mobiliario_{sector_id}"),
            "ambiente": _field(f"aet_ambiente_{sector_id}"),
            "organizacao": _field(f"aet_organizacao_{sector_id}"),
            "equipamentos": _field(f"aet_equipamentos_{sector_id}"),
            "fatores_organizacionais": request.form.getlist(f"aet_fatores_organizacionais_{sector_id}"),
            "medidas_recomendadas": request.form.getlist(f"aet_medidas_recomendadas_{sector_id}"),
            "queixas": _field(f"aet_queixas_{sector_id}"),
            "observacoes": _field(f"aet_observacoes_{sector_id}"),
            "recomendacoes": _field(f"aet_recomendacoes_{sector_id}"),
            "prioridade": _field(f"aet_prioridade_{sector_id}"),
            "prazo": _field(f"aet_prazo_{sector_id}"),
            "responsavel": _field(f"aet_responsavel_{sector_id}"),
            "conclusao_setor": _field(f"aet_conclusao_setor_{sector_id}"),
        }
    return {"general": general, "by_sector": by_sector}


def _gerar_form_state_from_request() -> dict[str, Any]:
    sector_ids = request.form.getlist("pgr_sector_ids")
    risks_by_sector = {sector_id: request.form.getlist(f"sector_risk_ids_{sector_id}") for sector_id in sector_ids}
    risk_groups_by_sector = {sector_id: request.form.getlist(f"sector_risk_group_ids_{sector_id}") for sector_id in sector_ids}
    exams_by_sector = {sector_id: request.form.getlist(f"sector_exam_ids_{sector_id}") for sector_id in sector_ids}
    return {
        "company_id": _field("company_id"),
        "data_criacao_laudo": _field("data_criacao_laudo"),
        "ajuste_psicossocial": "1" if request.form.get("ajuste_psicossocial") == "1" else "",
        "data_da_revisao": _field("data_da_revisao"),
        "profile_id": _field("profile_id"),
        "profile_name": _field("profile_name"),
        "selected_sector_ids": sector_ids,
        "selected_risk_ids_by_sector": risks_by_sector,
        "selected_risk_group_ids_by_sector": risk_groups_by_sector,
        "selected_exam_ids_by_sector": exams_by_sector,
        "aet": _aet_form_data_from_request(),
    }




def _report_profile_state_from_request() -> dict[str, Any]:
    """Estado recarregável da tela Gerar laudos."""
    return _gerar_form_state_from_request()


def _sorted_report_profiles() -> list[dict[str, Any]]:
    return [profile.to_dict() for profile in ReportProfile.query.order_by(ReportProfile.updated_at.desc()).all()]


def _sorted_imported_laudo_templates() -> list[dict[str, Any]]:
    return [item.to_dict() for item in ImportedLaudoTemplate.query.order_by(ImportedLaudoTemplate.updated_at.desc()).all()]


def _upsert_import_risk(detail: dict[str, Any] | None, fallback_name: str = "") -> Risk | None:
    detail = detail or {}
    name = re.sub(r"\s+", " ", str(detail.get("risco") or fallback_name or "").strip())
    if not name or _is_noise_import_risk_name(name):
        return None
    existing = Risk.query.filter(db.func.lower(Risk.risco) == name.lower()).first()
    if existing:
        return existing
    risk = Risk(
        risco=name,
        acoes=detail.get("acoes") or "Revisar ações preventivas/corretivas conforme atividade e atualizar com treinamentos NR aplicáveis.",
        indicador="Acompanhar implementação das medidas, registros de orientação e ausência de ocorrências relacionadas.",
        tipo_risco=(detail.get("tipo_risco") or "ERGONÔMICO").strip().upper(),
        descricao_agente=detail.get("risco") or name,
        possiveis_lesoes=detail.get("possiveis_lesoes") or "Revisar possíveis lesões ou agravos conforme o risco importado.",
        fontes_circunstancias=detail.get("fontes_circunstancias") or "Informação importada de laudo antigo; revisar fontes ou circunstâncias.",
        epis=detail.get("epis") or "A definir conforme avaliação técnica.",
        epcs=detail.get("epcs") or "A definir conforme avaliação técnica.",
        grau_severidade=(detail.get("grau_severidade") or "MÉDIO").strip().upper(),
        grau_possibilidade=(detail.get("grau_possibilidade") or "POSSÍVEL").strip().upper(),
        grau_nivel_risco=(detail.get("grau_nivel_risco") or "MODERADO").strip().upper(),
    )
    db.session.add(risk)
    db.session.flush()
    return risk


def _upsert_import_exam(name: str) -> Exam | None:
    name = re.sub(r"\s+", " ", str(name or "").strip())
    if not name:
        return None
    existing = Exam.query.filter(db.func.lower(Exam.exame) == name.lower()).first()
    if existing:
        return existing
    exam = Exam(exame=name, periodicidade="Conforme PCMSO")
    db.session.add(exam)
    db.session.flush()
    return exam


def _upsert_import_sector(item: dict[str, Any], group_id: str | None = None) -> Sector | None:
    setor_name = re.sub(r"\s+", " ", str(item.get("setor", "")).strip()).upper()
    if not setor_name:
        return None
    cargos_raw = item.get("cargos") or []
    cargos: list[dict[str, str]] = []
    for cargo in cargos_raw:
        cargo_nome, cbo_limpo = _clean_imported_cargo_and_cbo(cargo.get("cargo", ""), cargo.get("cbo", ""))
        cargos.append({
            "id": uuid.uuid4().hex,
            "cargo": cargo_nome or "A DEFINIR",
            "cbo": cbo_limpo or "A DEFINIR",
            "n_func": str(cargo.get("n_func") or "1").strip(),
            "descricao": str(cargo.get("descricao") or "Atividade importada de laudo antigo; revisar e detalhar conforme função.").strip(),
        })
    if not cargos:
        cargos = [{"id": uuid.uuid4().hex, "cargo": "A DEFINIR", "cbo": "A DEFINIR", "n_func": "1", "descricao": "Atividades importadas de laudo antigo; revisar e detalhar conforme função."}]
    query = Sector.query.filter(db.func.lower(Sector.setor) == setor_name.lower())
    if group_id:
        query = query.filter(Sector.group_id == group_id)
    else:
        query = query.filter((Sector.group_id == None) | (Sector.group_id == ""))  # noqa: E711
    sector = query.first()
    if not sector:
        sector = Sector(setor=setor_name, group_id=group_id, cargos=cargos)
        db.session.add(sector)
        db.session.flush()
        return sector
    if group_id:
        sector.group_id = group_id
    old_cargos = sector.cargos or []
    old_keys = {(_simple_norm(str(c.get("cargo", ""))) + "|" + _simple_norm(str(c.get("cbo", "")))) for c in old_cargos}
    merged = list(old_cargos)
    changed = False
    if not old_cargos or all(str(c.get("cargo", "")).upper() == "A DEFINIR" for c in old_cargos):
        sector.cargos = cargos
        changed = True
    else:
        for cargo_item in cargos:
            key = _simple_norm(str(cargo_item.get("cargo", ""))) + "|" + _simple_norm(str(cargo_item.get("cbo", "")))
            if key not in old_keys:
                merged.append(cargo_item)
                old_keys.add(key)
                changed = True
        if changed:
            sector.cargos = merged
    db.session.flush()
    return sector


def _apply_imported_template_to_company(template: ImportedLaudoTemplate, company_id: str, group_id: str | None = None) -> ReportProfile:
    state = template.state or {}
    sectors_data = state.get("setores_cargos") or []
    if not sectors_data:
        sectors_data = [{"setor": name, "cargos": []} for name in (state.get("setores") or [])]

    sector_risks_state = state.get("sector_risks") or {}
    if not isinstance(sector_risks_state, dict):
        sector_risks_state = {}

    risk_details = state.get("riscos_detalhados") or []
    risk_names = state.get("riscos") or []
    # Une riscos detalhados, riscos simples e riscos vinculados a setores.
    details_by_norm = {_simple_norm(str(item.get("risco", ""))): item for item in risk_details if item.get("risco")}
    merged_risk_details: list[dict[str, Any]] = []
    seen_risks: set[str] = set()
    for item in risk_details:
        key = _simple_norm(str(item.get("risco", "")))
        if key and key not in seen_risks:
            merged_risk_details.append(item)
            seen_risks.add(key)
    for _sector_name, sector_risk_items in sector_risks_state.items():
        if not isinstance(sector_risk_items, list):
            continue
        for item in sector_risk_items:
            if not isinstance(item, dict):
                continue
            key = _simple_norm(str(item.get("risco", "")))
            if key and key not in seen_risks:
                merged_risk_details.append(item)
                seen_risks.add(key)
                details_by_norm[key] = item
    for name in risk_names:
        key = _simple_norm(name)
        if key and key not in seen_risks:
            merged_risk_details.append(details_by_norm.get(key) or {"risco": name})
            seen_risks.add(key)

    sectors: list[Sector] = []
    sector_by_norm_name: dict[str, Sector] = {}
    for item in sectors_data:
        sector = _upsert_import_sector(item, group_id=group_id)
        if sector:
            sectors.append(sector)
            sector_by_norm_name[_simple_norm(sector.setor)] = sector

    # Se o inventário trouxe um setor com riscos mas a relação função x atividade não trouxe,
    # cria o setor mínimo para manter o vínculo risco/setor do modelo antigo.
    for sector_name in sector_risks_state.keys():
        key = _simple_norm(sector_name)
        if key and key not in sector_by_norm_name:
            sector = _upsert_import_sector({"setor": sector_name, "cargos": []}, group_id=group_id)
            if sector:
                sectors.append(sector)
                sector_by_norm_name[_simple_norm(sector.setor)] = sector

    risks: list[Risk] = []
    risk_by_norm_name: dict[str, Risk] = {}
    for detail in merged_risk_details:
        risk = _upsert_import_risk(detail)
        if risk:
            risks.append(risk)
            risk_by_norm_name[_simple_norm(risk.risco)] = risk

    exams: list[Exam] = []
    for exam_name in (state.get("exames") or []):
        exam = _upsert_import_exam(exam_name)
        if exam:
            exams.append(exam)

    company = db.session.get(Company, company_id)
    today = (company.data_atual if company else "") or state.get("data_atual") or datetime.now().strftime("%m/%Y")
    selected_sector_ids = [s.id for s in sectors]
    all_risk_ids = [r.id for r in risks]
    exam_ids = [e.id for e in exams]

    selected_risk_ids_by_sector: dict[str, list[str]] = {}
    selected_exam_ids_by_sector: dict[str, list[str]] = {}
    selected_risk_group_ids_by_sector: dict[str, list[str]] = {}
    for sector in sectors:
        sector_key = _simple_norm(sector.setor)
        matched_key = None
        for original_sector_name in sector_risks_state.keys():
            if _simple_norm(original_sector_name) == sector_key:
                matched_key = original_sector_name
                break
        sector_specific_risks = sector_risks_state.get(matched_key, []) if matched_key else []
        ids: list[str] = []
        if isinstance(sector_specific_risks, list) and sector_specific_risks:
            for detail in sector_specific_risks:
                if not isinstance(detail, dict):
                    continue
                name = str(detail.get("risco", ""))
                risk = risk_by_norm_name.get(_simple_norm(name)) or _upsert_import_risk(detail)
                if risk:
                    risk_by_norm_name[_simple_norm(risk.risco)] = risk
                    if risk.id not in ids:
                        ids.append(risk.id)
        else:
            # Compatibilidade para modelos importados antigos que ainda não tinham vínculo por setor.
            ids = list(all_risk_ids)
        selected_risk_ids_by_sector[sector.id] = ids
        selected_exam_ids_by_sector[sector.id] = list(exam_ids)
        selected_risk_group_ids_by_sector[sector.id] = []

    profile_state = {
        "company_id": company_id,
        "data_criacao_laudo": today,
        "ajuste_psicossocial": "",
        "data_da_revisao": "",
        "profile_id": "",
        "profile_name": f"Importado de {template.nome}",
        "selected_sector_ids": selected_sector_ids,
        "selected_risk_ids_by_sector": selected_risk_ids_by_sector,
        "selected_risk_group_ids_by_sector": selected_risk_group_ids_by_sector,
        "selected_exam_ids_by_sector": selected_exam_ids_by_sector,
        "aet": {"general": {}, "by_sector": {}},
        "imported_template_id": template.id,
    }
    profile = ReportProfile(
        company_id=company_id,
        nome=f"Importado: {template.nome}",
        data_criacao_laudo=today,
        ajuste_psicossocial="",
        data_da_revisao="",
        state=profile_state,
    )
    db.session.add(profile)
    db.session.commit()
    profile.state["profile_id"] = profile.id
    db.session.commit()
    return profile


def _save_report_profile_from_form(auto: bool = False) -> ReportProfile | None:
    company_id = _field("company_id")
    if not company_id:
        return None
    company = db.session.get(Company, company_id)
    if not company:
        return None

    state = _report_profile_state_from_request()
    data_criacao = state.get("data_criacao_laudo", "")
    timestamp = datetime.now().strftime("%d/%m/%Y %H:%M")
    default_name = f"Geração {timestamp}"
    nome = _field("profile_name") or default_name

    # Modo automático: mantém uma configuração padrão por empresa, atualizada
    # sempre que o usuário finaliza um laudo completo.
    profile_id = _field("profile_id")
    profile = db.session.get(ReportProfile, profile_id) if profile_id else None
    if profile is None and auto:
        profile = ReportProfile.query.filter_by(company_id=company_id, nome="Última geração salva").first()
        nome = "Última geração salva"

    if profile is None:
        profile = ReportProfile(id=uuid.uuid4().hex, company_id=company_id, nome=nome)
        db.session.add(profile)
    else:
        profile.company_id = company_id
        profile.nome = nome

    profile.data_criacao_laudo = data_criacao
    profile.ajuste_psicossocial = state.get("ajuste_psicossocial", "")
    profile.data_da_revisao = state.get("data_da_revisao", "")
    profile.state = state
    profile.updated_at = datetime.utcnow()
    db.session.commit()
    return profile


def _month_year_from_text(value: str) -> tuple[str, str]:
    """Extrai mês por extenso e ano de textos como 05/06/2026, 06/2026 ou Junho/2026."""
    months = {
        1: "JANEIRO", 2: "FEVEREIRO", 3: "MARÇO", 4: "ABRIL", 5: "MAIO", 6: "JUNHO",
        7: "JULHO", 8: "AGOSTO", 9: "SETEMBRO", 10: "OUTUBRO", 11: "NOVEMBRO", 12: "DEZEMBRO",
    }
    value = (value or "").strip()
    year_match = re.search(r"(20\d{2})", value)
    year = year_match.group(1) if year_match else str(datetime.now().year)
    nums = re.findall(r"\d+", value)
    month_num = None
    if len(nums) >= 2:
        # Em dd/mm/aaaa usa o segundo número; em mm/aaaa usa o primeiro.
        month_num = int(nums[1]) if len(nums[0]) <= 2 and len(nums[1]) <= 2 else int(nums[0])
    elif len(nums) == 1 and len(nums[0]) <= 2:
        month_num = int(nums[0])
    if month_num and 1 <= month_num <= 12:
        return months[month_num], year
    normalized = value.upper()
    for name in months.values():
        if name in normalized:
            return name, year
    return months[datetime.now().month], year


def _replace_simple_docx_text(doc, replacements: dict[str, str]) -> None:
    """Substituição simples preservando o estilo do primeiro run quando necessário."""
    from copy import deepcopy

    def replace_paragraph(paragraph):
        full_text = paragraph.text or ""
        if not any(key in full_text for key in replacements):
            return
        for run in paragraph.runs:
            text_value = run.text
            for key, val in replacements.items():
                text_value = text_value.replace(key, val)
            run.text = text_value
        full_text = paragraph.text or ""
        if any(key in full_text for key in replacements):
            for key, val in replacements.items():
                full_text = full_text.replace(key, val)
            rpr = deepcopy(paragraph.runs[0]._r.rPr) if paragraph.runs and paragraph.runs[0]._r.rPr is not None else None
            paragraph.clear()
            run = paragraph.add_run(full_text)
            if rpr is not None:
                run._r.insert(0, rpr)

    def walk_part(part):
        for paragraph in part.paragraphs:
            replace_paragraph(paragraph)
        for table in part.tables:
            for row in table.rows:
                for cell in row.cells:
                    walk_part(cell)

    walk_part(doc)
    for section in doc.sections:
        for part in [section.header, section.footer, section.first_page_header, section.first_page_footer, section.even_page_header, section.even_page_footer]:
            walk_part(part)


def _prepare_link_docx(template_path: Path, output_path: Path, empresa: str, data_criacao: str, mes_extenso: str | None = None) -> Path:
    from docx import Document
    doc = Document(str(template_path))
    mes, ano = _month_year_from_text(data_criacao)
    if mes_extenso:
        mes = mes_extenso.strip().upper()
    replacements = {
        "T E M NAKASHIMA - ME": empresa or "",
        "{{EMPRESA}}": empresa or "",
        "{{MES DE CRIAÇÃO POR EXTENSO}}": mes,
        "{{MES DE CRIAÇÃO POR\nEXTENSO}}": mes,
        "{{ANO DE CRIAÇÃO}}": ano,
    }
    _replace_simple_docx_text(doc, replacements)
    # Os modelos enviados tinham o ano 2026 fixo; ajusta apenas nas páginas de link.
    for paragraph in doc.paragraphs:
        if "INCLUSÃO NO PGR EM" in (paragraph.text or "") and f"DE {ano}" not in paragraph.text:
            for run in paragraph.runs:
                run.text = run.text.replace("DE 2026", f"DE {ano}")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def _clear_header_footer_part(part) -> None:
    """Remove cabeçalho/rodapé herdado ao inserir anexos no Word final."""
    try:
        part.is_linked_to_previous = False
    except Exception:
        pass
    element = getattr(part, "_element", None)
    if element is not None:
        for child in list(element):
            element.remove(child)
    try:
        part.add_paragraph()
    except Exception:
        pass


def _clear_section_headers_footers(section) -> None:
    try:
        section.different_first_page_header_footer = True
    except Exception:
        pass
    for part in [
        section.header,
        section.footer,
        section.first_page_header,
        section.first_page_footer,
        section.even_page_header,
        section.even_page_footer,
    ]:
        _clear_header_footer_part(part)


def _render_pdf_pages_to_docx_body(doc, pdf_path: Path, images_dir: Path) -> None:
    """Insere as páginas do PDF como imagens no corpo do DOCX atual.

    Usado para preservar o relatório psicossocial exatamente como o PDF.
    """
    try:
        import fitz  # PyMuPDF
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except Exception as exc:  # pragma: no cover
        raise RuntimeError("A biblioteca PyMuPDF/python-docx não está instalada. Verifique o requirements.txt no Render.") from exc

    images_dir.mkdir(parents=True, exist_ok=True)
    pdf = fitz.open(str(pdf_path))
    if pdf.page_count == 0:
        raise ValueError("O PDF do Relatório Psicossocial não possui páginas.")
    if pdf.page_count > PSICOSSOCIAL_MAX_PAGES:
        raise ValueError(f"O PDF possui {pdf.page_count} páginas. O limite atual é {PSICOSSOCIAL_MAX_PAGES} páginas para evitar travamentos no Render.")

    first_rect = pdf[0].rect
    section = doc.sections[-1]
    page_margin = 0.05
    section.page_width = Inches(first_rect.width / 72)
    section.page_height = Inches(first_rect.height / 72)
    section.top_margin = Inches(page_margin)
    section.bottom_margin = Inches(page_margin)
    section.left_margin = Inches(page_margin)
    section.right_margin = Inches(page_margin)
    section.header_distance = Inches(0)
    section.footer_distance = Inches(0)
    _clear_section_headers_footers(section)

    zoom = PSICOSSOCIAL_RENDER_DPI / 72
    matrix = fitz.Matrix(zoom, zoom)

    try:
        for page_index in range(pdf.page_count):
            page = pdf[page_index]
            image_path = images_dir / f"pagina_{page_index + 1:03d}.jpg"
            pix = page.get_pixmap(matrix=matrix, alpha=False, colorspace=fitz.csRGB)
            pix.save(str(image_path), output="jpeg", jpg_quality=PSICOSSOCIAL_JPEG_QUALITY)

            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1
            run = paragraph.add_run()

            rect = page.rect
            max_width_in = (first_rect.width / 72) - (page_margin * 2)
            max_height_in = (first_rect.height / 72) - (page_margin * 2) - 0.03
            width_in = rect.width / 72 if rect.width else max_width_in
            height_in = rect.height / 72 if rect.height else max_height_in
            scale = min(max_width_in / width_in, max_height_in / height_in, 1)
            run.add_picture(str(image_path), width=Inches(width_in * scale))
    finally:
        pdf.close()


def _append_pdf_as_images_to_docx(base_docx: Path, pdf_path: Path, output_docx: Path) -> Path:
    try:
        from docx import Document
        from docx.enum.section import WD_SECTION
    except Exception as exc:  # pragma: no cover
        raise RuntimeError("A biblioteca python-docx não está instalada. Verifique o requirements.txt no Render.") from exc

    doc = Document(str(base_docx))
    doc.add_section(WD_SECTION.NEW_PAGE)
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    # As imagens temporárias são incorporadas ao DOCX no momento do save.
    # Depois disso, podem ser removidas para não encher o disco do Render.
    with tempfile.TemporaryDirectory() as img_tmp:
        images_dir = Path(img_tmp) / "psicossocial_paginas"
        _render_pdf_pages_to_docx_body(doc, pdf_path, images_dir)
        doc.save(str(output_docx))
    return output_docx


def _convert_pdf_to_docx(pdf_path: Path, output_docx: Path) -> Path:
    """Converte PDF para DOCX preservando o visual de cada página.

    A conversão antiga por extração de texto/tabelas quebrava o layout dos
    relatórios psicossociais, criando sobreposições, faixas pretas e tabelas
    deformadas. Para documentos que precisam apenas ser anexados ao Word final,
    a forma mais fiel é rasterizar cada página do PDF e inserir a página inteira
    como imagem em um DOCX. O conteúdo fica não editável, mas visualmente igual
    ao PDF original.
    """
    try:
        import fitz  # PyMuPDF
        from docx import Document
        from docx.shared import Inches, Pt
    except Exception as exc:  # pragma: no cover
        raise RuntimeError("A biblioteca PyMuPDF/python-docx não está instalada. Verifique o requirements.txt no Render.") from exc

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    images_dir = output_docx.parent / f"{output_docx.stem}_paginas"
    images_dir.mkdir(parents=True, exist_ok=True)

    pdf = fitz.open(str(pdf_path))
    if pdf.page_count == 0:
        raise ValueError("O PDF do Relatório Psicossocial não possui páginas.")

    doc = Document()
    section = doc.sections[0]

    # Usa o tamanho real da primeira página do PDF. A maioria dos relatórios
    # psicossociais é gerada em tamanho único.
    first_rect = pdf[0].rect
    page_width = Inches(first_rect.width / 72)
    page_height = Inches(first_rect.height / 72)
    section.page_width = page_width
    section.page_height = page_height
    # Margem mínima para evitar que o Word crie páginas em branco por causa
    # da altura da imagem somada ao parágrafo. O PDF já possui suas próprias
    # margens visuais, então essa redução não altera a aparência do conteúdo.
    page_margin = 0.05
    section.top_margin = Inches(page_margin)
    section.bottom_margin = Inches(page_margin)
    section.left_margin = Inches(page_margin)
    section.right_margin = Inches(page_margin)
    section.header_distance = Inches(0)
    section.footer_distance = Inches(0)
    _clear_section_headers_footers(section)

    # 200 dpi é um bom equilíbrio entre fidelidade visual e tamanho do arquivo.
    zoom = PSICOSSOCIAL_RENDER_DPI / 72
    matrix = fitz.Matrix(zoom, zoom)

    try:
        for page_index in range(pdf.page_count):
            page = pdf[page_index]
            image_path = images_dir / f"pagina_{page_index + 1:03d}.jpg"
            pix = page.get_pixmap(matrix=matrix, alpha=False, colorspace=fitz.csRGB)
            pix.save(str(image_path), output="jpeg", jpg_quality=PSICOSSOCIAL_JPEG_QUALITY)

            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1
            run = paragraph.add_run()

            # Ajusta pelo tamanho útil da página. A pequena redução evita páginas
            # em branco causadas pelo parágrafo que o Word mantém após imagens.
            rect = page.rect
            max_width_in = (first_rect.width / 72) - (page_margin * 2)
            max_height_in = (first_rect.height / 72) - (page_margin * 2) - 0.03
            width_in = rect.width / 72 if rect.width else max_width_in
            height_in = rect.height / 72 if rect.height else max_height_in
            scale = min(max_width_in / width_in, max_height_in / height_in, 1)
            run.add_picture(str(image_path), width=Inches(width_in * scale))

    finally:
        pdf.close()

    doc.save(str(output_docx))
    return output_docx


def _merge_docx_files(paths: list[Path], output_path: Path) -> Path:
    if not paths:
        raise ValueError("Envie pelo menos um arquivo Word para juntar.")
    try:
        from docx import Document
        from docxcompose.composer import Composer
    except Exception as exc:  # pragma: no cover
        raise RuntimeError("As bibliotecas docxcompose/python-docx não estão instaladas. Verifique o requirements.txt no Render.") from exc

    master = Document(str(paths[0]))
    composer = Composer(master)
    for path in paths[1:]:
        # Garante que cada anexo comece em uma nova página.
        composer.doc.add_page_break()
        composer.append(Document(str(path)))
    output_path.parent.mkdir(parents=True, exist_ok=True)
    composer.save(str(output_path))
    return output_path


def _save_uploaded_file(file_storage, folder: Path, allowed_exts: set[str], label: str) -> Path:
    if not file_storage or not getattr(file_storage, "filename", ""):
        raise ValueError(f"Envie o arquivo: {label}.")
    filename = secure_filename(file_storage.filename)
    ext = Path(filename).suffix.lower()
    if ext not in allowed_exts:
        raise ValueError(f"O arquivo {label} precisa estar em: {', '.join(sorted(allowed_exts))}.")
    folder.mkdir(parents=True, exist_ok=True)
    path = folder / f"{uuid.uuid4().hex}_{filename}"
    file_storage.save(path)
    return path


def _build_combined_pgr_aet_psychosocial(pgr_docx: Path, aet_docx: Path, psicossocial_pdf: Path, output_path: Path, empresa: str, data_criacao: str, mes_extenso: str | None = None) -> Path:
    workdir = output_path.parent / f"merge_{uuid.uuid4().hex}"
    workdir.mkdir(parents=True, exist_ok=True)
    try:
        link_pgr = workdir / "link_pgr_para_aet.docx"
        link_aet = workdir / "link_aet_para_psicossocial.docx"
        merged_without_psy = workdir / "pgr_aet_links.docx"

        _prepare_link_docx(LINK_PGR_AET_TEMPLATE, link_pgr, empresa, data_criacao, mes_extenso)
        _prepare_link_docx(LINK_AET_PSICOSSOCIAL_TEMPLATE, link_aet, empresa, data_criacao, mes_extenso)

        # Primeiro junta os arquivos editáveis. Depois adiciona o Relatório
        # Psicossocial como páginas-imagem em uma nova seção sem cabeçalho/rodapé.
        # Isso evita as distorções geradas por conversão PDF -> DOCX por texto.
        _merge_docx_files([pgr_docx, link_pgr, aet_docx, link_aet], merged_without_psy)
        return _append_pdf_as_images_to_docx(merged_without_psy, psicossocial_pdf, output_path)
    finally:
        shutil.rmtree(workdir, ignore_errors=True)


def _send_docx_with_cleanup(path: Path, download_name: str):
    """Envia o DOCX e limpa o arquivo temporário quando a resposta terminar.

    Evita acúmulo de arquivos grandes no disco do Render após juntar documentos.
    """
    response = send_file(
        path,
        as_attachment=True,
        download_name=download_name,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    @response.call_on_close
    def _cleanup() -> None:
        try:
            path.unlink(missing_ok=True)
        except Exception:
            pass

    return response


def _is_ajax_request() -> bool:
    return request.headers.get("X-Requested-With") == "XMLHttpRequest"


def _ajax_error(message: str, status_code: int = 400):
    return jsonify({"ok": False, "error": message}), status_code


def _gerar_context(form_state: dict[str, Any] | None = None) -> dict[str, Any]:
    today = datetime.now().strftime("%m/%Y")
    next_year = datetime.now().replace(year=datetime.now().year + 1).strftime("%m/%Y")
    risks = _sorted_risks()
    ltcat_risks = [risk for risk in risks if str(risk.get("tipo_risco", "")).strip().upper() in TIPOS_RISCO_LTCAT]
    return {
        "risks": risks,
        "ltcat_risks": ltcat_risks,
        "risk_groups": _sorted_risk_groups(),
        "sectors": _sorted_sectors(),
        "exams": _sorted_exams(),
        "groups": _sorted_groups(),
        "grouped_sectors": _grouped_sectors(),
        "companies": _sorted_companies(),
        "aet_presets": AET_CNAE_PRESETS,
        "report_profiles": _sorted_report_profiles(),
        "imported_templates": _sorted_imported_laudo_templates(),
        "options": FORM_OPTIONS,
        "today": today,
        "next_year": next_year,
        "form_state": form_state or {},
    }


def _render_gerar_with_current_form():
    return render_template("gerar.html", **_gerar_context(_gerar_form_state_from_request()))


def _validate_complete_report_fields(company: dict[str, str], label: str) -> list[str]:
    errors: list[str] = []
    empresa = company.get("empresa") or company.get("nome", "")
    if not empresa:
        errors.append(f"Selecione ou cadastre uma empresa para gerar o {label} completo.")
    if not company.get("cnpj"):
        errors.append(f"Cadastre o CNPJ da empresa para gerar o {label} completo.")
    if not company.get("data_criacao_laudo"):
        errors.append("Preencha a Data de criação do laudo para finalizar o documento.")
    if company.get("ajuste_psicossocial") == "1" and not company.get("data_da_revisao"):
        errors.append("Preencha a Data da revisão psicossocial quando marcar que o laudo é apenas um ajuste.")
    return errors


@app.route("/modelos-aet")
def modelos_aet():
    """Catálogo de CNAEs e modelos técnicos padrão para AET.

    A tela serve como consulta para o usuário entender quais respostas padrão
    serão sugeridas quando clicar em Aplicar sugestões do CNAE na geração.
    """
    categorias: dict[str, list[dict[str, Any]]] = {}
    for item in AET_CNAE_PRESETS:
        categorias.setdefault(item.get("categoria", "Outros"), []).append(item)
    return render_template("modelos_aet.html", categorias=categorias, catalog=AET_CNAE_CATALOG)


@app.route("/gerar")
def gerar():
    _cleanup_expired_temp_import_groups()
    profile_id = request.args.get("profile_id", "").strip()
    if profile_id:
        profile = db.session.get(ReportProfile, profile_id)
        if profile:
            state = dict(profile.state or {})
            state["profile_id"] = profile.id
            state["profile_name"] = profile.nome
            return render_template("gerar.html", **_gerar_context(state))
    return render_template("gerar.html", **_gerar_context())




@app.post("/aplicar-modelo-importado")
def apply_imported_template():
    company_id = _field("company_id")
    template_id = _field("imported_template_id")
    group_id = _field("imported_template_group_id") or None
    if not company_id:
        flash("Selecione a empresa que receberá os dados do laudo antigo antes de aplicar o modelo.", "error")
        return _render_gerar_with_current_form()
    company = db.session.get(Company, company_id)
    if not company:
        flash("Empresa selecionada não foi encontrada.", "error")
        return _render_gerar_with_current_form()
    template = db.session.get(ImportedLaudoTemplate, template_id) if template_id else None
    if not template:
        flash("Selecione um modelo importado de laudo antigo para aplicar.", "error")
        return _render_gerar_with_current_form()
    temp_group = None
    if not group_id:
        # Quando o usuário não escolhe um grupo fixo, os setores importados vão para
        # um grupo temporário exclusivo dessa aplicação. Isso organiza a geração e
        # evita poluir os grupos permanentes de setores.
        temp_group = _create_temp_import_group(company, template)
        group_id = temp_group.id
    profile = _apply_imported_template_to_company(template, company_id, group_id=group_id)
    if temp_group:
        state = dict(profile.state or {})
        state["temporary_sector_group_id"] = temp_group.id
        state["temporary_sector_group_name"] = temp_group.nome
        profile.state = state
        db.session.commit()
        flash("Modelo importado aplicado em grupo temporário. Ele fica disponível para esta geração e será removido automaticamente depois.", "success")
    else:
        flash("Modelo importado aplicado à empresa. Os setores, cargos, riscos e exames foram carregados para revisão antes de gerar os laudos.", "success")
    return redirect(url_for("gerar", profile_id=profile.id))


@app.post("/salvar-configuracao-laudo")
def save_report_profile():
    profile = _save_report_profile_from_form(auto=False)
    if profile:
        flash("Configuração da empresa salva. Você poderá carregar essa seleção novamente depois.", "success")
    else:
        flash("Selecione uma empresa antes de salvar a configuração da geração.", "error")
    return _render_gerar_with_current_form()


@app.post("/configuracao-laudo/<profile_id>/excluir")
def delete_report_profile(profile_id: str):
    profile = db.session.get(ReportProfile, profile_id)
    if profile:
        db.session.delete(profile)
        db.session.commit()
        flash("Configuração salva excluída.", "success")
    return redirect(url_for("gerar"))


@app.post("/juntar-arquivos")
def merge_files_avulso():
    try:
        with tempfile.TemporaryDirectory() as tmp:
            tmpdir = Path(tmp)
            pgr_docx = _save_uploaded_file(request.files.get("pgr_docx"), tmpdir, {".docx"}, "PGR em Word (.docx)")
            aet_docx = _save_uploaded_file(request.files.get("aet_docx"), tmpdir, {".docx"}, "AET em Word (.docx)")
            psicossocial_pdf = _save_uploaded_file(request.files.get("psicossocial_pdf"), tmpdir, {".pdf"}, "Relatório Psicossocial em PDF")

            company = None
            company_id = _field("company_id")
            if company_id:
                company = db.session.get(Company, company_id)

            empresa = _field("empresa_avulsa")
            if company and not empresa:
                empresa = company.nome or ""
            empresa = empresa or "EMPRESA"
            data_criacao = _field("data_criacao_laudo") or datetime.now().strftime("%d/%m/%Y")
            mes_extenso = _field("mes_extenso")

            OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
            stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_empresa = _normalize_filename(empresa)
            output_path = OUTPUT_DIR / f"pgr_aet_psicossocial_{safe_empresa}_{stamp}.docx"

            _build_combined_pgr_aet_psychosocial(
                pgr_docx,
                aet_docx,
                psicossocial_pdf,
                output_path,
                empresa,
                data_criacao,
                mes_extenso,
            )
            return _send_docx_with_cleanup(output_path, f"PGR_AET_RELATORIO_PSICOSSOCIAL_{safe_empresa}.docx")
    except Exception as exc:
        message = f"Erro ao juntar arquivos: {exc}"
        if _is_ajax_request():
            return _ajax_error(message)
        flash(message, "error")
        return redirect(url_for("juntar_arquivos"))


@app.route("/juntar-arquivos")
def juntar_arquivos():
    return render_template("juntar.html", companies=_sorted_companies(), today=datetime.now().strftime("%d/%m/%Y"))



@app.route("/recibos-esocial", methods=["GET", "POST"])
def recibos_esocial():
    if request.method == "POST":
        try:
            with tempfile.TemporaryDirectory() as tmp:
                tmpdir = Path(tmp)
                planilha = _save_uploaded_file(
                    request.files.get("planilha"),
                    tmpdir,
                    {".xls", ".xlsx"},
                    "Planilha de recibos RELFUNCGERAL (.xls ou .xlsx)",
                )
                rows = _read_receipt_spreadsheet(planilha)
                empresa = rows[0].get("empresa", "RECIBOS") if rows else "RECIBOS"
                safe_empresa = _normalize_filename(empresa)
                stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
                output_path = OUTPUT_DIR / f"recibos_esocial_{safe_empresa}_{stamp}.pdf"
                _build_receipts_pdf(rows, output_path)
                return _send_pdf_with_cleanup(output_path, f"RECIBO_{safe_empresa}.pdf")
        except Exception as exc:
            flash(f"Erro ao formatar recibos: {exc}", "error")
    return render_template("recibos_esocial.html")


@app.route("/importar-laudo-antigo", methods=["GET", "POST"])
def importar_laudo_antigo():
    extracted = None
    draft_id = request.args.get("draft_id", "").strip()
    draft = _get_import_draft(draft_id)
    if draft:
        extracted = draft.state or {}
    if request.method == "POST":
        try:
            file_storage = request.files.get("arquivo")
            text_value = _extract_text_from_upload(file_storage)
            extracted = _smart_extract_laudo_data(text_value)
            draft = _create_import_draft(extracted)
            flash("Leitura concluída. Revise os dados extraídos antes de salvar no sistema.", "success")
            return redirect(url_for("importar_laudo_antigo", draft_id=draft.id))
        except Exception as exc:
            flash(f"Erro ao importar laudo antigo: {exc}", "error")
    return render_template(
        "importar_laudo.html",
        extracted=extracted or {},
        draft_id=draft.id if draft else "",
        groups=_sorted_groups(),
        risks=_sorted_risks(),
        exams=_sorted_exams(),
        imported_templates=_sorted_imported_laudo_templates(),
    )


@app.post("/importar-laudo-antigo/salvar")
def salvar_importacao_laudo_antigo():
    draft = _get_import_draft(_field("draft_id"))
    draft_state = dict(draft.state or {}) if draft else {}

    empresa = _field("empresa") or draft_state.get("empresa", "")
    cnpj = _field("cnpj") or draft_state.get("cnpj", "")
    setores_text = _field("setores_extraidos")
    riscos_text = _field("riscos_extraidos")
    exames_text = _field("exames_extraidos")
    group_id = _field("grupo_id") or None

    company_fields = {
        "endereco": _field("endereco") or draft_state.get("endereco", ""),
        "bairro_cidade": _field("bairro_cidade") or draft_state.get("bairro_cidade", ""),
        "cep": _field("cep") or draft_state.get("cep", ""),
        "cnae1": _field("cnae1") or draft_state.get("cnae1", ""),
        "descricao1": _field("descricao1") or draft_state.get("descricao1", ""),
        "grau1": _field("grau1") or draft_state.get("grau1", ""),
        "cnae2": _field("cnae2") or draft_state.get("cnae2", ""),
        "descricao2": _field("descricao2") or draft_state.get("descricao2", ""),
        "grau2": _field("grau2") or draft_state.get("grau2", ""),
        "funcionarios": _field("funcionarios") or draft_state.get("funcionarios", ""),
        "data_atual": _field("data_atual") or draft_state.get("data_atual", ""),
        "data_final": _field("data_final") or draft_state.get("data_final", ""),
        "email": _field("email") or draft_state.get("email", ""),
        "fone": _field("fone") or draft_state.get("fone", ""),
    }

    setores_json = draft_state.get("setores_cargos") or []
    riscos_json = draft_state.get("riscos_detalhados") or []
    sector_risks_json = draft_state.get("sector_risks") or {}
    if not isinstance(setores_json, list):
        setores_json = []
    if not isinstance(riscos_json, list):
        riscos_json = []
    if not isinstance(sector_risks_json, dict):
        sector_risks_json = {}

    setores_linhas = _unique_clean_lines((setores_text or "\n".join(draft_state.get("setores") or [])).splitlines())
    riscos_linhas = _unique_clean_lines((riscos_text or "\n".join(draft_state.get("riscos") or [])).splitlines())
    exames_linhas = _unique_clean_lines((exames_text or "\n".join(draft_state.get("exames") or [])).splitlines())

    # Se o usuário adicionou manualmente nomes no textarea, preserva como setor/risco simples.
    existing_sector_keys = {_simple_norm(item.get("setor", "")) for item in setores_json if isinstance(item, dict)}
    for setor_nome in setores_linhas:
        if _simple_norm(setor_nome) and _simple_norm(setor_nome) not in existing_sector_keys:
            setores_json.append({"setor": setor_nome.upper(), "cargos": []})
            existing_sector_keys.add(_simple_norm(setor_nome))

    existing_risk_keys = {_simple_norm(item.get("risco", "")) for item in riscos_json if isinstance(item, dict)}
    for risco_nome in riscos_linhas:
        if _simple_norm(risco_nome) and _simple_norm(risco_nome) not in existing_risk_keys:
            riscos_json.append({"risco": risco_nome})
            existing_risk_keys.add(_simple_norm(risco_nome))

    if not empresa and not cnpj and not setores_json and not riscos_json and not sector_risks_json and not exames_linhas:
        flash("Nenhum dado foi informado para salvar.", "error")
        return redirect(url_for("importar_laudo_antigo"))

    # Salva/atualiza a empresa de origem apenas para consulta; o modelo é reutilizável em qualquer empresa.
    if empresa or cnpj:
        company = None
        if cnpj:
            company = Company.query.filter(db.func.lower(Company.cnpj) == cnpj.lower()).first()
        if not company and empresa:
            company = Company.query.filter(db.func.lower(Company.nome) == empresa.lower()).first()
        if not company:
            company = Company(nome=empresa or "Empresa importada", cnpj=cnpj or "")
            db.session.add(company)
        else:
            if empresa:
                company.nome = empresa
            if cnpj:
                company.cnpj = cnpj
        for key, value in company_fields.items():
            if value:
                setattr(company, key, value)

    template_name = _field("template_name") or f"Modelo importado - {empresa or cnpj or datetime.now().strftime('%d/%m/%Y %H:%M')}"
    template_state = {
        "empresa": empresa,
        "cnpj": cnpj,
        **company_fields,
        "setores": setores_linhas,
        "setores_cargos": setores_json,
        "riscos": riscos_linhas,
        "riscos_detalhados": riscos_json,
        "sector_risks": sector_risks_json,
        "exames": exames_linhas,
        "grupo_id_preferencial": group_id or "",
    }
    template = ImportedLaudoTemplate(
        nome=template_name,
        source_company=empresa or "",
        source_cnpj=cnpj or "",
        state=template_state,
    )
    db.session.add(template)
    if draft:
        db.session.delete(draft)
    db.session.commit()
    sector_total = len(template_state["setores_cargos"] or template_state["setores"])
    risk_total = len(template_state["riscos_detalhados"] or template_state["riscos"])
    linked_total = sum(len(v) for v in (template_state.get("sector_risks") or {}).values() if isinstance(v, list))
    exam_total = len(template_state["exames"])
    flash(
        f"Modelo reutilizável salvo: {template.nome}. Guardado no banco: {sector_total} setor(es)/grupo(s), {risk_total} risco(s), {linked_total} vínculo(s) risco/setor e {exam_total} exame(s). Agora aplique em qualquer empresa pela aba Gerar laudos.",
        "success",
    )
    return redirect(url_for("importar_laudo_antigo"))


@app.post("/modelos-importados/<template_id>/excluir")
def delete_imported_template(template_id: str):
    template = db.session.get(ImportedLaudoTemplate, template_id)
    if template:
        db.session.delete(template)
        db.session.commit()
        flash("Modelo importado excluído. Os cadastros manuais de riscos, setores e exames não foram apagados.", "success")
    else:
        flash("Modelo importado não encontrado.", "error")
    return redirect(url_for("importar_laudo_antigo"))


@app.post("/modelos-importados/excluir-todos")
def delete_all_imported_templates():
    count = ImportedLaudoTemplate.query.delete()
    ImportedLaudoDraft.query.delete()
    db.session.commit()
    flash(f"{count} modelo(s) importado(s) excluído(s). Os cadastros manuais foram preservados.", "success")
    return redirect(url_for("importar_laudo_antigo"))


@app.route("/grupos-riscos")
def risk_groups():
    return render_template("risk_groups.html", risk_groups=_sorted_risk_groups(), risks=_sorted_risks())


def _risk_group_form_data(existing_id: str | None = None) -> dict[str, Any]:
    return {
        "id": existing_id or uuid.uuid4().hex,
        "nome": _field("nome"),
        "descricao": _field("descricao"),
        "risk_ids": request.form.getlist("risk_ids"),
    }


def _validate_risk_group(data: dict[str, Any], existing_id: str | None = None) -> list[str]:
    errors: list[str] = []
    if not data.get("nome"):
        errors.append("Preencha o nome do grupo de riscos.")
    if not data.get("risk_ids"):
        errors.append("Selecione pelo menos um risco para o grupo.")
    if data.get("nome"):
        query = RiskGroup.query.filter(db.func.lower(RiskGroup.nome) == data["nome"].strip().lower())
        if existing_id:
            query = query.filter(RiskGroup.id != existing_id)
        if query.first():
            errors.append("Já existe um grupo de riscos com esse nome.")
    return errors


def _apply_risk_group_data(model: RiskGroup, data: dict[str, Any]) -> RiskGroup:
    model.nome = data["nome"].strip()
    model.descricao = data.get("descricao", "").strip()
    risk_ids = _dedupe_preserve_order(data.get("risk_ids", []))
    model.risks = Risk.query.filter(Risk.id.in_(risk_ids)).all() if risk_ids else []
    model.updated_at = datetime.utcnow()
    return model


@app.post("/grupo-risco/novo")
def create_risk_group():
    data = _risk_group_form_data()
    errors = _validate_risk_group(data)
    if errors:
        for error in errors:
            flash(error, "error")
        return redirect(url_for("risk_groups"))
    model = RiskGroup(id=data["id"])
    _apply_risk_group_data(model, data)
    db.session.add(model)
    db.session.commit()
    flash("Grupo de riscos cadastrado com sucesso.", "success")
    return redirect(url_for("risk_groups"))


@app.route("/grupo-risco/<group_id>/editar", methods=["GET", "POST"])
def edit_risk_group(group_id: str):
    model = db.session.get(RiskGroup, group_id)
    if not model:
        flash("Grupo de riscos não encontrado.", "error")
        return redirect(url_for("risk_groups"))
    if request.method == "POST":
        data = _risk_group_form_data(existing_id=group_id)
        errors = _validate_risk_group(data, existing_id=group_id)
        if errors:
            for error in errors:
                flash(error, "error")
            return render_template("risk_group_form.html", group={**model.to_dict(), **data}, risks=_sorted_risks(), title="Editar grupo de riscos")
        _apply_risk_group_data(model, data)
        db.session.commit()
        flash("Grupo de riscos atualizado com sucesso.", "success")
        return redirect(url_for("risk_groups"))
    return render_template("risk_group_form.html", group=model.to_dict(), risks=_sorted_risks(), title="Editar grupo de riscos")


@app.post("/grupo-risco/<group_id>/excluir")
def delete_risk_group(group_id: str):
    model = db.session.get(RiskGroup, group_id)
    if model:
        db.session.delete(model)
        db.session.commit()
        flash("Grupo de riscos excluído.", "success")
    return redirect(url_for("risk_groups"))


@app.post("/api/risco/novo")
def api_create_risk():
    """Cadastro rápido usado na tela Gerar laudos, sem recarregar a página."""
    new_risk = _form_to_risk()
    errors = _validate_risk(new_risk)

    if new_risk.get("risco"):
        existing = Risk.query.filter(db.func.lower(Risk.risco) == new_risk["risco"].strip().lower()).first()
        if existing:
            errors.append("Já existe um risco cadastrado com esse nome.")

    if errors:
        return jsonify({"ok": False, "errors": errors}), 400

    risk_model = _risk_from_dict(new_risk)
    db.session.add(risk_model)
    db.session.commit()
    return jsonify({"ok": True, "message": "Risco cadastrado com sucesso.", "risk": risk_model.to_dict()})


@app.post("/risco/novo")
def create_risk():
    new_risk = _form_to_risk()
    errors = _validate_risk(new_risk)
    if errors:
        for error in errors:
            flash(error, "error")
        return redirect(url_for("cadastro"))
    db.session.add(_risk_from_dict(new_risk))
    db.session.commit()
    flash("Risco cadastrado com sucesso.", "success")
    return redirect(url_for("cadastro"))


@app.route("/risco/<risk_id>/editar", methods=["GET", "POST"])
def edit_risk(risk_id: str):
    risk_model = db.session.get(Risk, risk_id)
    if not risk_model:
        flash("Risco não encontrado.", "error")
        return redirect(url_for("cadastro"))

    if request.method == "POST":
        updated = _form_to_risk(existing_id=risk_id)
        errors = _validate_risk(updated)
        if errors:
            for error in errors:
                flash(error, "error")
            return render_template("edit.html", risk={**risk_model.to_dict(), **updated}, options=FORM_OPTIONS)
        _risk_from_dict(updated, risk_model)
        db.session.commit()
        flash("Risco atualizado com sucesso.", "success")
        return redirect(url_for("cadastro"))

    return render_template("edit.html", risk=risk_model.to_dict(), options=FORM_OPTIONS)


@app.post("/risco/<risk_id>/excluir")
def delete_risk(risk_id: str):
    risk_model = db.session.get(Risk, risk_id)
    if risk_model:
        for group in RiskGroup.query.all():
            if risk_model in (group.risks or []):
                group.risks.remove(risk_model)
        db.session.delete(risk_model)
        db.session.commit()
        flash("Risco excluído.", "success")
    return redirect(url_for("cadastro"))


@app.post("/setor/novo")
def create_sector():
    new_sector = _form_to_sector()
    errors = _validate_sector(new_sector)
    if errors:
        for error in errors:
            flash(error, "error")
        return redirect(url_for("setores"))
    db.session.add(_sector_from_dict(new_sector))
    db.session.commit()
    flash("Setor e cargos cadastrados com sucesso.", "success")
    return redirect(url_for("setores"))


@app.route("/setor/<sector_id>/editar", methods=["GET", "POST"])
def edit_sector(sector_id: str):
    sector_model = db.session.get(Sector, sector_id)
    if not sector_model:
        flash("Setor não encontrado.", "error")
        return redirect(url_for("setores"))

    if request.method == "POST":
        updated = _form_to_sector(existing_id=sector_id)
        errors = _validate_sector(updated)
        if errors:
            for error in errors:
                flash(error, "error")
            return render_template("edit_setor.html", sector={**sector_model.to_dict(), **updated}, groups=_sorted_groups())
        _sector_from_dict(updated, sector_model)
        db.session.commit()
        flash("Setor atualizado com sucesso.", "success")
        return redirect(url_for("setores"))

    return render_template("edit_setor.html", sector=sector_model.to_dict(), groups=_sorted_groups())


@app.post("/setor/<sector_id>/excluir")
def delete_sector(sector_id: str):
    sector_model = db.session.get(Sector, sector_id)
    if sector_model:
        db.session.delete(sector_model)
        db.session.commit()
        flash("Setor excluído.", "success")
    return redirect(url_for("setores"))


@app.post("/setores/apagar-todos")
def delete_all_sectors():
    count = Sector.query.count()
    Sector.query.delete()
    db.session.commit()
    flash(f"{count} setor(es) e seus cargos foram apagados. Os riscos cadastrados foram mantidos.", "success")
    return redirect(url_for("setores"))


def _selected_sector_ltcat_groups() -> tuple[list[dict[str, Any]], list[str]]:
    selected_sector_ids = request.form.getlist("pgr_sector_ids")
    risks = {risk.id: risk.to_dict() for risk in Risk.query.all()}
    sectors = {sector.id: sector.to_dict() for sector in Sector.query.all()}

    groups: list[dict[str, Any]] = []
    errors: list[str] = []

    if not selected_sector_ids:
        errors.append("Selecione pelo menos um setor para gerar o LTCAT.")
        return groups, errors

    for sector_id in selected_sector_ids:
        sector = sectors.get(sector_id)
        if not sector:
            continue
        # A LTCAT volta a ser automática: todos os riscos ambientais
        # selecionados no bloco "Riscos por setor" entram no laudo.
        # Tipos não ambientais são ignorados, e setor sem risco ambiental
        # gera o bloco de AUSÊNCIA DE RISCOS.
        risk_ids = request.form.getlist(f"sector_risk_ids_{sector_id}")
        group_ids = request.form.getlist(f"sector_risk_group_ids_{sector_id}")
        risk_ids = _dedupe_preserve_order(risk_ids + _risk_ids_from_group_ids(group_ids))
        selected_risks = [risks[risk_id] for risk_id in risk_ids if risk_id in risks]
        selected_risks = [
            risk for risk in selected_risks
            if str(risk.get("tipo_risco", "")).strip().upper() in TIPOS_RISCO_LTCAT
        ]
        groups.append({"sector": sector, "risks": selected_risks, "exams": []})

    return groups, errors


def _company_extra_from_form() -> dict[str, str]:
    return _company_payload_from_form()


def _send_generated_docx(generator, selected: list[dict[str, Any]], stem: str, download_name: str, *args):
    if not selected:
        flash("Selecione pelo menos um item para gerar o arquivo Word.", "error")
        return _render_gerar_with_current_form()

    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    output_path = OUTPUT_DIR / f"{stem}_{stamp}.docx"
    generator(selected, output_path, *args)
    return send_file(
        output_path,
        as_attachment=True,
        download_name=download_name,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )




@app.post("/api/sugerir-exames")
def api_sugerir_exames():
    groups, errors = _selected_sector_risk_groups()
    # Para sugestão, não bloqueia setor sem risco; só retorna avisos.
    suggestions = _exam_rule_suggestions_for_groups(groups)
    return jsonify({"ok": True, "suggestions": suggestions, "warnings": errors, "rules": EXAM_RULES})


@app.post("/api/previsualizar-geracao")
def api_previsualizar_geracao():
    groups, errors = _selected_sector_risk_groups()
    company = _company_payload_from_form()
    preview = _build_generation_preview(groups, company)
    preview["erros"] = errors
    return jsonify({"ok": True, "preview": preview})


@app.post("/gerar-aet-completa")
def generate_complete_aet():
    groups, errors = _selected_sector_risk_groups()
    company = _company_payload_from_form()
    errors.extend(_validate_complete_report_fields(company, "AET"))
    if errors:
        for error in errors:
            flash(error, "error")
        return _render_gerar_with_current_form()
    try:
        empresa = company.get("empresa") or company.get("nome", "")
        cnpj = company.get("cnpj", "")
        data_atual = company.get("data_atual", "")
        data_final = company.get("data_final", "")
        _save_report_profile_from_form(auto=True)
        return _send_generated_docx(
            generate_aet_docx,
            groups,
            "aet_completa",
            "AET_COMPLETA.docx",
            empresa,
            cnpj,
            data_atual,
            data_final,
            company,
        )
    except Exception as exc:
        flash(f"Erro ao gerar AET completa: {exc}", "error")
        return _render_gerar_with_current_form()


@app.post("/gerar-pacote-empresa")
def generate_company_zip_package():
    groups, errors = _selected_sector_risk_groups()
    ltcat_groups, ltcat_errors = _selected_sector_ltcat_groups()
    company = _company_payload_from_form()
    errors.extend(ltcat_errors)
    errors.extend(_validate_complete_report_fields(company, "pacote ZIP"))
    if errors:
        for error in errors:
            flash(error, "error")
        return _render_gerar_with_current_form()
    try:
        empresa = company.get("empresa") or company.get("nome", "")
        cnpj = company.get("cnpj", "")
        data_atual = company.get("data_atual", "")
        data_final = company.get("data_final", "")
        safe_empresa = _normalize_filename(empresa)
        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        zip_path = OUTPUT_DIR / f"pacote_{safe_empresa}_{stamp}.zip"
        with tempfile.TemporaryDirectory() as tmp:
            tmpdir = Path(tmp)
            pgr_path = tmpdir / "PGR_COMPLETO.docx"
            pcmso_path = tmpdir / "PCMSO_COMPLETO.docx"
            ltcat_path = tmpdir / "LTCAT_COMPLETO.docx"
            aet_path = tmpdir / "AET_COMPLETA.docx"
            generate_complete_pgr_docx(groups, pgr_path, empresa, cnpj, data_atual, data_final, company)
            generate_complete_pcmso_docx(groups, pcmso_path, empresa, cnpj, data_atual, data_final, company)
            generate_complete_ltcat_docx(ltcat_groups, ltcat_path, empresa, cnpj, data_atual, data_final, company)
            generate_aet_docx(groups, aet_path, empresa, cnpj, data_atual, data_final, company)
            with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
                for path in [pgr_path, pcmso_path, ltcat_path, aet_path]:
                    zf.write(path, arcname=path.name)
                resumo = tmpdir / "RESUMO_DA_GERACAO.txt"
                preview = _build_generation_preview(groups, company)
                resumo.write_text(
                    f"Empresa: {empresa}\nCNPJ: {cnpj}\nVigência: {data_atual} a {data_final}\n"
                    f"Setores: {preview['setores']}\nRiscos: {preview['riscos']}\nExames: {preview['exames']}\n"
                    f"Riscos psicossociais: {preview['psicossociais']}\nRiscos ambientais para LTCAT: {preview['ambientais']}\n"
                    f"Avisos:\n- " + "\n- ".join(preview.get('avisos') or ['Nenhum aviso.']),
                    encoding="utf-8"
                )
                zf.write(resumo, arcname="RESUMO_DA_GERACAO.txt")
        _save_report_profile_from_form(auto=True)
        return _send_zip_with_cleanup(zip_path, f"PACOTE_LAUDOS_{safe_empresa}.zip")
    except Exception as exc:
        flash(f"Erro ao gerar pacote ZIP da empresa: {exc}", "error")
        return _render_gerar_with_current_form()


@app.post("/gerar-pgr-aet-psicossocial")
def generate_pgr_aet_psychosocial():
    groups, errors = _selected_sector_risk_groups()
    company = _company_payload_from_form()
    errors.extend(_validate_complete_report_fields(company, "PGR + AET + Relatório Psicossocial"))
    if errors:
        if _is_ajax_request():
            return _ajax_error("\n".join(errors))
        for error in errors:
            flash(error, "error")
        return _render_gerar_with_current_form()
    try:
        with tempfile.TemporaryDirectory() as tmp:
            tmpdir = Path(tmp)
            psicossocial_pdf = _save_uploaded_file(request.files.get("psicossocial_pdf"), tmpdir, {".pdf"}, "Relatório Psicossocial em PDF")
            empresa = company.get("empresa") or company.get("nome", "")
            cnpj = company.get("cnpj", "")
            data_atual = company.get("data_atual", "")
            data_final = company.get("data_final", "")
            pgr_docx = tmpdir / "pgr_gerado.docx"
            aet_upload = request.files.get("aet_docx")
            if aet_upload and getattr(aet_upload, "filename", ""):
                aet_docx = _save_uploaded_file(aet_upload, tmpdir, {".docx"}, "AET em Word (.docx)")
            else:
                aet_docx = tmpdir / "aet_gerada.docx"
                generate_aet_docx(groups, aet_docx, empresa, cnpj, data_atual, data_final, company)
            generate_complete_pgr_docx(groups, pgr_docx, empresa, cnpj, data_atual, data_final, company)
            _save_report_profile_from_form(auto=True)
            stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
            output_path = OUTPUT_DIR / f"pgr_aet_psicossocial_{stamp}.docx"
            _build_combined_pgr_aet_psychosocial(
                pgr_docx,
                aet_docx,
                psicossocial_pdf,
                output_path,
                empresa,
                company.get("data_criacao_laudo") or company.get("data_avaliacao") or company.get("data_atual", ""),
                _field("mes_extenso"),
            )
            return _send_docx_with_cleanup(output_path, "PGR_AET_RELATORIO_PSICOSSOCIAL.docx")
    except Exception as exc:
        message = f"Erro ao gerar PGR + AET + Relatório Psicossocial: {exc}"
        if _is_ajax_request():
            return _ajax_error(message)
        flash(message, "error")
        return _render_gerar_with_current_form()

@app.post("/gerar-pgr-completo")
def generate_complete_pgr():
    groups, errors = _selected_sector_risk_groups()
    company = _company_payload_from_form()
    errors.extend(_validate_complete_report_fields(company, "PGR"))
    if errors:
        for error in errors:
            flash(error, "error")
        return _render_gerar_with_current_form()
    try:
        empresa = company.get("empresa") or company.get("nome", "")
        cnpj = company.get("cnpj", "")
        data_atual = company.get("data_atual", "")
        data_final = company.get("data_final", "")
        _save_report_profile_from_form(auto=True)
        return _send_generated_docx(
            generate_complete_pgr_docx,
            groups,
            "pgr_completo",
            "PGR_COMPLETO.docx",
            empresa,
            cnpj,
            data_atual,
            data_final,
            company,
        )
    except Exception as exc:
        flash(f"Erro ao gerar PGR completo: {exc}", "error")
        return _render_gerar_with_current_form()


@app.post("/gerar-plano-acao")
def generate_action_plan():
    try:
        if request.form.getlist("pgr_sector_ids"):
            groups, errors = _selected_sector_risk_groups()
            if errors:
                for error in errors:
                    flash(error, "error")
                return redirect(url_for("gerar"))
            selected = groups
        else:
            selected = _selected_risks()
        return _send_generated_docx(
            generate_action_plan_docx,
            selected,
            "plano_de_acao_riscos",
            "PLANO_DE_ACAO_RISCOS.docx",
            _field("data_atual"),
            _field("data_final"),
        )
    except Exception as exc:
        flash(f"Erro ao gerar Plano de Ação: {exc}", "error")
        return redirect(url_for("gerar"))


@app.post("/gerar-risco-pgr")
def generate_risco_pgr():
    groups, errors = _selected_sector_risk_groups()
    if errors:
        for error in errors:
            flash(error, "error")
        return redirect(url_for("gerar"))
    try:
        return _send_generated_docx(generate_pgr_docx, groups, "risco_pgr", "RISCO_PGR.docx")
    except Exception as exc:
        flash(f"Erro ao gerar Risco PGR: {exc}", "error")
        return redirect(url_for("gerar"))


@app.post("/gerar-riscos-pcmso")
def generate_riscos_pcmso():
    groups, errors = _selected_sector_risk_groups()
    if errors:
        for error in errors:
            flash(error, "error")
        return redirect(url_for("gerar"))
    try:
        return _send_generated_docx(generate_riscos_pcmso_docx, groups, "riscos_pcmso", "RISCOS_PCMSO.docx")
    except Exception as exc:
        flash(f"Erro ao gerar riscos/exames do PCMSO: {exc}", "error")
        return redirect(url_for("gerar"))


@app.post("/gerar-ltcat-completo")
def generate_complete_ltcat():
    groups, errors = _selected_sector_ltcat_groups()
    company = _company_payload_from_form()
    errors.extend(_validate_complete_report_fields(company, "LTCAT"))
    if errors:
        for error in errors:
            flash(error, "error")
        return _render_gerar_with_current_form()
    try:
        empresa = company.get("empresa") or company.get("nome", "")
        cnpj = company.get("cnpj", "")
        data_atual = company.get("data_atual", "")
        data_final = company.get("data_final", "")
        _save_report_profile_from_form(auto=True)
        return _send_generated_docx(
            generate_complete_ltcat_docx,
    generate_aet_docx,
            groups,
            "ltcat_completo",
            "LTCAT_COMPLETO.docx",
            empresa,
            cnpj,
            data_atual,
            data_final,
            company,
        )
    except Exception as exc:
        flash(f"Erro ao gerar LTCAT completo: {exc}", "error")
        return _render_gerar_with_current_form()


@app.post("/gerar-pcmso-completo")
def generate_complete_pcmso():
    groups, errors = _selected_sector_risk_groups()
    company = _company_payload_from_form()
    errors.extend(_validate_complete_report_fields(company, "PCMSO"))
    if errors:
        for error in errors:
            flash(error, "error")
        return _render_gerar_with_current_form()
    try:
        empresa = company.get("empresa") or company.get("nome", "")
        cnpj = company.get("cnpj", "")
        data_atual = company.get("data_atual", "")
        data_final = company.get("data_final", "")
        _save_report_profile_from_form(auto=True)
        return _send_generated_docx(
            generate_complete_pcmso_docx,
            groups,
            "pcmso_completo",
            "PCMSO_COMPLETO.docx",
            empresa,
            cnpj,
            data_atual,
            data_final,
            company,
        )
    except Exception as exc:
        flash(f"Erro ao gerar PCMSO completo: {exc}", "error")
        return _render_gerar_with_current_form()


@app.post("/gerar-pcmso")
def generate_pcmso():
    return generate_riscos_pcmso()


@app.post("/gerar-relacao-funcao-atividade")
def generate_relacao_funcao_atividade():
    try:
        return _send_generated_docx(
            generate_relacao_funcao_atividade_docx,
            _selected_sectors(),
            "relacao_funcao_atividade",
            "RELACAO_FUNCAO_X_ATIVIDADE.docx",
            _field("data_atual"),
            _field("data_final"),
        )
    except Exception as exc:
        flash(f"Erro ao gerar Relação Função x Atividade: {exc}", "error")
        return redirect(url_for("gerar"))


@app.post("/gerar-descritivo-setor")
def generate_descritivo_setor():
    try:
        return _send_generated_docx(generate_descritivo_setor_docx, _selected_sectors(), "descritivo_setor", "DESCRITIVO_SETOR.docx")
    except Exception as exc:
        flash(f"Erro ao gerar Descritivo Setor: {exc}", "error")
        return redirect(url_for("gerar"))


@app.post("/gerar-word")
def generate_word_legacy():
    return generate_action_plan()


with app.app_context():
    init_database()


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=int(os.environ.get("PORT", 5000)), debug=True)
